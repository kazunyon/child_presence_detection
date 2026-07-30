from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(".")
OUT = ROOT / "outputs"
FONT = "BIZ UDゴシック"
DATE = "2026年7月30日"
ACCENT = "0F766E"
TEXT = "0F172A"
MUTED = "64748B"
HEAD_FILL = "E0F2F1"
NOTE_FILL = "FEF3C7"


def font(run, size=9, bold=False, color=TEXT):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts or OxmlElement("w:rFonts")
    if rfonts.getparent() is None:
        rpr.append(rfonts)
    for k in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(k), FONT)


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd")) or OxmlElement("w:shd")
    if shd.getparent() is None:
        pr.append(shd)
    shd.set(qn("w:fill"), fill)


def ctext(cell, text, bold=False, fill=None, align=None, size=8.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade(cell, fill)


def table(doc, headers, rows, widths=None, size=8.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0]
    trpr = hdr._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)
    for i, h in enumerate(headers):
        ctext(hdr.cells[i], h, bold=True, fill=HEAD_FILL, align=WD_ALIGN_PARAGRAPH.CENTER, size=8.4)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            ctext(cells[i], v, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else None, size=size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    font(r, size=15 if level == 1 else 11.5, bold=True, color=ACCENT if level == 1 else "164E63")


def para(doc, text, fill=None):
    if fill:
        return table(doc, [text], [], [26.4], size=9)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    font(r, 9)
    return p


def setup(title, subtitle):
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.2)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.header_distance = Cm(0.7)
    sec.footer_distance = Cm(0.7)
    for s in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        st = doc.styles[s]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.styles["Normal"].font.size = Pt(9)
    h = sec.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(h.add_run(f"{title} / 1.0 / {DATE}"), 8, color=MUTED)
    f = sec.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(f.add_run("Confidential / SV - Mamoru Bus Design Documents"), 8, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(45)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(f"まもるバス\n{title}"), 24, True, ACCENT)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(s.add_run(subtitle), 12, True, "164E63")
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.paragraph_format.space_before = Pt(16)
    font(m.add_run(f"作成日：{DATE}　版数：1.0　確認元：README.md / backend/main.py / src/App.tsx / render.yaml"), 9)
    doc.add_page_break()
    heading(doc, "改訂履歴")
    table(doc, ["版", "日付", "区分", "内容", "作成"], [["1.0", DATE, "新規作成", f"{title}を新規作成。確定、未実装、要確認を区別して記載。", "Codex"]], [1.4, 3.0, 2.2, 16.0, 3.0])
    return doc


def routes():
    lines = (ROOT / "backend/main.py").read_text(encoding="utf-8").splitlines()
    out, decs = [], []
    for line in lines:
        s = line.strip()
        if s.startswith("@app."):
            decs.append(s)
            continue
        m = re.match(r"(async def|def)\s+(\w+)\(", s)
        if m and decs:
            for d in decs:
                dm = re.match(r"@app\.(get|post|put|delete|patch)\(\"([^\"]+)\"(?:,\s*status_code=([^\)]+))?", d)
                if dm:
                    auth = "JWT"
                    if dm.group(2) in ("/health", "/api/auth/login"):
                        auth = "不要"
                    if "line/webhook" in dm.group(2):
                        auth = "LINE署名"
                    if "admin-recovery" in dm.group(2):
                        auth = "復旧トークン"
                    out.append([dm.group(1).upper(), dm.group(2), m.group(2), auth])
            decs = []
    return out


def db_tables():
    text = (ROOT / "backend/main.py").read_text(encoding="utf-8")
    items = []
    for cm in re.finditer(r'class\s+(\w+)\(Base\):(?P<body>.*?)(?=\nclass\s+\w+\(|\n@app\.|\Z)', text, re.S):
        body = cm.group("body")
        tn = re.search(r'__tablename__\s*=\s*"([^"]+)"', body)
        if not tn:
            continue
        cols = []
        uniq = ", ".join(re.findall(r"UniqueConstraint\((.*?)\)", body, re.S)).replace("\n", " ") or "なし"
        for line in body.splitlines():
            mm = re.search(r'(\w+):\s*Mapped\[(.*?)\]\s*=\s*mapped_column\((.*)\)', line.strip())
            if not mm:
                continue
            name, py, d = mm.group(1), mm.group(2), mm.group(3)
            sql = "INTEGER" if "int" in py else "BOOLEAN" if "bool" in py else "DATETIME" if "datetime" in py else "TEXT" if "Text" in d else "VARCHAR"
            sm = re.search(r"String\((\d+)\)", d)
            if sm:
                sql = f"VARCHAR({sm.group(1)})"
            fk = re.search(r'ForeignKey\("([^"]+)"\)', d)
            cols.append([name, sql, "PK" if "primary_key=True" in d else "", fk.group(1) if fk else "", "可" if "nullable=True" in d or "| None" in py else "不可", "あり" if "index=True" in d else "なし", "あり" if "default=" in d else "なし"])
        items.append({"class": cm.group(1), "table": tn.group(1), "unique": uniq, "cols": cols})
    return items


SYSTEM = [
    ["利用者端末", "職員スマートフォン / ブラウザ", "React PWAでQR読取・動画撮影・送迎操作を行う。"],
    ["フロントエンド", "GitHub Pages / React / TypeScript / Vite", "画面表示、JWT保持、API呼び出し、オフライン保留を担当。"],
    ["バックエンド", "Render / FastAPI / Python", "認証、業務ロジック、DB更新、通知、LINE Webhook、動画保存を担当。"],
    ["DB", "PostgreSQL（本番想定） / SQLite（開発）", "園、職員、園児、車両、便、送迎、通知、監査ログを保存。"],
    ["動画保存", "Render Persistent Disk / UPLOAD_DIR", "車内動画ファイルを保存。保存期限・暗号化・バックアップは要確認。"],
    ["外部連携", "LINE Messaging API / メールWebhook", "保護者LINE連携、降車通知、メール連携を担当。"],
]

SCREENS = [
    ["SCR-01", "ログイン", "全職員", "職員ID、PIN、管理者PIN復旧導線", "JWT取得後ホームへ。"],
    ["SCR-02", "ホーム", "全職員", "園名、本日の送迎状況、未同期件数、運行導線", "進行中送迎を把握。"],
    ["SCR-03", "運行", "operator/admin", "便選択、当日名簿、QR読取、手動乗降、車内撮影、完了", "中心業務画面。"],
    ["SCR-04", "園児", "admin中心", "園児一覧、登録、編集、QR文字列確認", "APIでも権限制御。"],
    ["SCR-05", "記録", "全職員", "履歴、再開、乗降詳細、安全確認、動画・AI補助", "動画は認証付き取得。"],
    ["SCR-06", "LINE", "admin", "保護者通知先、LINE連携QR、通知履歴、再送", "実運用前に結合試験。"],
    ["SCR-07", "設定", "admin", "園、職員、車両、便、通常名簿", "車両・便は非表示化。"],
]

FUNCTIONS = [
    ["F-01", "認証・権限", "JWT、職員ロール、PINハッシュ、園単位制御。", "実装確認"],
    ["F-02", "マスタ管理", "園、園児、職員、車両、便、通常名簿。", "実装確認"],
    ["F-03", "送迎開始", "便・車両・方向を選び、通常名簿から当日送迎を作成。", "実装確認"],
    ["F-04", "当日名簿変更", "欠席・臨時乗車を反映。確認済み園児の除外は拒否。", "実装確認"],
    ["F-05", "乗降記録", "QR、手入力、手動操作で乗車・降車を記録。", "実装確認"],
    ["F-06", "安全確認", "車内確認、GPS、5〜30秒動画、完了条件を制御。", "実装確認"],
    ["F-07", "第三者確認", "別職員ID・PINで任意の確認記録を保存。", "任意"],
    ["F-08", "記録・監査", "送迎履歴、動画、監査ログを参照。", "実装確認"],
    ["F-09", "オフライン同期", "通信不能時の乗降記録を端末に保留し同期。", "一部実装"],
    ["F-10", "保護者通知", "同意済み保護者へLINE・メール通知をキュー化・再送。", "実運用要確認"],
    ["F-11", "AI補助", "動画AI補助要求と結果表示。実AIは未接続。", "土台のみ"],
    ["F-12", "PDF・CSV出力 / 未確認アラーム", "帳票出力と時間超過通知。", "未実装"],
]

PROCESSES = [
    ["P-01", "ログイン", "職員IDとPINを受け取り、有効職員とPINハッシュを確認。JWTを返す。", "失敗時は認証エラー。"],
    ["P-02", "初期表示", "トークン検証後、便・車両・園児・ダッシュボードを取得。", "一部失敗時は画面メッセージ。"],
    ["P-03", "送迎開始", "通常名簿からTripAttendanceを作成。帰り便は乗車済みで開始。", "削除済み便は開始不可。"],
    ["P-04", "乗降記録", "QRまたは手動で園児を特定し、乗車/降車時刻と担当者を更新。", "対象外・重複・不正工程は拒否。"],
    ["P-05", "人数照合", "乗車数、降車数、未確認数、完了前チェック可否を計算。", "ACTIVEは安全完了ではない。"],
    ["P-06", "車内撮影完了", "動画アップロード、AI補助要求、車内確認、送迎完了を順に実行。", "動画失敗時は完了不可。"],
    ["P-07", "LINE連携", "QR案内発行後、署名・期限・未使用・同意を検証して紐づけ。", "平文トークンは保存しない。"],
    ["P-08", "降車通知", "event_keyで保護者・チャネル単位に通知を冪等作成。", "自動配信はフラグ有効時のみ。"],
    ["P-09", "オフライン同期", "client_event_id付きイベントを同期し、既処理は再適用しない。", "長時間圏外は導入前試験。"],
]

VALIDATIONS = [
    ["職員PIN", "ログイン時4〜128文字、登録時8〜128文字。", "Pydantic / ハッシュ照合"],
    ["職員ロール", "admin / operator / verifier のみ。", "Literal"],
    ["園児QR", "1〜100文字、園内一意。", "UniqueConstraint"],
    ["車両名", "1〜100文字、園内一意。", "UniqueConstraint"],
    ["乗降区分", "乗車 / 降車 のみ。", "Literal"],
    ["保護者メール", "3〜254文字。trim・小文字化し園内一意。", "正規化 / UniqueConstraint"],
    ["LINE連携", "メール通知有効かつ同意済みが必要。", "業務ロジック"],
    ["同期イベント", "client_event_id 1〜80文字、1回100件まで。", "Pydantic / UniqueConstraint"],
    ["動画形式", "Content-Typeがvideo/で始まること。", "APIロジック"],
    ["動画時間", "5〜30秒。", "validate_video_duration"],
    ["動画容量", "100MB以下。", "アップロード時検査"],
]

ENV = [
    ["DATABASE_URL", "DB接続先。未指定時はSQLite。", "Secret", "本番必須"],
    ["JWT_SECRET", "JWT署名鍵。", "Secret", "本番必須"],
    ["TOKEN_EXPIRE_MINUTES", "JWT有効期限。既定480分。", "通常", "要確認"],
    ["CORS_ORIGINS", "許可するフロントURL。", "通常", "本番必須"],
    ["VITE_API_BASE_URL", "フロントのAPI接続先。", "公開", "本番必須"],
    ["UPLOAD_DIR", "動画保存先。Renderでは/var/data/mamoru-bus-uploads。", "通常", "本番必須"],
    ["ADMIN_PIN_RECOVERY_TOKEN", "管理者PIN緊急復旧用。", "Secret", "緊急時のみ"],
    ["LINE_CHANNEL_ACCESS_TOKEN", "LINE Messaging APIアクセストークン。", "Secret", "通知時必須"],
    ["LINE_CHANNEL_SECRET", "LINE Webhook署名検証Secret。", "Secret", "通知時必須"],
    ["LINE_ORGANIZATION_ID", "LINE連携対象の園ID。", "通常", "通知時必須"],
    ["LINE_BASIC_ID", "LINE公式アカウントBasic ID。@785ntzvy。", "公開", "通知時必須"],
    ["LINE_LINK_TOKEN_PEPPER", "LINE連携トークンハッシュ強化。", "Secret", "通知時必須"],
    ["EMAIL_WEBHOOK_URL", "メール配信アダプターWebhook URL。", "Secret", "通知時必須"],
    ["EMAIL_FROM_ADDRESS", "メール送信元。", "通常", "通知時必須"],
    ["NOTIFICATION_FEATURE_ENABLED", "降車時自動配信フラグ。", "通常", "承認後true"],
]

CODES = [
    ["role", "admin / operator / verifier", "管理者、運転担当、第三者確認担当。"],
    ["trip.status", "運行中 / 完了 / 中止", "送迎の進行状態。中止は履歴一覧から除外。"],
    ["event_type", "乗車 / 降車", "園児の乗降区分。"],
    ["vehicle_check.check_type", "tail_qr / third_party", "車内確認、第三者確認。"],
    ["line_status", "not_requested / pending / linked / expired / unfollowed / revoked / error", "LINE連携状態。"],
    ["notification.status", "queued / sent / failed ほか", "通知キューの送信状態。"],
    ["video.ai_status", "queued / needs_human_review", "動画AI補助状態。"],
]


def purpose(name):
    return {
        "organizations": "園テナント情報。",
        "staff": "職員、ロール、PINハッシュ、有効状態。",
        "vehicles": "送迎車両。過去記録保持のため非表示化。",
        "bus_routes": "便、方向、車両紐づけ。",
        "children": "園児、クラス、QR文字列。",
        "route_children": "便ごとの通常名簿。",
        "bus_trips": "日々の送迎記録。",
        "trip_attendance": "送迎ごとの園児別乗降記録。",
        "vehicle_safety_checks": "車内確認、第三者確認、GPS。",
        "guardian_contacts": "保護者連絡先、同意、LINE希望。",
        "child_guardians": "保護者と園児の通知対象関係。",
        "line_link_requests": "LINE連携案内トークンのハッシュと状態。",
        "notification_queue": "LINE・メール通知キューと送信履歴。",
        "line_contacts": "LINEユーザーIDと保護者の紐づけ。",
        "audit_logs": "操作監査ログ。",
        "admin_pin_recoveries": "管理者PIN復旧トークン使用履歴。",
        "sync_events": "オフライン同期イベントの重複防止。",
        "video_evidence": "車内動画証跡のメタデータ。",
    }.get(name, "要確認")


def props(doc, title):
    doc.core_properties.title = title
    doc.core_properties.subject = "まもるバス 設計書"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "A4 landscape / BIZ UDゴシック 9pt"
    doc.core_properties.created = datetime(2026, 7, 30, 0, 0, 0)


def build_basic():
    doc = setup("基本設計書", "Basic Design / 画面・機能・構成・権限")
    heading(doc, "1. 文書概要")
    para(doc, "本書は、まもるバスの基本設計を利用者・業務担当者・開発者が同じ目線で確認するための資料である。システム構成、画面、機能、権限、通知、エラー、未決事項を整理する。")
    para(doc, "重要: 本システムは開発途中の試作版であり、送迎バスの安全装置、法令上必要な装置、職員の目視確認を代替しない。", NOTE_FILL)
    heading(doc, "2. 全体構成")
    para(doc, "図1 全体構成: 職員端末 → GitHub PagesのReact PWA → RenderのFastAPI → DB/動画保存/外部通知、という構成である。GitHub Pages単体ではAPIやDBは動作しない。")
    table(doc, ["区分", "構成要素", "役割"], SYSTEM, [4.0, 8.0, 14.0])
    heading(doc, "3. 業務フロー")
    table(doc, ["順序", "工程", "主な操作", "システム制御"], [
        ["1", "ログイン", "職員IDとPINでログインする。", "JWTを発行し、ロールを保持する。"],
        ["2", "送迎開始", "担当便・車両・方向を選ぶ。", "通常名簿から当日送迎を作成する。"],
        ["3", "当日名簿調整", "欠席・臨時乗車を反映する。", "確認済み園児の除外を拒否する。"],
        ["4", "乗車確認", "園児QRまたは手動で乗車を記録する。", "担当者・時刻を保存する。"],
        ["5", "降車確認", "到着後、園児ごとに降車を記録する。", "未降車数を算出する。"],
        ["6", "車内確認・動画", "最後尾まで目視確認し、5〜30秒撮影する。", "動画1件と車内確認記録を保存する。"],
        ["7", "送迎完了", "完了条件を満たして完了する。", "全員降車、車内確認、動画1件が必須。"],
    ], [1.5, 4.0, 10.0, 11.0])
    heading(doc, "4. 画面一覧")
    table(doc, ["画面ID", "画面名", "利用者", "主な機能", "補足"], SCREENS, [2.0, 3.0, 4.0, 12.0, 5.0])
    heading(doc, "5. 画面遷移")
    table(doc, ["遷移ID", "遷移元", "操作", "遷移先", "条件"], [
        ["TR-01", "ログイン", "ログイン成功", "ホーム", "JWT取得成功。"],
        ["TR-02", "ホーム", "運行を開く", "運行", "進行中送迎があれば再表示。"],
        ["TR-03", "運行", "送迎開始", "運行（乗車/降車）", "便選択後。帰り便は降車工程から開始。"],
        ["TR-04", "運行", "全員乗車", "運行（降車）", "行き便で乗車数が対象人数に一致。"],
        ["TR-05", "運行", "送迎完了", "ホーム", "完了条件を満たした場合。"],
        ["TR-06", "記録", "進行中送迎を選択", "運行", "運行中送迎を再開。"],
    ], [2.0, 4.0, 5.0, 5.0, 10.0])
    heading(doc, "6. 機能一覧")
    table(doc, ["機能ID", "機能", "概要", "状態"], FUNCTIONS, [2.0, 4.5, 15.5, 4.0])
    heading(doc, "7. 権限設計")
    table(doc, ["ロール", "想定利用者", "許可する主な操作", "制限"], [
        ["admin", "園管理者", "設定管理、職員、通知、監査ログ、強制完了、PIN復旧後確認", "MFAや通常PIN変更は要検討。"],
        ["operator", "運転担当・添乗職員", "送迎開始、乗降、当日名簿、車内確認、動画撮影、記録閲覧", "管理設定変更不可。"],
        ["verifier", "第三者確認担当", "第三者確認、記録確認、動画AI補助要求", "第三者確認を完了必須にするか要確認。"],
        ["保護者", "園児の保護者", "LINE連携、通知受信", "管理画面は利用しない。"],
    ], [2.6, 4.5, 13.0, 6.0])
    heading(doc, "8. 通知・帳票・エラー")
    table(doc, ["区分", "設計", "状態"], [
        ["降車通知", "同意済み保護者へLINE/メール通知。安全確認の最終判断は代替しない。", "実運用要確認"],
        ["LINE連携案内", "期限付き・一回限りQR/リンクをメール送信。平文トークンは保存しない。", "実装確認"],
        ["PDF/CSV", "送迎記録帳票とCSV出力。", "未実装"],
        ["未確認アラーム", "時間超過時警告・管理者通知。", "未実装"],
        ["エラー表示", "認証失敗、権限不足、動画不足、LINE署名不正などを日本語で返す。", "実装確認"],
    ], [4.0, 17.0, 5.0])
    heading(doc, "9. 未決定・要確認事項")
    table(doc, ["ID", "事項", "内容"], [
        ["BD-OPEN-01", "本番利用承認", "試作版。実送迎前に園責任者、関係機関、専門家の確認が必要。"],
        ["BD-OPEN-02", "動画保存運用", "保存期限、暗号化、削除、アクセス記録、バックアップ方式が未決定。"],
        ["BD-OPEN-03", "第三者確認", "任意のままか、完了必須条件にするか未決定。"],
        ["BD-OPEN-04", "AI判定", "AIプロバイダー、精度基準、責任分界が未決定。"],
    ], [3.0, 6.0, 17.0])
    props(doc, "まもるバス 基本設計書")
    return doc


def build_detail():
    doc = setup("詳細設計書", "Detailed Design / API・処理・入力チェック・実装規約")
    heading(doc, "1. 文書概要")
    para(doc, "本書は、開発者が実装・レビュー・テストを行うための詳細設計書である。バックエンドAPI、フロント処理、入力チェック、ログ、環境変数、ライブラリ、既存テストを整理する。")
    para(doc, "注意: Secretや本番トークンの実値は記載しない。設計書には環境変数名と用途のみを記載する。", NOTE_FILL)
    heading(doc, "2. プログラム構成")
    table(doc, ["領域", "ファイル", "役割"], [
        ["フロント", "src/App.tsx", "主要画面、状態管理、API呼び出し、QR読取、動画撮影、通知設定UI。"],
        ["バックエンド", "backend/main.py", "FastAPI、SQLAlchemyモデル、Pydanticモデル、業務ロジック、API。"],
        ["テスト", "backend/test_main.py", "中止、削除、動画、保護者通知、LINE連携、通知冪等性の単体テスト。"],
        ["設定", ".env.example / render.yaml", "環境変数、Render DB、Persistent Disk。"],
        ["ビルド", "package.json / vite.config.ts", "React/Vite/PWAのビルドと依存関係。"],
    ], [4.0, 7.0, 15.0])
    heading(doc, "3. 主要処理設計")
    table(doc, ["処理ID", "処理", "処理内容", "異常・注意"], PROCESSES, [2.0, 4.0, 15.0, 5.0])
    heading(doc, "4. API設計")
    table(doc, ["Method", "Path", "関数", "認証"], routes(), [2.0, 10.0, 6.0, 4.0], size=7.3)
    heading(doc, "5. 入力チェック仕様")
    table(doc, ["対象", "チェック内容", "実装箇所"], VALIDATIONS, [5.5, 14.5, 6.0])
    heading(doc, "6. 認証・権限制御詳細")
    table(doc, ["項目", "設計", "補足"], [
        ["PIN保存", "PBKDF2-SHA256でハッシュ化する。", "平文PINは保存しない。"],
        ["JWT", "HS256、TOKEN_EXPIRE_MINUTESで期限設定。", "JWT_SECRETは環境変数で管理。"],
        ["current_staff", "Bearerトークンを検証し、職員を取得。", "無効・期限切れは401。"],
        ["require_roles", "APIごとに許可ロールを指定。", "管理APIはadmin中心。"],
        ["テナント分離", "actor.organization_idで全クエリを絞る。", "動画取得も園一致を確認。"],
        ["管理者PIN復旧", "復旧トークンと新PINで職員ID 3を復旧。", "一度だけ使用し監査ログへ。"],
    ], [5.0, 14.0, 7.0])
    heading(doc, "7. フロントエンド処理詳細")
    table(doc, ["機能", "状態・関数", "処理"], [
        ["認証状態", "token / operator", "localStorageのmamoru-bus-tokenを復元し、/api/auth/meで職員情報を取得。"],
        ["初期データ", "loadBootstrap", "便、車両、園児を並列取得。"],
        ["運行開始", "startTrip", "route_id、vehicle_id、directionをPOST /api/tripsへ送信。"],
        ["QR読取", "Scanner", "BarcodeDetectorまたはjsQRで読み取る。手入力も可能。"],
        ["オフライン", "queue / sync", "通信不能時にlocalStorageへ保留し、onlineイベントで/api/syncへ送信。"],
        ["動画撮影", "VehicleVideoRecorder", "5秒後に停止可、30秒で自動停止。FormDataでアップロード。"],
        ["LINE設定", "GuardianNotificationSettings", "保護者、同意、LINE希望、QR案内、再送、解除を管理。"],
    ], [5.0, 6.0, 15.0])
    heading(doc, "8. ログ・監査設計")
    table(doc, ["ログ種別", "保存先", "出力契機", "主な項目"], [
        ["監査ログ", "audit_logs", "管理操作、乗降、安全確認、通知、LINE連携、動画操作", "actor_id、action、resource_type、detail"],
        ["通知履歴", "notification_queue", "通知作成・送信・失敗・再送", "channel、event_key、status、attempt_count"],
        ["同期結果", "sync_events", "オフラインイベント同期", "client_event_id、outcome"],
        ["動画メタ", "video_evidence", "動画アップロード・AI補助", "storage_key、content_type、ai_status"],
    ], [4.0, 5.0, 9.0, 8.0])
    heading(doc, "9. 外部連携・環境変数")
    table(doc, ["環境変数", "用途", "秘匿", "状態"], ENV, [5.5, 12.0, 4.0, 4.5], size=7.6)
    heading(doc, "10. ライブラリ・テスト")
    table(doc, ["分類", "技術・観点", "用途・確認内容"], [
        ["フロント", "React / TypeScript / Vite / Tailwind / PWA / jsQR", "画面、型、ビルド、スタイル、PWA、QR読取。"],
        ["バックエンド", "FastAPI / SQLAlchemy / Pydantic / python-jose", "API、ORM、入力チェック、JWT。"],
        ["既存テスト", "中止・削除・動画・通知・LINE", "中止送迎除外、非表示化、動画完了条件、LINE署名、通知冪等性を確認。"],
    ], [4.0, 9.0, 13.0])
    heading(doc, "11. 実装上の注意・未決事項")
    table(doc, ["ID", "事項", "内容"], [
        ["DD-OPEN-01", "backend/main.py肥大化", "API、モデル、業務ロジックが単一ファイル。今後は分割を検討。"],
        ["DD-OPEN-02", "DBマイグレーション", "本番向けマイグレーション運用は未整備。"],
        ["DD-OPEN-03", "AIプロバイダー", "analyze_videoは固定応答。実接続は未決定。"],
        ["DD-OPEN-04", "通知実運用", "LINE・メールは事業者設定、実機、実回線、保護者送信確認が必要。"],
        ["DD-OPEN-05", "セキュリティ強化", "MFA、通常PIN変更、トークン失効、動画暗号化が未実装。"],
    ], [3.0, 6.0, 17.0])
    props(doc, "まもるバス 詳細設計書")
    return doc


def build_db():
    tabs = db_tables()
    doc = setup("DB設計書", "Database Design / ER・テーブル・コード・移行・バックアップ")
    heading(doc, "1. 文書概要")
    para(doc, "本書は、SQLAlchemyモデルから確認できるDB設計を整理する資料である。テーブル、項目、キー、コード値、初期データ、移行・バックアップ上の注意を記載する。")
    para(doc, "注意: 本番向けDBマイグレーション、バックアップ、暗号化、保存期限は未実装または要確認である。", NOTE_FILL)
    heading(doc, "2. DB設計方針")
    table(doc, ["方針", "内容"], [
        ["園単位分離", "主要テーブルにorganization_idを持たせ、ログイン職員の園だけを操作する。"],
        ["履歴保持", "車両・便は削除ではなく非表示化し、過去送迎記録の名称参照を壊さない。"],
        ["監査性", "業務操作、通知、動画、復旧操作をaudit_logsへ保存する。"],
        ["通知冪等性", "event_key、guardian_contact_id、channelの一意制約で重複通知を防ぐ。"],
        ["動画分離", "動画ファイル本体はUPLOAD_DIR、メタデータはvideo_evidenceへ保存する。"],
        ["Secret非保存", "LINE連携トークン平文は保存せず、token_hashを保存する。"],
    ], [5.0, 21.0])
    heading(doc, "3. ER概要")
    table(doc, ["親", "子", "関係", "説明"], [
        ["organizations", "staff / children / vehicles / bus_routes / bus_trips", "1:N", "園ごとのマスタと送迎記録。"],
        ["bus_routes", "route_children", "1:N", "便ごとの通常名簿。"],
        ["bus_trips", "trip_attendance", "1:N", "当日の園児別乗降記録。"],
        ["bus_trips", "vehicle_safety_checks / video_evidence", "1:N", "車内確認と動画証跡。"],
        ["guardian_contacts", "child_guardians / line_contacts / notification_queue", "1:N", "保護者通知とLINE連携。"],
        ["staff", "audit_logs / video_evidence / vehicle_safety_checks", "1:N", "操作主体、アップロード者、確認者。"],
    ], [5.5, 8.0, 3.0, 10.0])
    heading(doc, "4. テーブル一覧")
    table(doc, ["No", "テーブル名", "モデル", "主な用途", "一意制約・補足"], [[i + 1, t["table"], t["class"], purpose(t["table"]), t["unique"]] for i, t in enumerate(tabs)], [1.2, 5.0, 4.5, 11.0, 5.0], size=7.4)
    heading(doc, "5. テーブル定義")
    for i, t in enumerate(tabs, 1):
        heading(doc, f"5.{i} {t['table']}", 2)
        para(doc, f"用途: {purpose(t['table'])}")
        table(doc, ["項目名", "DB型", "PK", "FK", "NULL", "索引", "既定"], t["cols"], [4.4, 3.4, 1.4, 6.0, 1.7, 1.7, 1.7], size=7.2)
    heading(doc, "6. コード定義")
    table(doc, ["区分", "値", "意味"], CODES, [5.0, 8.0, 13.0])
    heading(doc, "7. 初期データ・移行")
    table(doc, ["対象", "初期投入", "内容", "状態"], [
        ["organizations", "必要", "園名。例: バナナ幼稚園。", "要確認"],
        ["staff", "必要", "初期管理者、運転担当、第三者確認担当。", "要確認"],
        ["children", "必要", "園児名、クラス、QR文字列。", "要確認"],
        ["vehicles / bus_routes / route_children", "必要", "車両、便、通常名簿。", "要確認"],
        ["guardian_contacts / child_guardians", "任意", "保護者通知を使う場合に登録。", "実運用前要確認"],
        ["line_contacts / notification_queue", "自動作成", "LINE連携成功、通知イベントで作成。", "実装確認"],
    ], [5.5, 3.0, 13.0, 4.5])
    heading(doc, "8. バックアップ・保管")
    table(doc, ["対象", "方針", "現状", "未決事項"], [
        ["DB", "PostgreSQLの定期バックアップを利用する。", "Render DB設定あり。", "頻度、保持期間、復元試験。"],
        ["動画", "Persistent Diskまたは外部ストレージを使う。", "UPLOAD_DIRへ保存。", "暗号化、保存期限、削除、容量監視。"],
        ["監査ログ", "事故・問い合わせ対応に必要な期間を保管。", "audit_logsへ保存。", "保管期間、検索権限、出力。"],
        ["通知履歴", "問い合わせ対応のため保管。", "notification_queueへ保存。", "個人情報の保存期間、削除手順。"],
    ], [4.0, 9.0, 6.0, 7.0])
    heading(doc, "9. 未決定・要確認事項")
    table(doc, ["ID", "事項", "内容"], [
        ["DB-OPEN-01", "マイグレーション", "Alembic等による本番DBマイグレーション手順が未整備。"],
        ["DB-OPEN-02", "バックアップ復元", "DB・動画とも復元試験が未実施。"],
        ["DB-OPEN-03", "暗号化", "動画、保護者連絡先、LINEユーザーIDの暗号化方針が未決定。"],
        ["DB-OPEN-04", "保存期限", "動画、通知履歴、監査ログの保存期間と削除手順が未決定。"],
        ["DB-OPEN-05", "物理設計", "PostgreSQLのインデックス、容量見積、監視は本番前に要確認。"],
    ], [3.0, 6.0, 17.0])
    props(doc, "まもるバス DB設計書")
    return doc


def audit(path):
    d = Document(path)
    s = d.sections[0]
    with ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
        styles = z.read("word/styles.xml").decode("utf-8", "ignore")
    return {
        "path": str(path.resolve()),
        "size": path.stat().st_size,
        "tables": len(d.tables),
        "rows": sum(len(t.rows) for t in d.tables),
        "landscape": s.orientation == WD_ORIENT.LANDSCAPE,
        "page_cm": [round(s.page_width.cm, 2), round(s.page_height.cm, 2)],
        "font": "BIZ UDゴシック" in xml or "BIZ UDゴシック" in styles,
    }


if __name__ == "__main__":
    OUT.mkdir(exist_ok=True)
    files = [
        (build_basic(), OUT / "まもるバス_基本設計書_20260730.docx"),
        (build_detail(), OUT / "まもるバス_詳細設計書_20260730.docx"),
        (build_db(), OUT / "まもるバス_DB設計書_20260730.docx"),
    ]
    for doc, path in files:
        doc.save(path)
        print(audit(path))
