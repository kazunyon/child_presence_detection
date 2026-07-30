from __future__ import annotations
import importlib.util
import sys
from collections import OrderedDict
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
BASE=Path(r'C:\home\github\child_presence_detection\tmp\build_mamoru_bus_test_cases_docx.py')
OUT=Path(r'C:\home\github\child_presence_detection\outputs\まもるバス_総合テスト仕様書_上限なし版_20260729.docx')
FONT='BIZ UDゴシック'
spec=importlib.util.spec_from_file_location('base_cases', BASE); base=importlib.util.module_from_spec(spec); sys.modules[spec.name]=base; spec.loader.exec_module(base); Case=base.Case
def f(run,size=9,bold=False,color=None):
    run.font.name=FONT; run.font.size=Pt(size); run.font.bold=bold
    if color: run.font.color.rgb=RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); run._element.rPr.rFonts.set(qn('w:ascii'),FONT); run._element.rPr.rFonts.set(qn('w:hAnsi'),FONT)
def par(doc,text,style=None,size=9,bold=False,color=None,align=None):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.0
    if align is not None: p.alignment=align
    r=p.add_run(text); f(r,size,bold,color); return p
def st(*items): return '\n'.join(f'{i+1}. {x}' for i,x in enumerate(items))
def add(cases,ch,cat,item,pre,steps,expected): cases.append(Case(ch,cat,item,pre,steps,expected))
def extend(cases):
    endpoints=[('GET /api/dashboard','ホーム集計','current_staff','本日の件数、運行中、完了、未確認'),('GET /api/bootstrap','初期データ','current_staff','園児、職員、車両、便'),('GET /api/children','園児一覧','current_staff','自園の園児一覧'),('POST /api/children','園児登録','admin','園児名、クラス、QR'),('PUT /api/children/{id}','園児更新','admin','園児名、クラス、QR'),('GET /api/staff','職員一覧','admin','職員名、ロール、有効状態'),('POST /api/staff','職員登録','admin','職員名、ロール、PIN'),('PUT /api/staff/{id}','職員更新','admin','ロール、PIN、有効状態'),('GET /api/vehicles','車両一覧','current_staff','有効車両'),('POST /api/vehicles','車両登録','admin','車両名、ナンバー'),('DELETE /api/vehicles/{id}','車両非表示','admin','過去記録保持'),('GET /api/bus-routes','便一覧','current_staff','有効便と通常名簿'),('POST /api/bus-routes','便登録','admin','便名、方向、車両、名簿'),('PUT /api/bus-routes/{id}','便更新','admin','便情報と通常名簿'),('DELETE /api/bus-routes/{id}','便非表示','admin','通常名簿削除、過去保持'),('POST /api/trips','送迎開始','current_staff','便、車両、方向'),('PUT /api/trips/{id}/roster','当日名簿','operator/admin','欠席・臨時乗車'),('POST /api/trips/{id}/scans','QR乗降','current_staff','園児QR、乗車/降車'),('POST /api/trips/{id}/manual-attendance','QRなし乗降','operator/admin','園児ID、乗車/降車'),('POST /api/trips/{id}/cancel','送迎中止','operator/admin','実記録前のみ中止'),('GET /api/trips/{id}/record','記録詳細','current_staff','乗降、安全確認、動画'),('POST /api/trips/{id}/complete','送迎完了','operator/admin','全員降車、車内確認、動画1件'),('POST /api/trips/{id}/force-complete','強制完了','admin','管理者監査付き完了'),('POST /api/vehicle-checks','車内確認','current_staff','tail_qr、GPS'),('POST /api/sync','オフライン同期','current_staff','client_event_id冪等'),('GET /api/audit-logs','監査ログ','admin','検索、絞込、上限500'),('POST /api/trips/{id}/videos','動画アップロード','current_staff','動画ファイル、5〜30秒'),('GET /api/videos/{id}/download','動画取得','current_staff','同園動画ファイル'),('POST /api/videos/{id}/analyze','AI補助','admin/operator/verifier','needs_human_review'),('GET /api/guardian-contacts','保護者一覧','admin','保護者、対象園児、LINE状態'),('POST /api/guardian-contacts','保護者登録','admin','メール、同意、対象園児'),('PUT /api/guardian-contacts/{id}','保護者更新','admin','同意撤回、LINE解除、対象園児'),('POST /api/guardian-contacts/{id}/line-link-requests','LINE QR案内','admin','期限付き一回限り連携案内'),('DELETE /api/guardian-contacts/{id}/line-link','LINE解除','admin','LineContact紐付け解除'),('POST /api/notifications/{id}/retry','通知再送','admin','failed通知の再送'),('POST /api/integrations/line/webhook','LINE Webhook','署名検証','署名付きイベント、連携token')]
    for ep,name,role,data in endpoints:
        add(cases,14,'API認証',name+' 未認証拒否',f'対象API: {ep}。Authorizationヘッダーなし。',st('対象APIを認証なしで実行する','HTTPステータスと本文を確認する'),'HTTP 401。業務データ、個人情報、動画パス、通知先を返さない。')
        add(cases,14,'API権限',name+' 権限境界',f'対象API: {ep}。必要権限: {role}。operator、verifier、adminを準備。',st('各ロールで対象APIを実行する','許可/拒否結果を比較する'),f'{role}の仕様どおり許可される。権限外は403または401で拒否される。')
        add(cases,14,'園データ分離',name+' 別園データ拒否',f'対象API: {ep}。別organizationの対象IDまたはデータを準備。',st('自園トークンで別園対象を指定する','レスポンスとDB更新有無を確認する'),'別園データは取得・更新できない。存在を漏らす詳細情報を返さない。')
        add(cases,14,'入力・境界',name+' 入力境界',f'対象API: {ep}。入力データ: {data}。',st('必須不足、最大長超過、不正ID、不正形式を送信する','エラーとDB更新有無を確認する'),'422/404/409等の適切なエラーになり、不正データは保存されない。仕様未定の境界値は要確認として記録する。')
    groups={
    15:('PWA・端末','実機または対象ブラウザを準備。送迎中のテストデータを使用する。',[('PWAホーム画面追加','ホームアイコン起動でログインまたは初期画面が表示される'),('Service Worker更新','旧キャッシュが残り続けず新しい画面へ更新できる'),('iOS Safari QR','背面カメラでQRを読める。PWA単独起動時も要確認'),('Android Chrome QR','背面カメラ優先でQRを読める'),('カメラ権限拒否','手入力フォールバックが表示される'),('MediaRecorder非対応','動画録画非対応メッセージを出し完了扱いにしない'),('位置情報拒否','位置情報なしで車内確認を記録できる'),('端末回転','縦横切替でボタンや人数表示が重ならない'),('小型端末360px','主要操作がスクロールで実行できる'),('低速回線','QR記録は保留または再試行でき、動画はタイムアウト表示になる'),('複数タブ','片タブの更新後にもう片方も最新状態へ更新できる'),('端末再起動','未同期localStorageとログイン状態の復元可否を確認できる'),('省電力モード','録画中にスリープしないか、失敗時に再撮影できる'),('QR汚損','読み取り失敗時でも手入力で業務継続できる'),('QR連続読取','同一モーダル内で重複送信されない'),('通信切替','Wi-Fi/モバイル切替で中途半端な完了にならない'),('端末容量不足','動画保存失敗時に完了扱いにしない'),('ブラウザ戻る','送迎記録が消えず再開できるか確認する'),('長時間表示','カメラ・録画・タイマーのリソースが残り続けない'),('HTTPS確認','本番URLでカメラ・位置情報・録画が許可される')]),
    16:('UI・アクセシビリティ','スマートフォンとPCブラウザで対象画面を表示。',[('大ボタン視認性','主要ボタンが指で押しやすく、誤操作しにくい'),('未降車警告','未降車人数と完了不可理由が文字で分かる'),('ACTIVE表示','ACTIVEは安全確認完了ではなく開始可能状態として伝わる'),('QRなしバッジ','手動記録が通常QR記録と区別できる'),('確認ダイアログ','中止・QRなし・削除前に確認が出る'),('処理中disabled','二重送信しやすい操作は処理中に押せない'),('長い園児名','園児名がバッジやボタンと重ならない'),('長いメール','保護者メールが折返し表示される'),('空一覧','保護者・通知・便なし状態が分かる'),('キーボードログイン','Tab/Enterだけでログインできる'),('キーボード設定','フォーム入力と保存がキーボードでできる'),('色以外の状態識別','文字ラベルで状態を判別できる'),('QR画像alt','LINE QR画像に意味のあるaltがある'),('alertdialog','中止確認にroleと見出し関連付けがある'),('エラー表示','API detailが職員に分かる日本語で表示される'),('通知安全文言','通知が安全確認の最終判断を代替しない表現になっている'),('ログアウト動線','職員名ボタンの意味が運用上分かるか確認する'),('設定権限メッセージ','権限不足時に操作不可理由が表示される'),('録画秒数表示','5秒経過後STOP可、30秒自動終了が分かる'),('読み込み中','遅延時に空白やクラッシュではなく読み込み表示になる')]),
    17:('DB・整合性','テストDBに正常データと異常境界データを準備。',[('QR一意制約','同園同QRは登録できない'),('通常名簿重複','同じ園児が同一便に重複登録されない'),('当日名簿確認済み保護','確認済み園児を名簿から外せない'),('SyncEvent冪等','同一client_event_idは二重処理されない'),('NotificationQueue冪等','同一降車イベントの通知が二重作成されない'),('VideoEvidence保存キー','storage_keyが一意で園単位のパスになる'),('中止除外','中止送迎はホーム集計と履歴から除外される'),('削除済み車両履歴','過去記録の車両名が保持される'),('削除済み便履歴','過去記録の便名が保持される'),('UTC/JST集計','JST当日の境界で本日件数を計算する'),('旧DBマイグレーション','旧schemaにorganization_id等を補完する'),('マイグレーション冪等','再起動しても追加処理が重複しない'),('通知同意撤回','同意撤回でメール/LINEが停止しpending案内が失効する'),('LINE解除','LineContactの保護者紐付けが解除される'),('監査ログ検索','action/resource/query_text/from/to/limitで絞込できる'),('履歴上限','trips 200件、notifications 100件、audit 500件上限を確認する'),('強制完了監査','未確認数をdetailに残して通常完了と区別できる'),('0名送迎','対象0名時の完了可否と運用許可を要確認として残す'),('複数進行中','複数運行中が作成された場合の再開対象と運用を要確認する'),('直接DB異常','重複attendance等の異常データ時の画面挙動を要確認する')]),
    18:('外部連携・通知','LINE/メールのテストアカウント、Webhookモック、環境変数を準備。',[('LINE token未設定','failed/configurationとして通知履歴に残る'),('LINE API失敗','transport失敗として再試行時刻が設定される'),('メールWebhook未設定','failed/configurationとして残る'),('メールWebhook成功','sentとprovider_message_idを保存する'),('再試行1回目','next_attempt_atが約1分後になる'),('再試行2回目','next_attempt_atが約5分後になる'),('再試行3回目','next_attempt_atがNoneになる'),('LINE Webhook署名不正','401で連携状態を変えない'),('LINE Webhook JSON不正','400で連携状態を変えない'),('期限切れ連携token','expiredへ更新し再発行を案内する'),('同一Webhook再送','webhookEventIdで二重処理しない'),('別LINEアカウント','同一保護者への二重紐付けを拒否する'),('同一LINE別保護者','別保護者への上書き紐付けを拒否する'),('友だち解除','unfollowedへ更新し通知対象から外す'),('通知フラグOFF','キュー作成のみで外部送信しない'),('通知フラグON','キュー作成時に送信結果まで保存する'),('LINE未連携','メールのみ、または通知なしとして処理する'),('notify_alighted false','降車通知を作成しない'),('QR案内再表示不可','同じ平文tokenをDBから再表示できない'),('LINE公式アカウント変更','新Basic IDで案内され旧案内の扱いを要確認する')]),
    19:('運用・復旧・導入判定','本番相当環境、園責任者、運用手順書、バックアップ手順を準備。',[('Persistent Disk','動画が再デプロイ後も残る構成か確認する'),('動画容量監視','容量監視・削除基準は未実装のため運用で定める'),('動画保存期限','保存期限と削除手順を園規程で定める'),('動画暗号化','アプリ独自暗号化は未実装のため基盤側対策を確認する'),('事故時証跡保全','DB、動画、監査ログを紐付けて保全できる'),('DBバックアップ','DBバックアップと復元手順を別途確認する'),('動画バックアップ','UPLOAD_DIRもDBと同時点でバックアップする'),('管理者PIN復旧','一度限りの復旧tokenでID=3を復旧できる'),('復旧token削除','復旧後にRender環境変数から削除する'),('職員棚卸','退職者を無効化し最後の管理者を残す'),('園児卒園','園児削除/非表示未実装のため名簿除外と保持方針を確認する'),('訓練園分離','訓練データが本番園へ混入しない'),('紙運用併用','導入初期に紙名簿とアプリ記録を照合する'),('緊急時代替手順','API停止時も安全確認を継続できる紙手順がある'),('法令確認','アプリが安全装置や職員目視を代替しないことを責任者確認する'),('保護者問い合わせ','降車時刻、通知状態、動画証跡を説明できる'),('実機導入判定','端末、車内、回線、QR、録画が実運用で使える'),('通知導入判定','実在保護者送信前にテストアカウントで合格する'),('リリース後監視','通知失敗、動画容量、APIエラー、監査ログを定期確認する'),('総合受入判定','重大NGゼロ、要確認に責任者判断が付いた状態で受入可')])}
    for ch,(cat,pre,items) in groups.items():
        for item,exp in items:
            add(cases,ch,cat,item,pre,st('対象条件を準備する','画面またはAPIで該当操作を実施する','画面表示、DB更新、監査ログ、再操作可否を確認する'),exp)
    return cases
def make_doc():
    cases=extend(base.build_cases())
    titles=OrderedDict([(1,'第1章 起動・環境・認証'),(2,'第2章 ホーム・ダッシュボード'),(3,'第3章 設定・マスタ'),(4,'第4章 送迎開始・中止・再開'),(5,'第5章 当日名簿変更'),(6,'第6章 QR・手動乗降'),(7,'第7章 人数照合・工程遷移・完了条件'),(8,'第8章 車内撮影・動画証跡・AI補助/GPS'),(9,'第9章 第三者確認・強制完了'),(10,'第10章 過去記録・監査ログ'),(11,'第11章 オフライン同期・同時操作'),(12,'第12章 LINE・メール通知'),(13,'第13章 権限・園データ分離・セキュリティ・未実装確認'),(14,'第14章 API詳細・権限・園データ分離'),(15,'第15章 PWA・端末・ブラウザ・カメラ'),(16,'第16章 UI・アクセシビリティ・誤操作防止'),(17,'第17章 データ整合性・DB制約・マイグレーション'),(18,'第18章 外部連携障害・通知運用'),(19,'第19章 運用・復旧・導入判定')])
    doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width=Inches(11); sec.page_height=Inches(8.5)
    for m in ('top_margin','bottom_margin','left_margin','right_margin'): setattr(sec,m,Inches(0.45))
    for name in ('Normal','Heading 1','Heading 2','Heading 3'):
        style=doc.styles[name]; style.font.name=FONT; style.font.size=Pt(9); style._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); style._element.rPr.rFonts.set(qn('w:ascii'),FONT); style._element.rPr.rFonts.set(qn('w:hAnsi'),FONT)
    doc.core_properties.title='まもるバス 総合テスト仕様書 上限なし版'; doc.core_properties.author=''
    par(doc,'まもるバス 総合テスト仕様書',size=14,bold=True,color=(31,78,121),align=WD_ALIGN_PARAGRAPH.CENTER); par(doc,'送迎バス安全確認PWA／上限なし詳細版',size=10,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER); par(doc,'更新日：2026年7月29日',align=WD_ALIGN_PARAGRAPH.RIGHT)
    par(doc,f'更新概要：仕訳会計システム総合テスト仕様書の横向き8列構成を参考にしつつ、200件の圧縮上限を外し、まもるバスの実装・未実装・実園運用・障害対応まで広げた総合テストケース {len(cases)} 件を作成した。')
    par(doc,'作成目的：送迎開始、乗降、人数照合、車内目視確認、5〜30秒撮影、記録保存、保護者通知、権限、園単位データ分離、PWA端末挙動、セキュリティ、復旧、導入判定を、実施者が再現できる粒度で確認する。')
    par(doc,'実施方法：各行の確認を実施し、期待結果を満たす場合は結果欄にOK、差異がある場合はNGまたは保留を記入する。未実装・仕様未確定・運用判断が必要な行は、要確認事項として責任者判断を追記する。')
    par(doc,'確認元',style='Heading 2',size=10,bold=True,color=(31,78,121)); base.add_table(doc,['確認元','確認内容','テストへの反映','確認区分'],[('仕訳会計システム_総合テスト仕様書_200件_01.docx','横向き、8列テストケース表、結果欄運用','体裁・列構成の参考。200件上限は今回採用しない。','参考資料'),('README.md','実装状況、運用フロー、安全上の考え方、未実装範囲、環境変数','対象機能・未実装確認・導入判定に反映','仕様'),('backend/main.py','DBモデル、API、JWT、ロール、バリデーション、監査ログ、通知、動画','期待結果と異常系に反映','実装'),('src/App.tsx','画面、ナビゲーション、QR読取、当日名簿、動画撮影、LINE設定','操作手順と画面表示に反映','画面'),('backend/test_main.py','中止、削除、動画、LINE通知の既存単体テスト','重要回帰観点に反映','テスト')],[2.35,3.2,3.0,1.0])
    par(doc,'仕様・実装差異／要確認',style='Heading 2',size=10,bold=True,color=(31,78,121))
    for x in ['READMEには /api/health、/api/public-settings の記載があるが、現行 backend/main.py のルート一覧では /health のみ確認できた。公開設定APIの要否は要確認。','PDF・CSV出力、未確認アラーム、多要素認証、通常PIN変更、バックアップ・復元、実AIプロバイダー接続は未実装または土台のみ。実装済みとして扱わない。','第三者確認は任意機能であり、送迎完了の必須条件ではない。園運用で必須化する場合は仕様変更が必要。','PWA、カメラ、位置情報、MediaRecorder、低速回線、LINE/メール外部連携は実機・実回線・テストアカウントでの確認を必須とする。']: par(doc,'・'+x)
    par(doc,'1. 件数配分',style='Heading 2',size=10,bold=True,color=(31,78,121)); base.add_table(doc,['章','対象領域','件数','作成方針'],[(str(ch),title.replace(f'第{ch}章 ',''),str(sum(1 for c in cases if c.chapter==ch)),'200件上限を外し、正常・異常・権限・運用・導入前確認を詳細化') for ch,title in titles.items()],[0.55,2.8,0.7,5.0])
    par(doc,f'2. 総合テストケース一覧（{len(cases)}件）',style='Heading 2',size=10,bold=True,color=(31,78,121)); headers=['No','元No','分類','テスト項目','前提・確認データ','操作手順','期待結果','結果']; widths=[0.55,0.65,0.85,1.35,1.75,2.0,2.15,0.65]
    n=1
    for ch,title in titles.items():
        h=par(doc,title,style='Heading 3',size=10,bold=True,color=(31,78,121)); h.paragraph_format.keep_with_next=True; rows=[]
        for seq,case in enumerate([c for c in cases if c.chapter==ch],1): rows.append([f'{ch}-{seq:02d}',f'MAM-{n:03d}',case.category,case.item,case.precondition,case.steps,case.expected,'']); n+=1
        base.add_table(doc,headers,rows,widths); doc.add_paragraph()
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=footer.add_run(f'まもるバス 総合テスト仕様書 上限なし版（{len(cases)}件）'); f(r,8,False,(89,89,89))
    OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT); return OUT,len(cases)
if __name__=='__main__':
    out,count=make_doc(); print(out); print(count)
