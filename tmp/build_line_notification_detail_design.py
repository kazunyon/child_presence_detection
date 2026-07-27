from __future__ import annotations

import math
import os
from datetime import date
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(r"C:\home\github\child_presence_detection")
OUT_DIR = ROOT / "outputs"
ASSET_DIR = ROOT / "tmp" / "line_design_assets"
OUT_PATH = OUT_DIR / "LINE通知_QR連携_詳細設計書.docx"

FONT_NAME = "BIZ UD Gothic"
MONO_FONT = "Consolas"
FONT_FILE = Path(r"C:\Windows\Fonts\BIZ-UDGothicR.ttc")
FONT_BOLD_FILE = Path(r"C:\Windows\Fonts\BIZ-UDGothicB.ttc")

NAVY = "17324D"
TEAL = "176B6F"
TEAL_DARK = "0F5256"
MINT = "E8F5F2"
BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D8DEE5"
TEXT_GRAY = "4B5563"
CORAL = "B84A3A"
AMBER = "8A6500"
WHITE = "FFFFFF"
BLACK = "111827"

# Named landscape override of standard_business_brief.
# A4 landscape, 14 mm margins, table body 15010 DXA + 120 DXA indent.
TABLE_WIDTH_DXA = 15010
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name=FONT_NAME, size=None, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=5, line=1.10) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def distribute_widths(ratios: list[float], total=TABLE_WIDTH_DXA) -> list[int]:
    raw = [int(total * ratio / sum(ratios)) for ratio in ratios]
    raw[-1] += total - sum(raw)
    return raw


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def add_table(doc, headers: list[str], rows: list[list[str]], ratios: list[float], font_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, 0, 0, 1.0)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, size=font_size, color=WHITE, bold=True)
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ridx % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 14 and idx != len(row_data) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, 0, 0, 1.03)
            for line_idx, line in enumerate(str(text).split("\n")):
                if line_idx:
                    p.add_run().add_break()
                run = p.add_run(line)
                set_run_font(run, size=font_size, color=BLACK)
    set_table_geometry(table, distribute_widths(ratios))
    after = doc.add_paragraph()
    set_paragraph_spacing(after, 0, 4, 1.0)
    return table


def add_code_block(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 0, 0, 1.0)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name=MONO_FONT, size=7.8, color=NAVY)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 0, 4, 1.0)


def add_callout(doc, label: str, text: str, fill=MINT, accent=TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 0, 0, 1.08)
    lead = p.add_run(label + "　")
    set_run_font(lead, size=9.3, color=accent, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=9.1, color=BLACK)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 0, 5, 1.0)


def add_heading(doc, text: str, level=1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    # Built-in style provides the paragraph structure. Run settings keep Japanese font stable.
    set_run_font(run, size={1: 15, 2: 11.5, 3: 10}[level], color={1: NAVY, 2: TEAL, 3: TEAL_DARK}[level], bold=True)


def add_body(doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 5, 1.10)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, size=9.2, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, size=9.2)
    else:
        run = p.add_run(text)
        set_run_font(run, size=9.2)


def add_bullet(doc, text: str, level=0) -> None:
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.50 if level == 0 else 0.78)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, 0, 3, 1.10)
    run = p.add_run(text)
    set_run_font(run, size=9.1)


def create_decimal_numbering(doc) -> int:
    root = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in root.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)

    first_num_index = next((i for i, node in enumerate(root) if node.tag == qn("w:num")), len(root))
    root.insert(first_num_index, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    root.append(num)
    return num_id


def add_number(doc, num_id: int, text: str) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.insert(0, num_pr)
    set_paragraph_spacing(p, 0, 4, 1.10)
    run = p.add_run(text)
    set_run_font(run, size=9.1)


def add_page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_footer_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, 0, 0, 1.0)
    run = p.add_run("まもるバス｜LINE通知 QR連携 詳細設計書　")
    set_run_font(run, size=7.5, color=TEXT_GRAY)
    field_run = p.add_run()
    set_run_font(field_run, size=7.5, color=TEXT_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char1)
    field_run._r.append(instr_text)
    field_run._r.append(fld_char2)


def add_header(section) -> None:
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, 0, 0, 1.0)
    run = p.add_run("CHILD PRESENCE DETECTION / まもるバス")
    set_run_font(run, size=7.5, color=TEXT_GRAY, bold=True)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in (
        (1, 15, NAVY, 13, 6),
        (2, 11.5, TEAL, 9, 4),
        (3, 10, TEAL_DARK, 7, 3),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(9.1)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.10


def rounded_box(draw, xy, fill, outline, text, font, text_fill=BLACK, radius=24):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=4)
    x1, y1, x2, y2 = xy
    lines = []
    for chunk in text.split("\n"):
        lines.extend(wrap(chunk, width=12) or [""])
    text_value = "\n".join(lines)
    box = draw.multiline_textbbox((0, 0), text_value, font=font, spacing=8, align="center")
    tw, th = box[2] - box[0], box[3] - box[1]
    normalized_text_fill = text_fill if str(text_fill).startswith("#") else "#" + str(text_fill)
    draw.multiline_text(((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2), text_value, font=font, fill=normalized_text_fill, spacing=8, align="center")


def arrow(draw, start, end, color=TEAL_DARK, width=6):
    color_rgb = "#" + color
    draw.line([start, end], fill=color_rgb, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (2.55, -2.55):
        point = (end[0] + length * math.cos(angle + delta), end[1] + length * math.sin(angle + delta))
        draw.line([end, point], fill=color_rgb, width=width)


def create_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (2200, 720), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype(str(FONT_BOLD_FILE), 39)
    small = ImageFont.truetype(str(FONT_FILE), 30)
    boxes = [
        ((60, 180, 360, 470), "管理者画面\n保護者・通知設定", "E8F5F2", TEAL),
        ((480, 180, 800, 470), "FastAPI\n連携要求・通知生成", "E8EEF5", NAVY),
        ((920, 80, 1260, 310), "DB\n保護者・連携・通知", "F4F6F8", TEXT_GRAY),
        ((920, 390, 1260, 620), "メール配信サービス\nQR・リンク／通知", "FFF6DF", AMBER),
        ((1400, 80, 1740, 310), "LINE Webhook\n署名検証・紐付け", "E8F5F2", TEAL),
        ((1400, 390, 1740, 620), "LINE Messaging API\nPush送信", "E8F5F2", TEAL),
        ((1880, 180, 2160, 470), "保護者\nメール＋LINE", "FDEEEB", CORAL),
    ]
    for xy, label, fill, outline in boxes:
        rounded_box(draw, xy, "#" + fill, "#" + outline, label, font)
    arrow(draw, (360, 325), (480, 325))
    arrow(draw, (800, 260), (920, 195))
    arrow(draw, (800, 390), (920, 505))
    arrow(draw, (1260, 195), (1400, 195))
    arrow(draw, (1260, 505), (1400, 505))
    arrow(draw, (1740, 195), (1880, 270))
    arrow(draw, (1740, 505), (1880, 380))
    draw.text((60, 30), "提案構成：メールを本人確認の起点とし、QRでLINEユーザーIDを紐付ける", font=small, fill="#" + NAVY)
    img.save(path)


def create_link_flow_diagram(path: Path) -> None:
    img = Image.new("RGB", (2200, 820), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype(str(FONT_BOLD_FILE), 34)
    small = ImageFont.truetype(str(FONT_FILE), 28)
    role_y = 70
    roles = [("管理者", 130), ("API / DB", 690), ("メール", 1250), ("保護者 / LINE", 1810)]
    for label, x in roles:
        draw.text((x - 70, role_y), label, font=font, fill="#" + NAVY)
        draw.line([(x, 130), (x, 780)], fill="#" + MID_GRAY, width=3)
    steps = [
        (130, 690, 190, "1. 保護者メール・園児・LINE希望を登録"),
        (690, 1250, 300, "2. 一回限りの連携トークンを生成し、QR付き案内を送信"),
        (1250, 1810, 410, "3. QRを読む／リンクをタップ"),
        (1810, 690, 520, "4. LINEトークへ「連携 <token>」を送信"),
        (690, 1810, 630, "5. 署名・期限・未使用を検証し、LINE userIdを紐付け"),
        (1810, 130, 740, "6. 管理画面へ「連携済み」を表示"),
    ]
    for sx, ex, y, label in steps:
        arrow(draw, (sx, y), (ex, y), NAVY if sx < ex else TEAL_DARK, 5)
        mx = (sx + ex) / 2
        bbox = draw.textbbox((0, 0), label, font=small)
        tw = bbox[2] - bbox[0]
        draw.rectangle((mx - tw / 2 - 12, y - 44, mx + tw / 2 + 12, y - 5), fill="white")
        draw.text((mx - tw / 2, y - 41), label, font=small, fill="#" + BLACK)
    img.save(path)


def create_dispatch_flow_diagram(path: Path) -> None:
    img = Image.new("RGB", (2200, 720), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.truetype(str(FONT_BOLD_FILE), 35)
    boxes = [
        ((50, 225, 360, 500), "降車記録\n成立", "E8EEF5", NAVY),
        ((470, 225, 790, 500), "通知対象を\n保護者単位で抽出", "F4F6F8", TEXT_GRAY),
        ((900, 225, 1220, 500), "event_keyで\n二重生成を防止", "FFF6DF", AMBER),
        ((1340, 80, 1680, 310), "LINEキュー\nPush送信", "E8F5F2", TEAL),
        ((1340, 410, 1680, 640), "メールキュー\nメール送信", "E8EEF5", NAVY),
        ((1830, 225, 2160, 500), "保護者へ併送\n失敗は個別再送", "FDEEEB", CORAL),
    ]
    for xy, label, fill, outline in boxes:
        rounded_box(draw, xy, "#" + fill, "#" + outline, label, font)
    arrow(draw, (360, 360), (470, 360))
    arrow(draw, (790, 360), (900, 360))
    arrow(draw, (1220, 315), (1340, 195))
    arrow(draw, (1220, 405), (1340, 525))
    arrow(draw, (1680, 195), (1830, 315))
    arrow(draw, (1680, 525), (1830, 405))
    img.save(path)


def add_picture(doc, path: Path, width=10.0, alt_text: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 2, 4, 1.0)
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    if alt_text:
        inline_shape._inline.docPr.set("descr", alt_text)
        inline_shape._inline.docPr.set("title", alt_text)


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture = ASSET_DIR / "architecture.png"
    link_flow = ASSET_DIR / "line_link_flow.png"
    dispatch_flow = ASSET_DIR / "notification_dispatch_flow.png"
    create_architecture_diagram(architecture)
    create_link_flow_diagram(link_flow)
    create_dispatch_flow_diagram(dispatch_flow)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(7)
    configure_styles(doc)
    add_header(section)
    add_footer_page_number(section)

    # Cover / memo masthead
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 10, 0, 1.0)
    r = p.add_run("DETAIL DESIGN")
    set_run_font(r, size=10, color=TEAL, bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 2, 5, 1.0)
    r = p.add_run("LINE通知（QR連携）詳細設計書")
    set_run_font(r, size=24, color=NAVY, bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 16, 1.0)
    r = p.add_run("保護者メールアドレス登録・LINE連携・LINE／メール併送")
    set_run_font(r, size=13, color=TEXT_GRAY, bold=True)

    add_table(
        doc,
        ["項目", "内容", "項目", "内容"],
        [
            ["対象システム", "まもるバス（child_presence_detection）", "文書版", "1.1"],
            ["作成日", str(date(2026, 7, 27)), "文書状態", "実装前レビュー用"],
            ["対象機能", "LINE通知／QR連携／保護者メール", "想定読者", "開発者・園管理者・レビュー担当"],
            ["設計基準", "現行コード・README（2026-07-26時点）", "実装判定", "現行／追加／要決定を区分"],
        ],
        [1.2, 4.1, 1.2, 4.1],
        9.0,
    )
    add_callout(
        doc,
        "重要前提",
        "「LINE通知を希望する保護者はメールアドレスを必須登録し、QRでLINEを紐付ける。通知はLINEとメールへ併送する」と解釈した設計である。"
        "QR案内はメール本文にQR画像とタップ可能なリンクを併記し、スマートフォンだけでも連携できるようにする。"
        "採用するLINE公式アカウントは「バナナ幼稚園」（LINE ID：@408mrkbk）とする。",
    )
    add_body(doc, "現行のLINE API実装は土台として再利用するが、園児・保護者・メール・LINE userIdの対応付け、同意、画面、再送、到達確認は追加実装対象である。")

    add_heading(doc, "改訂履歴", 1)
    add_table(
        doc,
        ["版", "日付", "変更内容", "作成／確認"],
        [
            ["1.0", "2026-07-27", "初版。QR連携、保護者メール、LINE／メール併送の詳細設計を作成", "Codex作成／人による承認待ち"],
            ["1.1", "2026-07-27", "採用LINE公式アカウント「バナナ幼稚園」（@408mrkbk）を追記", "利用アカウントは依頼者指定"],
        ],
        [0.8, 1.3, 6.8, 1.7],
    )

    add_page_break(doc)
    add_heading(doc, "1. 目的・範囲・前提", 1)
    add_heading(doc, "1.1 目的", 2)
    add_body(doc, "降車等の記録を、通知を希望する保護者へLINEとメールで確実に案内するため、保護者のメールアドレスを起点としてLINEアカウントをQRで安全に紐付ける。")
    add_bullet(doc, "園児ごとに通知先となる保護者を登録できること")
    add_bullet(doc, "メールで受け取ったQRまたはリンクからLINE連携できること")
    add_bullet(doc, "LINEとメールの各送信結果を独立して記録し、一方だけ失敗した場合に個別再送できること")
    add_bullet(doc, "QR、ログ、画面にメールアドレスやLINE userId等の個人情報を不要に露出しないこと")

    add_heading(doc, "1.2 対象範囲", 2)
    add_table(
        doc,
        ["区分", "対象", "対象外／別途決定"],
        [
            ["画面", "管理者向け保護者・通知設定、連携状況、再送・解除", "保護者専用ポータル、園児QRの変更"],
            ["API", "保護者連絡先、LINE連携要求、Webhook拡張、通知イベント、再送", "決済、外部CRM連携"],
            ["DB", "保護者、園児紐付け、連携要求、LINE宛先拡張、通知冪等性", "長期分析用DWH"],
            ["通知", "降車記録を初期対象としたLINE／メール併送", "安否・安全の最終判定、緊急通報の代替"],
            ["運用", "登録、同意、リンク期限、解除、テスト、障害時手順", "法務最終判断、実園のリリース承認"],
        ],
        [1.0, 5.0, 4.6],
    )

    add_heading(doc, "1.3 用語", 2)
    add_table(
        doc,
        ["用語", "定義"],
        [
            ["保護者連絡先", "メールアドレスを必須とし、1人以上の園児と紐付く通知先。兄弟姉妹は同一連絡先へ複数紐付け可能。"],
            ["LINE連携要求", "保護者連絡先とLINE userIdを結ぶための、期限付き・一回限りのトークン。"],
            ["QR連携", "QR内のLINEトーク起動URLから、事前入力された連携トークンを公式アカウントへ送信して紐付ける方式。"],
            ["採用LINE公式アカウント", "「バナナ幼稚園」（LINE ID：@408mrkbk）。QR連携、Webhook受信、Push送信に使用する。"],
            ["併送", "同一通知イベントについてLINE用とメール用の通知キューを別レコードで作成し、独立送信すること。"],
            ["event_key", "同一園児・同一運行・同一通知種別の二重生成を防ぐ一意キー。"],
        ],
        [2.1, 8.5],
    )

    add_heading(doc, "1.4 現行実装と追加範囲", 2)
    add_table(
        doc,
        ["機能", "現行状態", "詳細設計での扱い"],
        [
            ["署名付きWebhook", "実装済み（X-Line-Signature / HMAC-SHA256）", "維持。messageイベントの連携トークン処理を追加"],
            ["LINE宛先登録", "友だち追加／解除でline_contactsへ登録", "保護者連絡先との紐付け、状態・監査項目を追加"],
            ["LINE Push", "通知キューからPush送信、Retry Keyあり", "保護者単位・event_key・再送回数を追加"],
            ["メール", "channel=emailは汎用Webhook経由のみ", "メール配信サービス仕様と結果記録を明確化"],
            ["管理画面", "LINE通知の操作画面なし", "保護者・通知設定画面を追加"],
            ["同意・解除", "未実装", "同意日時、停止、LINE解除、再連携を追加"],
            ["実運用確認", "未実施", "テスト仕様と導入判定を定義。結果は未実施のまま記載"],
        ],
        [2.1, 3.7, 4.8],
    )

    add_heading(doc, "2. システム構成", 1)
    add_picture(doc, architecture, 10.3, "LINE通知QR連携の構成図。管理画面、FastAPI、DB、メール配信サービス、LINE Messaging API、保護者の接続関係を示す。")
    add_table(
        doc,
        ["構成要素", "責務", "現行／追加"],
        [
            ["React管理画面", "保護者・園児・メール・通知希望の登録、連携状況表示、再送・解除", "追加"],
            ["FastAPI", "組織・権限検証、連携トークン発行、Webhook処理、通知イベント生成", "現行拡張"],
            ["DB", "保護者、園児対応、連携要求、LINE宛先、通知送信履歴、監査ログ", "現行拡張"],
            ["メール配信サービス", "QR連携案内とイベント通知メールの送信、配信結果返却", "追加／事業者要決定"],
            ["LINE Messaging API", "Webhook受信、Push送信", "現行再利用"],
            ["保護者", "メール受信、QR読取またはリンクタップ、LINEトークで連携確定", "新運用"],
        ],
        [2.0, 6.7, 1.9],
    )
    add_callout(doc, "構成上の注意", "現在のLINE_ORGANIZATION_IDは環境変数で1園を固定する方式である。複数園運用では、Webhookパスまたはチャネル設定からorganization_idを安全に解決する方式へ変更する。", "FFF6DF", AMBER)

    add_heading(doc, "2.1 採用LINE公式アカウントとQR方式", 2)
    add_table(
        doc,
        ["項目", "設定値", "設計上の扱い"],
        [
            ["公式アカウント名", "バナナ幼稚園", "保護者に表示される通知元・連携先"],
            ["LINE ID（Basic ID）", "@408mrkbk", "LINE_BASIC_IDへ設定"],
            ["アカウントページ", "https://page.line.biz/account/@408mrkbk", "管理・確認用。保護者向け案内はQR／連携リンクを使用"],
            ["トーク起動先", "https://line.me/R/oaMessage/@408mrkbk/", "QRとメール内の連携リンク生成に使用"],
        ],
        [2.2, 4.0, 4.4],
        8.2,
    )
    add_body(doc, "QRにはメールアドレス、園児名、LINE userIdを含めず、Messaging API公式アカウントのトークを起動し、一回限りの連携トークンを入力済みにするURLだけを格納する。")
    add_code_block(doc, "QR payload（採用値）:\nhttps://line.me/R/oaMessage/@408mrkbk/?{URL_ENCODED('連携 ' + raw_token)}")
    add_bullet(doc, "保護者は送信前に「連携 <token>」を確認し、送信操作で確定する。")
    add_bullet(doc, "メールにはQR画像と同じ遷移先のリンクを併記し、同じ端末でメールを開いた場合も連携できる。")
    add_bullet(doc, "LINE_BASIC_IDは@408mrkbkを設定する。LINE側URL仕様は実装時にLINE Developers公式仕様と実機で再確認する。")
    add_bullet(doc, "LINE_CHANNEL_ACCESS_TOKENとLINE_CHANNEL_SECRETは、この公式アカウントへ接続された同一Messaging APIチャネルの値を使用する。対応関係は実装・本番設定時にLINE Developersで確認する。")

    add_heading(doc, "3. 業務フロー・ユースケース", 1)
    add_picture(doc, link_flow, 10.3, "LINE通知QR連携の業務フロー。管理者による保護者登録から案内メール、QRまたはリンク、署名付きWebhook、連携完了までを示す。")
    add_heading(doc, "3.1 ユースケース一覧", 2)
    add_table(
        doc,
        ["ID", "ユースケース", "実行者", "開始条件", "完了条件"],
        [
            ["UC-01", "保護者連絡先を登録", "管理者", "園児が登録済み", "メール・園児・同意・通知希望が保存"],
            ["UC-02", "QR連携案内を送る", "管理者", "メール有効、LINE希望ON", "案内メール送信、状態=pending"],
            ["UC-03", "LINEを連携", "保護者", "要求が未使用・期限内", "line_user_id紐付け、状態=linked"],
            ["UC-04", "降車通知を併送", "システム", "降車記録成立、通知ON", "LINE／メールキューが作成・送信"],
            ["UC-05", "失敗通知を再送", "管理者", "channel単位でfailed", "新しい試行として送信結果を記録"],
            ["UC-06", "LINE連携を解除", "管理者／保護者", "linked", "LINE送信停止、メールは設定に従い継続"],
        ],
        [0.8, 2.6, 1.3, 3.0, 2.9],
    )

    add_heading(doc, "3.2 業務ルール", 2)
    add_table(
        doc,
        ["ID", "ルール"],
        [
            ["BR-01", "LINE通知希望がONの場合、正規化済みメールアドレスを必須とする。"],
            ["BR-02", "通知同意日時・同意取得者・同意方法を保存し、同意なしでは通知を有効化しない。"],
            ["BR-03", "連携要求は24時間有効、1回のみ使用可。再発行時は旧要求をrevokedへ変更する。"],
            ["BR-04", "1つのLINE userIdは同一園内で1つの保護者連絡先だけに紐付ける。"],
            ["BR-05", "1つの保護者連絡先は複数園児に紐付け可能。園児ごとに通知種別をON/OFFできる。"],
            ["BR-06", "LINE unfollow受信時はLINEのみ停止し、メール通知はemail_enabledの設定に従い継続する。"],
            ["BR-07", "降車通知文言は記録成立を伝える。安全確認完了・所在保証等の断定はしない。"],
            ["BR-08", "通知生成はevent_keyで冪等化し、再試行時に同一イベントの新規通知を重複生成しない。"],
            ["BR-09", "送信済みLINE通知はMessaging API上の既読・到達を保証しない。sentはAPI受付成功を意味する。"],
            ["BR-10", "実在保護者への本番送信は、テストアカウントでの結合試験と園責任者の承認後に行う。"],
        ],
        [1.1, 9.5],
    )

    add_heading(doc, "3.3 状態遷移", 2)
    add_table(
        doc,
        ["対象", "状態", "遷移契機", "次状態"],
        [
            ["LINE連携", "not_requested", "LINE希望ON・案内送信", "pending"],
            ["LINE連携", "pending", "有効トークンのmessage受信", "linked"],
            ["LINE連携", "pending", "期限超過", "expired"],
            ["LINE連携", "linked", "unfollow／管理者解除", "unfollowed／revoked"],
            ["通知", "queued", "送信開始", "sending"],
            ["通知", "sending", "プロバイダー2xx", "sent"],
            ["通知", "sending", "通信・4xx・5xx", "failed"],
            ["通知", "failed", "再送操作・自動再送", "queued（attempt+1）"],
        ],
        [1.8, 1.6, 4.8, 2.4],
    )

    add_heading(doc, "4. 画面詳細設計", 1)
    add_heading(doc, "4.1 画面一覧", 2)
    add_table(
        doc,
        ["画面ID", "画面名", "権限", "主な機能"],
        [
            ["SC-LN-01", "保護者・通知設定一覧", "admin", "園児別の保護者、メール、通知希望、LINE状態、最終送信結果を表示"],
            ["SC-LN-02", "保護者連絡先登録／編集", "admin", "保護者名、メール、園児、関係、通知種別、同意を登録"],
            ["SC-LN-03", "LINE連携案内", "admin", "QRプレビュー、期限、案内メール送信、再発行"],
            ["SC-LN-04", "通知履歴", "admin", "LINE／メールの状態、provider応答、失敗理由、再送"],
        ],
        [1.3, 2.6, 1.2, 5.5],
    )

    add_heading(doc, "4.2 SC-LN-02 入力項目", 2)
    add_table(
        doc,
        ["項目", "UI", "必須", "検証", "保存先"],
        [
            ["保護者名", "テキスト", "任意", "100文字以内。画面表示用", "guardian_contacts.name"],
            ["メールアドレス", "email", "必須", "前後空白除去、小文字化、254文字以内、形式確認", "guardian_contacts.email"],
            ["対象園児", "複数選択", "必須", "同一organization_idの有効な園児", "child_guardians"],
            ["続柄", "選択＋任意入力", "任意", "50文字以内", "child_guardians.relationship"],
            ["メール通知", "チェック", "必須", "LINE希望時も原則ON。個別停止可能", "guardian_contacts.email_enabled"],
            ["LINE通知", "チェック", "必須", "ONならメール必須。未連携はpending表示", "guardian_contacts.line_enabled"],
            ["降車通知", "チェック", "必須", "園児単位。初期値ON", "child_guardians.notify_alighted"],
            ["同意確認", "チェック", "必須", "ONでないと保存不可", "consented_at / consented_by"],
        ],
        [2.0, 1.5, 0.8, 4.4, 1.9],
    )

    add_heading(doc, "4.3 一覧表示・操作", 2)
    add_table(
        doc,
        ["列／操作", "仕様"],
        [
            ["園児・保護者", "園児名、保護者名、続柄を表示。複数園児はタグ表示。"],
            ["メール", "表示時はk***@example.jpのように一部マスク。編集画面のみ全体表示。"],
            ["LINE状態", "未希望／案内待ち／連携待ち／連携済み／解除／期限切れ／エラー。色だけに依存せず文言を表示。"],
            ["案内送信", "LINE希望ONかつメール有効時のみ活性。送信前に宛先・期限を確認。"],
            ["再発行", "既存pendingを失効させ、新tokenを生成。二重連携を防ぐ確認ダイアログを表示。"],
            ["解除", "LINE紐付けだけを解除。メール通知を停止する場合は別チェックを変更。"],
            ["通知履歴", "直近のLINE／メール状態、送信時刻、失敗理由の要約、再送ボタンを表示。"],
        ],
        [2.2, 8.4],
    )

    add_heading(doc, "4.4 メッセージ・エラー表示", 2)
    add_table(
        doc,
        ["コード", "表示文言", "利用者対応"],
        [
            ["VAL-EMAIL-001", "メールアドレスの形式を確認してください。", "入力修正"],
            ["VAL-LINE-001", "LINE通知を希望する場合はメールアドレスを登録してください。", "メール入力"],
            ["LINK-EXPIRED", "連携期限が切れています。園へ再発行をご依頼ください。", "管理者が再発行"],
            ["LINK-USED", "この連携案内はすでに使用されています。", "状態確認、必要時解除・再発行"],
            ["LINK-MISMATCH", "連携情報を確認できませんでした。", "詳細を画面へ出さず監査ログ確認"],
            ["SEND-PARTIAL", "LINEまたはメールの一方が送信できませんでした。通知履歴を確認してください。", "失敗channelのみ再送"],
        ],
        [1.8, 5.7, 3.1],
    )

    add_heading(doc, "5. API詳細設計", 1)
    add_heading(doc, "5.1 API一覧", 2)
    add_table(
        doc,
        ["Method", "Path", "権限", "用途"],
        [
            ["GET", "/api/guardian-contacts", "admin", "保護者・園児・連携状態一覧"],
            ["POST", "/api/guardian-contacts", "admin", "保護者連絡先を登録"],
            ["PUT", "/api/guardian-contacts/{id}", "admin", "メール、園児、通知希望、同意を更新"],
            ["POST", "/api/guardian-contacts/{id}/line-link-requests", "admin", "連携要求生成＋案内メール送信"],
            ["GET", "/api/guardian-contacts/{id}/line-link-status", "admin", "現在の連携状態取得"],
            ["DELETE", "/api/guardian-contacts/{id}/line-link", "admin", "LINE紐付け解除"],
            ["POST", "/api/integrations/line/webhook", "LINE署名", "follow/message/unfollow処理（現行拡張）"],
            ["POST", "/api/notification-events", "admin/operator/internal", "通知イベントを冪等生成"],
            ["GET", "/api/notifications", "admin", "通知履歴（現行拡張）"],
            ["POST", "/api/notifications/{id}/retry", "admin", "失敗channelを再送"],
        ],
        [1.0, 5.3, 2.0, 2.3],
        8.0,
    )

    add_heading(doc, "5.2 保護者連絡先登録", 2)
    add_code_block(
        doc,
        'POST /api/guardian-contacts\n'
        '{\n'
        '  "name": "山田 花子",\n'
        '  "email": "hanako@example.jp",\n'
        '  "email_enabled": true,\n'
        '  "line_enabled": true,\n'
        '  "consent": true,\n'
        '  "children": [{"child_id": 12, "relationship": "母", "notify_alighted": true}]\n'
        '}\n'
        '201 Created -> {"id": 31, "line_status": "not_requested"}',
    )
    add_table(
        doc,
        ["検証", "HTTP", "detail"],
        [
            ["未認証／権限不足", "401 / 403", "既存認証規約に従う"],
            ["別園のchild_id", "404", "園児が見つかりません"],
            ["メール不正", "422", "メールアドレスの形式を確認してください"],
            ["同一園で同一正規化メール", "409", "既存連絡先へ園児を追加するよう案内"],
            ["LINE希望ON・同意false", "422", "通知同意を確認してください"],
        ],
        [4.4, 1.2, 5.0],
    )

    add_heading(doc, "5.3 LINE連携要求発行", 2)
    add_code_block(
        doc,
        'POST /api/guardian-contacts/31/line-link-requests\n'
        '{"delivery": "email"}\n\n'
        '202 Accepted -> {\n'
        '  "request_id": 88,\n'
        '  "status": "pending",\n'
        '  "expires_at": "2026-07-28T10:00:00+09:00",\n'
        '  "email_delivery_status": "queued"\n'
        '}',
    )
    add_bullet(doc, "raw_tokenはsecrets.token_urlsafe(32)相当の128bit以上とし、応答・ログ・DBへ平文保存しない。")
    add_bullet(doc, "DBにはSHA-256(token + server_pepper)を保存し、Webhookで受信したtokenを同じ方法で照合する。")
    add_bullet(doc, "QR画像は案内メール作成時だけ生成し、永続保存しない。再表示が必要な場合は再発行する。")

    add_heading(doc, "5.4 LINE Webhook拡張", 2)
    add_table(
        doc,
        ["event.type", "処理", "応答"],
        [
            ["follow", "line_contactsをupsertしis_active=true。tokenはまだないため未紐付けでも可。", "200"],
            ["message(text)", "「連携 <token>」形式のみ連携処理。署名・期限・状態・organizationを検証。", "200（業務エラーも再送防止のため原則200）"],
            ["unfollow", "line_contacts.is_active=false、line_status=unfollowed、監査ログ記録。", "200"],
            ["その他", "無視し、個人情報を含まないイベント種別のみ監査集計。", "200"],
        ],
        [1.7, 7.7, 1.2],
    )
    add_code_block(
        doc,
        "Webhook検証順序\n"
        "1. LINE_CHANNEL_SECRET / organization解決設定を確認\n"
        "2. raw bodyでX-Line-SignatureをHMAC-SHA256検証\n"
        "3. event.webhookEventIdで重複受信を排除\n"
        "4. message textからtokenを抽出・hash化\n"
        "5. pendingかつexpires_at > nowの要求をSELECT FOR UPDATE\n"
        "6. line_user_id重複、園、保護者状態を検証\n"
        "7. line_contactsとguardian_contactを紐付け、要求をusedへ更新\n"
        "8. commit後、連携完了Replyを可能なら送信",
    )

    add_heading(doc, "5.5 通知イベント生成", 2)
    add_code_block(
        doc,
        'POST /api/notification-events\n'
        '{"event_type":"child.alighted","trip_id":501,"child_id":12,"occurred_at":"2026-07-27T17:42:03+09:00"}\n\n'
        '200 OK -> {"event_key":"org:1:trip:501:child:12:alighted","created":2,"skipped":0}',
    )
    add_bullet(doc, "保護者連絡先が2件ある場合は、各保護者についてLINE／メール最大2channelを生成する。")
    add_bullet(doc, "LINE未連携・unfollowedはLINEをskippedとし、メールは継続する。")
    add_bullet(doc, "同一event_key・guardian_contact_id・channelの一意制約で二重キューを防ぐ。")

    add_heading(doc, "6. データベース詳細設計", 1)
    add_heading(doc, "6.1 guardian_contacts（追加）", 2)
    add_table(
        doc,
        ["列", "型", "NULL", "制約／初期値", "説明"],
        [
            ["id", "BIGINT", "NO", "PK", "保護者連絡先ID"],
            ["organization_id", "BIGINT", "NO", "FK, INDEX", "園ID。全検索で必ず条件指定"],
            ["name", "VARCHAR(100)", "YES", "-", "保護者表示名"],
            ["email", "VARCHAR(254)", "NO", "-", "入力値（表示・送信用）"],
            ["email_normalized", "VARCHAR(254)", "NO", "UQ(org,email_normalized)", "trim＋lower。重複判定"],
            ["email_enabled", "BOOLEAN", "NO", "TRUE", "メール通知ON/OFF"],
            ["line_enabled", "BOOLEAN", "NO", "FALSE", "LINE通知希望"],
            ["line_status", "VARCHAR(30)", "NO", "not_requested", "pending/linked/expired/unfollowed/revoked/error"],
            ["consented_at", "TIMESTAMPTZ", "YES", "-", "通知同意日時"],
            ["consented_by", "BIGINT", "YES", "FK staff", "同意登録者"],
            ["is_active", "BOOLEAN", "NO", "TRUE", "論理無効"],
            ["created_at / updated_at", "TIMESTAMPTZ", "NO", "now", "作成／更新日時"],
        ],
        [2.8, 2.0, 0.8, 2.5, 2.5],
        7.7,
    )

    add_heading(doc, "6.2 child_guardians（追加）", 2)
    add_table(
        doc,
        ["列", "型", "NULL", "制約／初期値", "説明"],
        [
            ["id", "BIGINT", "NO", "PK", "紐付けID"],
            ["organization_id", "BIGINT", "NO", "FK, INDEX", "園ID（結合時の漏えい防止）"],
            ["child_id", "BIGINT", "NO", "FK children", "園児ID"],
            ["guardian_contact_id", "BIGINT", "NO", "FK guardian_contacts", "保護者連絡先ID"],
            ["relationship", "VARCHAR(50)", "YES", "-", "続柄"],
            ["notify_alighted", "BOOLEAN", "NO", "TRUE", "降車通知"],
            ["created_at / updated_at", "TIMESTAMPTZ", "NO", "now", "作成／更新日時"],
        ],
        [2.9, 1.9, 0.8, 2.5, 2.5],
        7.8,
    )
    add_body(doc, "一意制約：UNIQUE (organization_id, child_id, guardian_contact_id)。childとguardian_contactが同一organization_idであることをサービス層でも検証する。")

    add_heading(doc, "6.3 line_link_requests（追加）", 2)
    add_table(
        doc,
        ["列", "型", "NULL", "制約／初期値", "説明"],
        [
            ["id", "BIGINT", "NO", "PK", "連携要求ID"],
            ["organization_id", "BIGINT", "NO", "FK, INDEX", "園ID"],
            ["guardian_contact_id", "BIGINT", "NO", "FK, INDEX", "対象保護者"],
            ["token_hash", "CHAR(64)", "NO", "UNIQUE", "pepper付きSHA-256。平文禁止"],
            ["status", "VARCHAR(20)", "NO", "pending", "pending/used/expired/revoked"],
            ["expires_at", "TIMESTAMPTZ", "NO", "-", "発行から24時間"],
            ["requested_by", "BIGINT", "NO", "FK staff", "発行管理者"],
            ["email_notification_id", "BIGINT", "YES", "FK notification_queue", "QR案内メールの送信履歴"],
            ["used_at", "TIMESTAMPTZ", "YES", "-", "連携成立日時"],
            ["created_at", "TIMESTAMPTZ", "NO", "now", "作成日時"],
        ],
        [2.9, 1.9, 0.8, 2.5, 2.5],
        7.6,
    )

    add_heading(doc, "6.4 line_contacts（現行拡張）", 2)
    add_table(
        doc,
        ["列", "変更", "説明"],
        [
            ["guardian_contact_id", "NULL許可で追加、FK", "未紐付けのfollowを許容。連携成立時に設定"],
            ["last_webhook_event_id", "NULL許可で追加、INDEX", "Webhook重複判定。別途イベント表でも可"],
            ["last_event_at", "NULL許可で追加", "follow/message/unfollowの最終受信日時"],
            ["display_name", "現行維持", "必要時のみProfile APIで取得。通知に不要なら未取得"],
            ["UNIQUE(org,line_user_id)", "現行維持", "1 LINE userIdの園内重複を防止"],
        ],
        [3.1, 3.4, 4.1],
    )

    add_heading(doc, "6.5 notification_queue（現行拡張）", 2)
    add_table(
        doc,
        ["列", "型", "説明"],
        [
            ["guardian_contact_id", "BIGINT NULL", "通知先保護者。既存手動通知との互換でNULL可"],
            ["event_key", "VARCHAR(160) NULL", "同一業務イベントの識別子"],
            ["template_key", "VARCHAR(60) NULL", "child.alighted.v1 / line.link.v1等"],
            ["subject", "VARCHAR(200) NULL", "メール件名。LINEではNULL"],
            ["status", "VARCHAR(30)", "queued/sending/sent/failed/skipped"],
            ["attempt_count", "INTEGER", "送信試行回数。初期値0"],
            ["next_attempt_at", "TIMESTAMPTZ NULL", "自動再送予定"],
            ["provider_message_id", "VARCHAR(200) NULL", "メール事業者等の受付ID"],
            ["error_code", "VARCHAR(60) NULL", "分類済みエラー"],
            ["dedupe制約", "-", "UNIQUE(event_key, guardian_contact_id, channel)"],
        ],
        [3.4, 2.6, 4.6],
    )
    add_callout(doc, "マイグレーション方針", "現行は起動時ALTER TABLEの軽量移行である。今回の複数テーブル・一意制約・外部キー追加はAlembic等で版管理し、バックアップ、dry-run、ロールバック手順を用意する。create_allだけを本番移行手段にしない。", "FFF6DF", AMBER)

    add_heading(doc, "7. 処理詳細", 1)
    add_heading(doc, "7.1 QR連携処理", 2)
    qr_num_id = create_decimal_numbering(doc)
    add_number(doc, qr_num_id, "管理者が保護者連絡先、対象園児、LINE希望、同意を保存する。")
    add_number(doc, qr_num_id, "APIは同一園の保護者・園児であることを確認し、既存pendingをrevokedにする。")
    add_number(doc, qr_num_id, "暗号学的乱数でraw_tokenを生成し、hashと期限だけをline_link_requestsへ保存する。")
    add_number(doc, qr_num_id, "raw_tokenからLINEトーク起動URLとQRを一時生成し、登録メールへ送る。")
    add_number(doc, qr_num_id, "保護者がQR／リンクを開き、事前入力済み連携メッセージを送る。")
    add_number(doc, qr_num_id, "Webhookはraw body署名検証後、token要求を排他取得し、期限・状態・園・LINE重複を確認する。")
    add_number(doc, qr_num_id, "line_contacts.guardian_contact_idを設定し、要求をused、保護者をlinkedへ更新する。")
    add_number(doc, qr_num_id, "監査ログを保存し、保護者へ連携完了メッセージ、管理画面へlinked状態を返す。")

    add_heading(doc, "7.2 通知生成・併送処理", 2)
    add_picture(doc, dispatch_flow, 10.3, "通知生成と併送処理図。降車記録成立後、保護者単位で通知対象を抽出し、event_keyで重複を防止してLINEとメールへ個別送信する。")
    add_table(
        doc,
        ["段階", "処理", "失敗時"],
        [
            ["イベント受付", "降車記録成立後にevent_keyを生成", "イベント呼出しを再試行。冪等制約で重複防止"],
            ["宛先抽出", "child_guardiansとguardian_contactsを園ID付きで検索", "対象なしはskipped相当の監査集計"],
            ["キュー作成", "LINE／メールを別レコードで同一トランザクション作成", "一意違反は既存通知を返す"],
            ["LINE送信", "現行Push API＋X-Line-Retry-Key。active/linked確認", "failed。401/403は設定エラー、429/5xxは再送候補"],
            ["メール送信", "選定済みプロバイダーへ送信しmessage_id保存", "failed。恒久エラーと一時エラーを分類"],
            ["結果表示", "channel別statusを画面・監査ログへ反映", "一方sentでも他方failedを隠さない"],
        ],
        [1.7, 6.1, 2.8],
    )

    add_heading(doc, "7.3 再送制御", 2)
    add_bullet(doc, "一時エラー（timeout、429、5xx）は1分、5分、30分の最大3回を候補とし、Retry-Afterを優先する。")
    add_bullet(doc, "恒久エラー（無効メール、LINE block/invalid user、認証設定不正）は自動再送せず管理者確認とする。")
    add_bullet(doc, "管理者再送は元通知を上書きせず、attempt_countと監査ログで追跡できる形にする。")
    add_bullet(doc, "LINE Retry Keyはnotification_queue.idから決定的に生成する現行方式を維持し、同一試行の重複受付を抑止する。")

    add_heading(doc, "8. 通知文面・テンプレート", 1)
    add_heading(doc, "8.1 LINE連携案内メール", 2)
    add_table(
        doc,
        ["項目", "テンプレート"],
        [
            ["件名", "【まもるバス】LINE通知の連携をお願いします"],
            ["本文", "園からLINE通知のご案内です。下のQRを読み取るか「LINEで連携する」をタップし、表示された連携メッセージを送信してください。"],
            ["期限", "この案内は YYYY/MM/DD HH:mm まで有効です。期限後は園へ再発行をご依頼ください。"],
            ["注意", "このメールを転送しないでください。心当たりがない場合は操作せず園へご連絡ください。"],
            ["代替導線", "QR画像＋同一URLのボタン。アクセシビリティ用代替テキストを設定。"],
        ],
        [2.0, 8.6],
    )

    add_heading(doc, "8.2 降車通知", 2)
    add_table(
        doc,
        ["channel", "件名／本文例"],
        [
            ["LINE", "まもるバスからのお知らせです。\n{child_display_name}さんの降車記録を {occurred_at_jst} に受け付けました。\n※本通知は記録のお知らせであり、安全確認の最終判断を代替するものではありません。"],
            ["メール件名", "【まもるバス】降車記録のお知らせ"],
            ["メール本文", "{guardian_name} 様\n{child_display_name}さんの降車記録を {occurred_at_jst} に受け付けました。\n本メールは送信専用です。内容に心当たりがない場合は園へご連絡ください。"],
        ],
        [2.0, 8.6],
    )
    add_callout(doc, "文言ルール", "「安全確認が完了しました」「必ず降車しました」等の保証表現は使用しない。システムに保存された操作記録の通知であることを明示する。", "FDEEEB", CORAL)

    add_heading(doc, "8.3 テンプレート変数", 2)
    add_table(
        doc,
        ["変数", "生成元", "加工"],
        [
            ["guardian_name", "guardian_contacts.name", "未登録時は「保護者」"],
            ["child_display_name", "children.name", "園の通知方針により姓名／名のみを決定"],
            ["occurred_at_jst", "trip_attendance.alighted_at", "JST YYYY/MM/DD HH:mm"],
            ["organization_name", "organizations.name", "HTMLエスケープ"],
            ["contact_information", "園設定", "返信先電話等。運用決定後に設定項目化"],
        ],
        [2.8, 3.8, 4.0],
    )

    add_heading(doc, "9. セキュリティ・個人情報保護", 1)
    add_table(
        doc,
        ["観点", "設計"],
        [
            ["Webhook真正性", "受信raw bodyとX-Line-SignatureをHMAC-SHA256で比較。compare_digestを維持。"],
            ["QR秘密性", "QRに個人情報を含めない。raw_tokenは128bit以上、24時間、一回限り、DBにはhashのみ。"],
            ["組織分離", "全API・JOIN・一意制約にorganization_idを含め、別園データは404で隠す。"],
            ["権限", "登録、再発行、解除、履歴閲覧はadmin。通知イベントはadmin/operatorまたは内部認証。"],
            ["ログ", "メールはマスク、LINE userIdとtokenは出力禁止。provider応答も個人情報を除外・1000文字以内。"],
            ["同意", "同意日時、取得職員、取得方法を保存。撤回時はchannelを停止し監査ログを残す。"],
            ["メール", "SPF/DKIM/DMARCを設定した送信ドメインを使用。本文へraw token文字列を直接表示しない。"],
            ["LINE", "アクセストークン／Secretは環境変数またはSecret Manager。Git、README、画面へ出さない。"],
            ["保持", "連携要求は期限後30日以内に削除またはhashを消去。通知履歴の保存期間は園規程で決定。"],
            ["監査", "登録・更新・同意・発行・連携・解除・送信・再送・失敗をaction単位で保存。"],
        ],
        [2.2, 8.4],
    )

    add_heading(doc, "9.1 監査ログaction", 2)
    add_table(
        doc,
        ["action", "resource", "detail（個人情報を除く）"],
        [
            ["guardian_contact.create/update", "guardian_contact", "child_ids、channel設定、consent有無"],
            ["line.link_request.issue/revoke", "line_link_request", "expires_at、delivery notification id"],
            ["line.contact.link/unfollow/unlink", "line_contact", "guardian_contact_id、状態"],
            ["notification.event.create", "notification", "event_key、channel数"],
            ["notification.dispatch/retry", "notification", "status、channel、error_code、attempt_count"],
        ],
        [4.0, 2.5, 4.1],
    )

    add_heading(doc, "10. 障害・例外設計", 1)
    add_table(
        doc,
        ["事象", "システム処理", "管理者／保護者への案内"],
        [
            ["案内メール送信失敗", "line_link_requestはpendingのまま、email通知failed", "管理者がメール修正後に再発行"],
            ["連携token期限切れ", "連携拒否、status=expired", "保護者へ再発行依頼を案内"],
            ["token転送・別LINEで使用", "最初の有効使用のみ成立。以後usedで拒否", "心当たりがなければ管理者が解除・再発行"],
            ["LINE unfollow", "is_active=false、LINE送信停止", "管理画面に解除表示。メールは継続"],
            ["LINE API 429/5xx", "failed、一時エラーとして再送予定", "履歴に次回予定を表示"],
            ["LINE API 401/403", "自動再送停止、設定エラー", "アクセストークン・権限を確認"],
            ["メール恒久エラー", "email failed、error_code=invalid_recipient", "メール修正。LINEがlinkedならLINEは継続"],
            ["DB重複", "一意制約で重複を抑止し既存結果を返す", "二重表示しない"],
            ["Webhook JSON不正", "400。署名不正は401", "外部へ詳細非公開、内部ログを確認"],
        ],
        [3.0, 4.2, 3.4],
    )

    add_heading(doc, "10.1 可観測性", 2)
    add_bullet(doc, "channel別のqueued/sent/failed/skipped件数、送信遅延、再送回数を日次で集計する。")
    add_bullet(doc, "Webhook署名不正、連携token不正試行、同一tokenの再使用をセキュリティ指標として集計する。")
    add_bullet(doc, "アラート候補：10分間にfailedが5件以上、LINE 401/403が1件以上、Webhook受信が24時間ゼロ。")
    add_bullet(doc, "本文・メール・LINE userId・raw tokenをメトリクスラベルや例外ログへ含めない。")

    add_heading(doc, "11. テスト仕様・受入条件", 1)
    add_callout(doc, "判定の扱い", "以下は実施予定の試験であり、現時点の合否ではない。実在保護者を使わず、専用のテスト用LINEアカウントとメールアドレスで実施する。")
    add_heading(doc, "11.1 API・単体テスト", 2)
    add_table(
        doc,
        ["No.", "試験観点", "期待結果", "状態"],
        [
            ["UT-01", "LINE希望ON・メールなし", "422で登録拒否", "未実施"],
            ["UT-02", "別園child_id", "404、データ非開示", "未実施"],
            ["UT-03", "token生成", "128bit以上、DBに平文なし、期限24h", "未実施"],
            ["UT-04", "署名不正Webhook", "401、DB変更なし", "未実施"],
            ["UT-05", "期限内pending token", "1回だけlinkedへ遷移", "未実施"],
            ["UT-06", "使用済み／期限切れtoken", "紐付け変更なし", "未実施"],
            ["UT-07", "同一Webhook再送", "webhookEventIdで1回だけ処理", "未実施"],
            ["UT-08", "同一通知イベント2回", "channelごとに1レコードのみ", "未実施"],
            ["UT-09", "LINE未連携＋メール有効", "LINE skipped、メールqueued", "未実施"],
            ["UT-10", "一方だけ送信失敗", "sent/failedを独立保持", "未実施"],
        ],
        [1.0, 4.3, 4.2, 1.1],
        8.0,
    )

    add_heading(doc, "11.2 画面・結合・運用テスト", 2)
    add_table(
        doc,
        ["No.", "試験観点", "期待結果", "主担当", "状態"],
        [
            ["IT-01", "PCメールでQR読取", "LINEトークが開き、送信後linked", "開発＋園", "未実施"],
            ["IT-02", "スマホメールでリンクタップ", "同一端末で連携完了", "開発＋園", "未実施"],
            ["IT-03", "兄弟姉妹を同一保護者へ登録", "各園児の降車通知を同一LINE／メールへ送信", "開発", "未実施"],
            ["IT-04", "unfollow後の降車", "LINE送信なし、メール継続", "開発", "未実施"],
            ["IT-05", "通知履歴・再送", "失敗channelだけ再送可能", "開発＋園", "未実施"],
            ["IT-06", "メール誤入力修正・再発行", "旧token失効、新宛先のみ有効", "園", "未実施"],
            ["IT-07", "実機・低速回線", "重複連携・重複通知なし", "開発", "未実施"],
            ["IT-08", "文言レビュー", "安全保証の誤解がなく、問い合わせ先が明確", "園責任者", "未実施"],
            ["IT-09", "同意撤回", "両channel停止、監査ログ保存", "園", "未実施"],
            ["IT-10", "本番前テスト送信", "テストアカウントのみ、結果記録・承認", "園責任者", "未実施"],
        ],
        [1.0, 4.0, 3.8, 1.2, 0.8],
        7.7,
    )
    add_body(doc, "受入条件：P0/P1不具合ゼロ、署名・組織分離・token再使用防止・通知冪等性の試験が全件合格し、園責任者が文面・同意・障害時運用を承認していること。")

    add_heading(doc, "12. リリース・移行・ロールバック", 1)
    add_heading(doc, "12.1 リリース順序", 2)
    release_num_id = create_decimal_numbering(doc)
    add_number(doc, release_num_id, "本番DBバックアップと復元確認を行う。")
    add_number(doc, release_num_id, "新規テーブル・列・制約をマイグレーションし、既存LINE宛先はguardian_contact_id=NULLで保持する。")
    add_number(doc, release_num_id, "メール配信、LINE_BASIC_ID=@408mrkbk、token pepper等の環境変数を登録する。Secret値はSecret Manager等で管理する。")
    add_number(doc, release_num_id, "バックエンドを先行リリースし、新APIをfeature flagで無効のまま疎通確認する。")
    add_number(doc, release_num_id, "フロントエンドをリリースし、管理者画面をテスト園だけ有効化する。")
    add_number(doc, release_num_id, "テスト用LINE／メールで連携・併送・再送・unfollowを確認する。")
    add_number(doc, release_num_id, "園責任者承認後、対象園で段階的に有効化する。")

    add_heading(doc, "12.2 ロールバック", 2)
    add_bullet(doc, "feature flagで新規発行と自動通知イベント生成を停止する。")
    add_bullet(doc, "既存LINE Push APIは手動通知用途として維持できるが、未紐付け宛先へは送らない。")
    add_bullet(doc, "DB列・テーブルは即時削除せず、アプリを旧版互換状態へ戻す。データ削除は別承認とバックアップ後に行う。")
    add_bullet(doc, "誤通知が疑われる場合は送信キューを停止し、園の事故・個人情報対応手順を優先する。")

    add_heading(doc, "12.3 必要な環境変数", 2)
    add_table(
        doc,
        ["変数", "用途", "取扱い"],
        [
            ["LINE_CHANNEL_ACCESS_TOKEN", "Push送信", "現行。Secret"],
            ["LINE_CHANNEL_SECRET", "Webhook署名検証", "現行。Secret"],
            ["LINE_BASIC_ID", "QR／トーク起動URL生成", "採用値：@408mrkbk。公開可能値だが設定管理"],
            ["LINE_ORGANIZATION_ID", "単一園の暫定紐付け", "現行。複数園化時は廃止検討"],
            ["LINE_LINK_TOKEN_PEPPER", "token hash強化", "追加。Secret、ローテーション手順必要"],
            ["EMAIL_PROVIDER_API_KEY", "メール送信", "追加。Secret"],
            ["EMAIL_FROM_ADDRESS", "送信元", "追加。SPF/DKIM/DMARC対象"],
            ["NOTIFICATION_FEATURE_ENABLED", "段階リリース", "追加。園別flagへの発展を検討"],
        ],
        [4.0, 4.0, 2.6],
    )

    add_heading(doc, "13. 未決事項・レビュー依頼", 1)
    add_table(
        doc,
        ["ID", "決定事項", "推奨案", "決定者", "期限／状態"],
        [
            ["D-01", "通知トリガー", "初期は降車記録成立時のみ", "園責任者", "要決定"],
            ["D-02", "LINEとメールの関係", "LINE希望者はメール必須・両方併送", "園責任者", "本設計の前提"],
            ["D-03", "メール配信事業者", "API、配信結果、DKIM対応を比較して選定", "開発責任者", "要決定"],
            ["D-04", "QR有効期限", "24時間", "園＋開発", "要承認"],
            ["D-05", "保護者1人に複数園児", "許可。園児ごと通知ON/OFF", "園責任者", "推奨"],
            ["D-06", "LINEアカウントを複数保護者で共有", "1 LINE userId = 1保護者連絡先。兄弟は同連絡先へ追加", "園＋開発", "要承認"],
            ["D-07", "保護者名・園児名の通知表示", "必要最小限。名のみ等を園方針で決定", "園＋個人情報担当", "要決定"],
            ["D-08", "通知履歴・連携要求の保存期間", "園規程・法務確認後に設定", "園＋専門家", "要決定"],
            ["D-09", "複数園LINEチャネル", "園別チャネルまたはWebhook識別子で組織解決", "開発責任者", "要設計確定"],
        ],
        [0.9, 3.4, 4.0, 1.4, 1.3],
        7.8,
    )
    add_callout(doc, "レビュー上の要点", "D-02の解釈が依頼意図と異なる場合は、①LINEのみ、②メールのみ、③LINE＋メール併送、④QR案内だけメール送信、のどれかを確定し、通知生成・画面・文面を更新する。", "FFF6DF", AMBER)

    add_heading(doc, "14. 変更影響・トレーサビリティ", 1)
    add_table(
        doc,
        ["成果物／領域", "影響", "対応"],
        [
            ["業務フロー", "保護者登録、同意、案内、解除、障害時連絡を追加", "運用マニュアル更新"],
            ["画面一覧・画面設計", "SC-LN-01～04追加", "React設定画面へ実装"],
            ["DB設計", "3テーブル追加、2テーブル拡張", "版管理マイグレーション"],
            ["API設計", "保護者・連携・通知イベント・再送を追加", "FastAPI＋OpenAPI更新"],
            ["バリデーション", "メール、同意、園ID、token期限・状態", "APIと画面の両方で実装"],
            ["テストケース", "署名、冪等性、組織分離、実機、併送、再送", "自動＋結合＋運用試験"],
            ["監査ログ", "登録、同意、連携、解除、送信、再送", "action定義追加"],
            ["運用マニュアル", "テストアカウント、再発行、誤登録、障害時手順", "本番前に整備"],
            ["GitHub Issue / PR", "機能をDB/API/UI/運用に分割", "各PRに本設計の該当節を紐付け"],
            ["リリース手順", "Secret、migration、feature flag、段階有効化", "チェックリスト化"],
        ],
        [2.7, 4.3, 3.6],
    )

    add_heading(doc, "15. 実装根拠", 1)
    add_table(
        doc,
        ["根拠", "確認内容"],
        [
            ["README.md（2026-07-26時点）", "LINEはAPI実装済み、画面操作・同意・再送・到達確認・実運用は未完成"],
            ["園運用決定（2026-07-27）", "LINE公式アカウント「バナナ幼稚園」（LINE ID：@408mrkbk）を採用"],
            ["backend/main.py:35-37", "LINE_CHANNEL_ACCESS_TOKEN、LINE_CHANNEL_SECRET、LINE_ORGANIZATION_ID"],
            ["backend/main.py:139-162", "notification_queue、line_contactsの現行データモデル"],
            ["backend/main.py:920-980", "通知キュー、LINE Push、署名付きWebhook、宛先一覧API"],
            ["src/App.tsx:331-380", "園児画面・設定画面の現行管理者UI。通知設定は未実装"],
            ["backend/test_main.py", "LINE通知・Webhook専用の自動テストは現時点で確認できず、追加が必要"],
        ],
        [4.0, 6.6],
    )
    add_body(doc, "本書はコードレビューに基づく実装前詳細設計であり、LINE Developers設定、メール事業者、実機動作、園運用、個人情報保護の最終判断は人による確認と承認を必要とする。")

    # Core properties
    doc.core_properties.title = "LINE通知（QR連携）詳細設計書"
    doc.core_properties.subject = "保護者メールアドレス登録・LINE連携・LINE／メール併送"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "LINE, QR, メール, 通知, 詳細設計, まもるバス"
    doc.core_properties.comments = "現行実装を根拠に作成した実装前レビュー用設計書"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()


