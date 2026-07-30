from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

TARGET = Path(r'C:\home\SV\2025\15_まもるバス\まもるバス_基本設計書_20260730.docx')
OUTPUT = Path(r'C:\home\SV\2025\15_まもるバス\まもるバス_基本設計書_20260730_画面オブジェクト追加.docx')
BACKUP = Path(r'C:\home\SV\2025\15_まもるバス\まもるバス_基本設計書_20260730_BACKUP_before_screen_objects.docx')
FONT = 'BIZ UDゴシック'
DATE = '2026年7月30日'

SCREEN_OBJECT_ROWS = [
    ['SCR-01', 'OBJ-LOGIN-01', '入力欄', '職員ID', 'ログイン画面中央', 'ログイン時に職員IDを入力する。', '空欄・数値以外・存在しないIDの表示を確認。'],
    ['SCR-01', 'OBJ-LOGIN-02', '入力欄', 'PIN', 'ログイン画面中央', 'PINを入力する。', '伏字、桁数、誤PIN時の表示を確認。'],
    ['SCR-01', 'OBJ-LOGIN-03', 'ボタン', 'ログイン', 'PIN入力欄の下', '認証APIを呼び出し、成功時にホームへ遷移する。', '成功・失敗・通信不可を取得。'],
    ['SCR-01', 'OBJ-LOGIN-04', '導線', '管理者PINを復旧する', 'ログイン画面下部', '緊急復旧画面または入力欄を開く。', '緊急時のみ使用する注意表示を確認。'],
    ['SCR-02', 'OBJ-HOME-01', 'カード', '本日の送迎状況', 'ホーム上部', '園名、日付、送迎件数、運行中、完了、未確認数を表示する。', '0件、未確認あり、運行中ありを取得。'],
    ['SCR-02', 'OBJ-HOME-02', 'ボタン', '運行を開く', 'ホーム主要操作', '進行中送迎の確認または便選択へ進む。', '進行中あり/なしの遷移を取得。'],
    ['SCR-02', 'OBJ-HOME-03', '通知バー', '未同期の記録', 'ホームまたは共通上部', '端末内保留件数と同期ボタンを表示する。', 'オフライン保留ありの状態を取得。'],
    ['SCR-02', 'OBJ-HOME-04', 'ナビ', '下部ナビ', '画面下部', 'ホーム、運行、園児、記録、LINE、設定を切り替える。', '選択中表示と各遷移を取得。'],
    ['SCR-03', 'OBJ-OP-01', 'カード/ボタン', '便・車両選択', '運行開始前', '便名、方向、車両を選んで送迎開始する。', '往路、帰り、車両未設定を取得。'],
    ['SCR-03', 'OBJ-OP-02', '一覧', '当日名簿', '運行中', '対象園児、乗車済み、降車済み、未確認を表示する。', '欠席・臨時乗車・確認済み除外不可を取得。'],
    ['SCR-03', 'OBJ-OP-03', 'ボタン', 'QRを読み取る', '乗降操作エリア', 'Scannerモーダルを開き、QRで乗車/降車を記録する。', 'カメラ許可、読取成功、手入力を取得。'],
    ['SCR-03', 'OBJ-OP-04', 'ボタン', 'QRなしで記録', '園児行ごと', '園児本人を目視確認した場合だけ手動乗降を記録する。', '確認ダイアログと監査対象文言を取得。'],
    ['SCR-03', 'OBJ-OP-05', '表示', '人数照合', '運行中上部/完了前', '対象、乗車、降車、未確認人数を表示する。', '未降車あり/全員降車を取得。'],
    ['SCR-03', 'OBJ-OP-06', '表示/ボタン', '帰りの完了前チェック ACTIVE', '全員降車後', '車内撮影と完了操作を有効化する。', 'ACTIVEは安全確認完了ではない表示を確認。'],
    ['SCR-03', 'OBJ-OP-07', 'モーダル', '車内撮影', '動画撮影時', '5秒後停止可能、30秒自動終了、動画アップロードを行う。', '5秒未満不可、30秒自動終了、失敗時表示を取得。'],
    ['SCR-03', 'OBJ-OP-08', '入力/ボタン', '第三者確認', '完了前チェック付近', '別職員IDとPINで第三者確認を記録する。', '任意機能であること、誤PIN時表示を取得。'],
    ['SCR-03', 'OBJ-OP-09', 'ボタン', '送迎を中止して選び直す', '送迎開始直後', '乗降・車内撮影前のみ中止し便選択へ戻る。', '中止可/不可の表示を取得。'],
    ['SCR-04', 'OBJ-CH-01', '一覧', '園児一覧', '園児画面', '園児名、クラス、QR文字列を表示する。', 'QR文字列の見え方と長文折返しを取得。'],
    ['SCR-04', 'OBJ-CH-02', '入力欄', '園児名・クラス・QR文字列', '園児登録/編集フォーム', '園児情報を登録・更新する。', '必須、重複QR、更新成功を取得。'],
    ['SCR-04', 'OBJ-CH-03', 'ボタン', '登録・保存', '園児フォーム下部', '園児登録または更新APIを呼ぶ。', 'admin以外の扱いを取得。'],
    ['SCR-05', 'OBJ-REC-01', '検索/一覧', '送迎記録一覧', '記録画面', '期間・状態に応じて送迎履歴を表示する。', '運行中、完了、中止除外を取得。'],
    ['SCR-05', 'OBJ-REC-02', 'ボタン', '進行中送迎を再開', '記録一覧行', '運行中の送迎を運行画面へ戻す。', '便名・車両名が分かる状態を取得。'],
    ['SCR-05', 'OBJ-REC-03', '詳細', '乗降・安全確認・GPS', '記録詳細', '園児別乗降、担当者、安全確認、GPSを表示する。', '未記録、GPSなしを取得。'],
    ['SCR-05', 'OBJ-REC-04', '詳細/ボタン', '動画・AI補助', '記録詳細', '動画ID、保存キー、形式、AI状態、動画を開くを表示する。', '動画あり/なし、AI未接続メッセージを取得。'],
    ['SCR-06', 'OBJ-LINE-01', 'フォーム', '保護者通知先登録', 'LINE画面上部', '保護者名、続柄、メール、対象園児、LINE希望、同意を登録する。', '同意必須、メール必須、対象園児未選択を取得。'],
    ['SCR-06', 'OBJ-LINE-02', '一覧', '保護者・LINE通知一覧', 'LINE画面中央', '保護者、メール、対象園児、LINE状態を表示する。', '未案内、案内済み、連携済み、解除済みを取得。'],
    ['SCR-06', 'OBJ-LINE-03', 'ボタン', 'QR案内を発行/再発行', '保護者行', '期限付きLINE連携QRとメール案内を発行する。', '初回、期限切れ、再発行を取得。'],
    ['SCR-06', 'OBJ-LINE-04', 'カード', 'QR連携案内', 'QR発行後', 'QR画像、LINE公式アカウント、連携URL、連携メッセージ、有効期限を表示する。', 'Snipping対象として最重要。QRと有効期限を取得。'],
    ['SCR-06', 'OBJ-LINE-05', '一覧/ボタン', '通知履歴・再送', 'LINE画面下部', 'LINE/メール別の送信結果、試行回数、失敗理由、再送を表示する。', '失敗、送信済み、再送不可を取得。'],
    ['SCR-07', 'OBJ-SET-01', 'フォーム', '園情報', '設定画面', '園名称を表示・更新する。', '更新成功と権限制御を取得。'],
    ['SCR-07', 'OBJ-SET-02', 'フォーム/一覧', '職員設定', '設定画面', '職員名、ロール、PIN、有効状態を管理する。', 'admin/operator/verifierを取得。'],
    ['SCR-07', 'OBJ-SET-03', 'フォーム/一覧', '車両設定', '設定画面', '車両名、ナンバー、非表示化を管理する。', '削除ではなく非表示化の状態を取得。'],
    ['SCR-07', 'OBJ-SET-04', 'フォーム/一覧', '便・通常名簿', '設定画面', '便名、方向、車両、通常名簿を管理する。', '名簿保存、便削除、車両切離しを取得。'],
]

CAPTURE_ROWS = [
    ['CAP-01', 'SCR-01', 'ログイン初期表示', '職員ID/PIN/ログイン/復旧導線が1枚で分かること。'],
    ['CAP-02', 'SCR-02', 'ホーム通常表示', '本日の件数カード、運行導線、下部ナビが分かること。'],
    ['CAP-03', 'SCR-02', '未同期あり', '未同期件数と同期ボタンが分かること。'],
    ['CAP-04', 'SCR-03', '便選択', '便・方向・車両を選ぶ画面オブジェクトが分かること。'],
    ['CAP-05', 'SCR-03', '乗車確認中', '当日名簿、QR読取、手動記録、人数照合が分かること。'],
    ['CAP-06', 'SCR-03', '降車確認中', '未降車数、降車操作、完了前チェック待ちが分かること。'],
    ['CAP-07', 'SCR-03', '完了前チェックACTIVE', 'ACTIVE、車内撮影、送迎完了条件が分かること。'],
    ['CAP-08', 'SCR-03', '車内撮影モーダル', '5秒/30秒制御、停止、アップロード中表示が分かること。'],
    ['CAP-09', 'SCR-04', '園児一覧・登録', '園児名、クラス、QR文字列、登録/保存が分かること。'],
    ['CAP-10', 'SCR-05', '記録詳細', '乗降、安全確認、GPS、動画・AI補助が分かること。'],
    ['CAP-11', 'SCR-06', '保護者通知先登録', 'メール、対象園児、LINE希望、同意チェックが分かること。'],
    ['CAP-12', 'SCR-06', 'QR連携案内', 'QR画像、連携メッセージ、URL、有効期限が分かること。'],
    ['CAP-13', 'SCR-06', '通知履歴・再送', 'LINE/メール別の状態、失敗理由、再送ボタンが分かること。'],
    ['CAP-14', 'SCR-07', '設定', '園、職員、車両、便、通常名簿の管理オブジェクトが分かること。'],
]


def set_run_font(run, size=9, bold=False, color='0F172A'):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for key in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rfonts.set(qn(key), FONT)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell(cell, text, bold=False, fill=None, align=None, size=8.0):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    if align:
        p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    hdr = OxmlElement('w:tblHeader')
    hdr.set(qn('w:val'), 'true')
    tr_pr.append(hdr)


def add_table(doc, headers, rows, widths=None, size=7.7):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    set_repeat_header(t.rows[0])
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, fill='E0F2F1', align=WD_ALIGN_PARAGRAPH.CENTER, size=8.2)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1, 2) else None
            set_cell(cells[i], v, align=align, size=size)
    if widths:
        t.autofit = False
        for row in t.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return t


def add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=11.5 if level == 2 else 15, bold=True, color='164E63' if level == 2 else '0F766E')
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=9)
    return p


def move_new_content_before(doc, start_idx, target_para_text):
    body = doc._body._element
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == target_para_text:
            target = p._p
            break
    if target is None:
        return False
    new_elements = list(body)[start_idx:]
    for el in new_elements:
        body.remove(el)
    insert_at = list(body).index(target)
    for offset, el in enumerate(new_elements):
        body.insert(insert_at + offset, el)
    return True


if not TARGET.exists():
    raise SystemExit(f'target not found: {TARGET}')
if not BACKUP.exists():
    BACKUP.write_bytes(TARGET.read_bytes())

doc = Document(TARGET)

# Revision history table.
if doc.tables:
    rev = doc.tables[0]
    exists = any('画面オブジェクト' in ''.join(cell.text for cell in row.cells) for row in rev.rows)
    if not exists:
        row = rev.add_row().cells
        vals = ['1.1', DATE, '更新', '画面一覧の補足として、Snipping Toolで取得・確認する画面オブジェクト一覧とキャプチャ観点を追加。', 'Codex']
        for cell, val in zip(row, vals):
            set_cell(cell, val, size=8.0)

body = doc._body._element
start_idx = len(body) - 1
add_heading(doc, '4.1 画面オブジェクト一覧（Snipping Tool確認対象）', 2)
add_para(doc, '本節は、画面一覧を補足する画面オブジェクト一覧である。Snipping Toolで画面を取得するときは、画面全体だけでなく、ボタン、入力欄、一覧、警告、モーダル、状態表示が読み取れる状態で取得する。画像ファイルを別途取得した場合は、各画面IDまたはキャプチャIDに対応させて貼付する。')
add_table(
    doc,
    ['画面ID', 'オブジェクトID', '種別', '画面オブジェクト', '位置', '役割', 'Snipping Tool取得・確認観点'],
    SCREEN_OBJECT_ROWS,
    [1.8, 2.8, 2.2, 4.0, 3.8, 7.0, 7.0],
    size=7.0,
)
add_heading(doc, '4.2 Snipping Toolキャプチャ観点', 2)
add_para(doc, 'キャプチャは、設計レビューで画面の見た目と操作対象を確認するために使う。個人情報、実在する園児名、保護者メールアドレス、LINEユーザーID、Secret、PINが写る場合は、取得前にテストデータへ置き換える。')
add_table(
    doc,
    ['キャプチャID', '画面ID', '取得場面', '取得時に写すべき内容'],
    CAPTURE_ROWS,
    [2.5, 2.0, 6.0, 15.5],
    size=7.6,
)

moved = move_new_content_before(doc, start_idx, '5. 画面遷移')
if not moved:
    add_para(doc, '注記: 章5の位置を特定できなかったため、本節は文末に追加した。')

doc.core_properties.modified = datetime(2026, 7, 30, 0, 0, 0)
doc.save(OUTPUT)
print('updated', OUTPUT)
print('backup', BACKUP)



