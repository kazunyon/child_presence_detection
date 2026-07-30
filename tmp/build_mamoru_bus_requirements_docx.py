from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("outputs/まもるバス_要件定義書_20260730.docx")
FONT = "BIZ UDゴシック"
TODAY = "2026年7月30日"


def set_run_font(run, size=9, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)


def set_paragraph_font(paragraph, size=9):
    for run in paragraph.runs:
        set_run_font(run, size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, fill=None, color=None, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    set_run_font(run, 9, bold, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        set_cell_shading(cell, fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    size = 15 if level == 1 else 11
    set_run_font(run, size=size, bold=True, color="0F766E" if level == 1 else "334155")
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, 9)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_run_font(run, 9)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        set_cell_text(hdr.cells[idx], header, bold=True, fill="E0F2F1", color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_page_break(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.4)
section.bottom_margin = Cm(1.2)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.header_distance = Cm(0.7)
section.footer_distance = Cm(0.7)

styles = doc.styles
styles["Normal"].font.name = FONT
styles["Normal"].font.size = Pt(9)
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
    styles[style_name].font.name = FONT
    styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run("まもるバス 要件定義書")
set_run_font(run, 8, color="64748B")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Confidential / SV - Mamoru Bus Requirements Definition")
set_run_font(run, 8, color="64748B")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(55)
run = title.add_run("まもるバス\n要件定義書")
set_run_font(run, 24, bold=True, color="0F766E")
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("送迎バス安全確認PWA / Requirements Definition")
set_run_font(run, 12, bold=True, color="334155")
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(20)
run = meta.add_run(f"作成日：{TODAY}　版数：1.0　作成対象：まもるバス")
set_run_font(run, 9)

add_page_break(doc)

add_heading(doc, "改訂履歴", 1)
add_table(
    doc,
    ["版", "日付", "区分", "内容", "作成・更新"],
    [
        ["1.0", TODAY, "新規作成", "READMEおよび実装確認結果をもとに、まもるバスの要件定義書を作成。", "Codex"],
    ],
    [1.5, 3.2, 2.2, 15.5, 4.0],
)

add_heading(doc, "目次", 1)
toc_items = [
    "1. 文書概要",
    "2. システム化の目的・背景",
    "3. 対象範囲",
    "4. 現行課題",
    "5. 利用者・権限要件",
    "6. 業務要件",
    "7. 機能要件",
    "8. 非機能要件",
    "9. 画面要件",
    "10. 外部連携・環境要件",
    "11. データ要件",
    "12. 運用・移行要件",
    "13. 未決定・要確認事項",
    "14. 用語集",
    "15. 参考・確認元",
]
add_bullets(doc, toc_items)

add_page_break(doc)

add_heading(doc, "1. 文書概要", 1)
add_table(
    doc,
    ["項目", "内容"],
    [
        ["文書名", "まもるバス 要件定義書"],
        ["対象システム", "保育園・幼稚園の送迎バスにおける園児置き去り防止を支援する安全確認PWA"],
        ["作成目的", "業務上必要な機能・非機能・運用条件を整理し、設計・実装・総合テスト・導入判断の共通基準にする。"],
        ["前提", "現時点の実装は開発途中の試作版であり、実際の送迎業務・法令上必要な装置・職員の目視確認を代替しない。"],
        ["確認元", "README.md、backend/main.py、src/App.tsx、.env.example、render.yaml、package.json"],
        ["記載ルール", "実装で確認できる事項は「実装確認」、README上の方針は「方針」、未確定事項は「要確認」、未実装事項は「未実装」と表記する。"],
    ],
    [5.0, 21.0],
)

add_heading(doc, "2. システム化の目的・背景", 1)
add_body(doc, "まもるバスは、園にあるスマートフォンと園児QRコードを活用し、送迎バスでの乗車確認、降車確認、人数照合、車内目視確認、5〜30秒の車内動画撮影、記録保存を一連の流れとして支援するシステムである。")
add_body(doc, "安全確認の最終判断は職員が行う。システムは確認漏れを減らすための補助であり、アプリの表示、動画、AI判定のいずれか単独で安全確認完了とみなさない。")
add_table(
    doc,
    ["ID", "目的", "説明", "優先度"],
    [
        ["OBJ-01", "園児の置き去り防止", "乗車・降車・車内確認を段階的に記録し、未確認のまま完了できない仕組みにする。", "最優先"],
        ["OBJ-02", "職員が迷わない操作", "ITに詳しくない職員でも、当日の送迎で必要な操作を大きなボタンと分かりやすい日本語で進められるようにする。", "高"],
        ["OBJ-03", "記録の証跡化", "誰が、いつ、どの園児・車両・便に対して確認したかを保存し、後から確認できるようにする。", "高"],
        ["OBJ-04", "園ごとのデータ分離", "複数園で利用する場合に、園児・職員・車両・送迎・通知・動画・監査ログを園単位で分離する。", "高"],
        ["OBJ-05", "保護者通知の拡張", "降車記録を保護者へLINE・メールで通知できる土台を用意する。", "中"],
        ["OBJ-06", "AI補助の将来拡張", "動画確認にAI補助を追加できる構造を用意しつつ、AI未接続時は人による再確認を促す。", "中"],
    ],
    [2.4, 4.5, 16.5, 2.4],
)

add_heading(doc, "3. 対象範囲", 1)
add_table(
    doc,
    ["区分", "対象", "内容"],
    [
        ["対象業務", "日々の送迎確認", "送迎開始、当日名簿調整、乗車確認、降車確認、人数照合、車内確認、動画撮影、完了記録を対象とする。"],
        ["対象業務", "園基本設定", "園児、職員、車両、便、通常名簿、保護者通知先の登録・変更を対象とする。"],
        ["対象業務", "記録確認", "送迎記録、乗降記録、安全確認、動画証跡、監査ログ、通知履歴の確認を対象とする。"],
        ["対象外", "車載安全装置の代替", "法令上必要な安全装置、ブザー、センサー、園の緊急対応手順の代替は対象外とする。"],
        ["対象外", "実AI判定の本番運用", "現時点ではAIプロバイダー未接続のため、子ども検出や自動安全判定は対象外とする。"],
        ["対象外", "会計・勤怠・園務全般", "送迎安全確認に直接関係しない会計、勤怠、園児台帳全般、保育記録は対象外とする。"],
    ],
    [3.0, 5.0, 18.0],
)

add_heading(doc, "4. 現行課題", 1)
add_table(
    doc,
    ["ID", "課題", "影響", "対応方針"],
    [
        ["ISS-01", "紙や口頭の確認では、未確認者・確認時刻・担当者が即時に見えにくい。", "確認漏れや引き継ぎ漏れにつながる。", "乗降と安全確認をアプリ上で時刻・担当者付きで記録する。"],
        ["ISS-02", "送迎中に欠席や臨時乗車が発生すると、通常名簿だけでは当日の実態とずれる。", "人数照合が不正確になる。", "運行中の当日名簿変更を可能にし、確認済み園児の除外は不可とする。"],
        ["ISS-03", "降車後の車内目視確認が、実施証跡として残りにくい。", "後追い確認や監査が難しい。", "車内確認記録、GPS、5〜30秒動画を保存する。"],
        ["ISS-04", "通信不能時に記録できないと、現場運用が止まる。", "紙への一時退避や後入力が必要になる。", "端末内に乗降記録を保留し、復帰後に同期する。ただし長時間圏外等は導入前試験が必要。"],
        ["ISS-05", "保護者への降車連絡は手作業だと負荷が高く、送信履歴も分散しやすい。", "問い合わせ対応に時間がかかる。", "同意済み保護者に対してLINE・メール通知をキュー化し、履歴・再送を管理する。"],
        ["ISS-06", "動画保存や通知連携は個人情報を扱うため、本番前の運用設計が不足するとリスクが高い。", "情報漏えい、保存容量不足、削除漏れにつながる。", "保存期限、暗号化、アクセス記録、バックアップ、権限、実機試験を導入前必須事項にする。"],
    ],
    [2.4, 8.2, 7.0, 8.4],
)

add_heading(doc, "5. 利用者・権限要件", 1)
add_table(
    doc,
    ["ロール", "想定利用者", "主な操作", "制限・補足"],
    [
        ["admin", "園管理者、システム管理担当", "園設定、園児・職員・車両・便・通常名簿、保護者通知先、通知履歴、監査ログ、管理者PIN復旧後の確認", "管理操作を行える。PINはハッシュ保存。"],
        ["operator", "運転担当、添乗職員", "ログイン、送迎開始、当日名簿変更、QR・手動乗降記録、車内確認、動画撮影、送迎完了、記録閲覧", "管理設定の変更は不可。"],
        ["verifier", "第三者確認担当", "第三者確認、動画AI補助要求、必要な記録確認", "管理設定の変更は不可。第三者確認は現状、完了必須条件ではない。"],
        ["保護者", "園児の保護者", "LINE連携、降車通知の受信", "管理画面の利用者ではない。通知同意・LINE希望・連携状態に基づき通知する。"],
    ],
    [3.0, 5.0, 11.5, 7.0],
)

add_table(
    doc,
    ["ID", "権限要件", "受入条件"],
    [
        ["AUTH-01", "職員は職員IDとPINでログインできること。", "ログイン成功時にJWTを発行し、以後のAPI操作にAuthorizationヘッダーを要求する。"],
        ["AUTH-02", "職員ロールは admin、operator、verifier のいずれかとすること。", "未定義ロールを登録・更新できない。"],
        ["AUTH-03", "管理機能はadminのみ利用できること。", "一般職員が園児・職員・車両・便などの管理APIを呼ぶと403になる。"],
        ["AUTH-04", "園ごとのデータを相互参照・操作できないこと。", "ログイン職員のorganization_idと異なるデータを取得・更新・動画取得できない。"],
        ["AUTH-05", "管理者PIN復旧は緊急時のみ、一時トークンと新PINで行うこと。", "復旧トークンは一度だけ使用可能で、操作は監査ログに残る。"],
        ["AUTH-06", "本番用Secret、PIN、トークンはリポジトリ・README・画面キャプチャへ記載しないこと。", "環境変数・Secret管理で保持し、平文の混入をレビュー対象にする。"],
    ],
    [2.5, 14.0, 9.0],
)

add_heading(doc, "6. 業務要件", 1)
business_rows = [
    ["BR-01", "送迎開始", "職員は担当する便・方向・車両を選択し、当日の送迎を開始できること。", "必須"],
    ["BR-02", "通常名簿引き継ぎ", "送迎開始時、選択した便の通常名簿を当日の送迎名簿へ引き継ぐこと。", "必須"],
    ["BR-03", "当日名簿変更", "欠席・臨時乗車に応じて、送迎中に当日の園児名簿を変更できること。", "必須"],
    ["BR-04", "確認済み園児保護", "すでに乗車または降車を確認した園児は、誤って当日名簿から外せないこと。", "必須"],
    ["BR-05", "乗車確認", "乗車対象の園児ごとに、QRまたは手動で乗車時刻・担当者を記録できること。", "必須"],
    ["BR-06", "降車確認", "降車対象の園児ごとに、QRまたは手動で降車時刻・担当者を記録できること。", "必須"],
    ["BR-07", "工程切替", "行き便では全員乗車確認後に降車確認へ自動的に切り替えること。", "必須"],
    ["BR-08", "人数照合", "対象人数、乗車済み、降車済み、未確認人数を工程に応じて表示すること。", "必須"],
    ["BR-09", "完了前チェック", "全員降車後に完了前チェックを開始可能な状態として表示すること。", "必須"],
    ["BR-10", "車内目視確認", "職員が座席、足元、座席下、荷物の陰、最後尾を目視確認した記録を残せること。", "必須"],
    ["BR-11", "動画証跡", "車内を5〜30秒撮影し、動画1件として保存できること。", "必須"],
    ["BR-12", "完了条件", "全員降車、車内確認記録、5〜30秒動画1件がそろうまで送迎完了できないこと。", "必須"],
    ["BR-13", "第三者確認", "必要に応じて別職員がID・PINで第三者確認を記録できること。", "任意"],
    ["BR-14", "送迎再開", "進行中の送迎をホームまたは記録から再表示・再開できること。", "必須"],
    ["BR-15", "送迎選び直し", "乗降・車内撮影前の送迎は中止し、便選択へ戻れること。", "必須"],
    ["BR-16", "過去記録確認", "期間指定で送迎記録、乗降、担当者、安全確認、GPS、動画証跡を確認できること。", "必須"],
    ["BR-17", "保護者通知", "同意済み保護者へ、降車記録をLINE・メールで通知できること。", "導入前要確認"],
    ["BR-18", "通知履歴・再送", "LINE・メールの送信結果、失敗理由、試行回数を記録し、個別再送できること。", "導入前要確認"],
    ["BR-19", "オフライン記録", "通信不能時の乗降記録は端末に保留し、通信復帰後に同期できること。", "一部実装"],
    ["BR-20", "AI補助", "動画に対してAI補助要求と結果表示ができること。ただしAI判定のみで完了させないこと。", "土台のみ"],
]
add_table(doc, ["ID", "業務要件", "内容", "優先度・状態"], business_rows, [2.3, 5.0, 17.2, 2.7])

add_heading(doc, "7. 機能要件", 1)
functional_rows = [
    ["FR-001", "ログイン", "職員IDとPINを入力し、認証成功時に職員名・ロール・JWTを取得する。", "実装確認"],
    ["FR-002", "ログアウト", "端末内の認証トークンを削除し、ログイン画面へ戻す。", "実装確認"],
    ["FR-003", "認証状態復元", "保存済みトークンがある場合、起動時に職員情報と初期データを取得する。", "実装確認"],
    ["FR-004", "園情報表示", "ホーム画面で園名、本日の日付、送迎件数、未確認件数を表示する。", "実装確認"],
    ["FR-005", "園情報更新", "管理者は園名称を更新できる。", "実装確認"],
    ["FR-006", "園児一覧", "園児名、クラス、QR文字列を一覧表示できる。", "実装確認"],
    ["FR-007", "園児登録", "管理者は園児名、クラス、QR文字列を登録できる。", "実装確認"],
    ["FR-008", "園児更新", "管理者は園児名、クラス、QR文字列を更新できる。", "実装確認"],
    ["FR-009", "職員一覧", "管理者は職員名、ロール、有効状態を一覧表示できる。", "実装確認"],
    ["FR-010", "職員登録", "管理者は職員名、ロール、PINを登録できる。", "実装確認"],
    ["FR-011", "職員更新", "管理者は職員名、ロール、PIN、有効状態を更新できる。", "実装確認"],
    ["FR-012", "車両一覧", "管理者は車両名、ナンバー、有効状態を確認できる。", "実装確認"],
    ["FR-013", "車両登録・更新", "管理者は車両名、ナンバーを登録・更新できる。", "実装確認"],
    ["FR-014", "車両非表示化", "車両削除時は過去記録を壊さないよう非表示化する。", "実装確認"],
    ["FR-015", "便一覧", "管理者は便名、方向、紐づく車両、通常名簿を確認できる。", "実装確認"],
    ["FR-016", "便登録・更新", "管理者は便名、方向、車両、通常名簿を登録・更新できる。", "実装確認"],
    ["FR-017", "便非表示化", "便削除時は過去記録を壊さないよう非表示化する。", "実装確認"],
    ["FR-018", "送迎開始", "便を選択すると送迎記録を作成し、通常名簿から園児別乗降レコードを作成する。", "実装確認"],
    ["FR-019", "帰り便開始", "帰り便は通常名簿の園児を乗車済みとして開始し、降車確認へ進む。", "実装確認"],
    ["FR-020", "行き便開始", "行き便は乗車確認から開始し、全員乗車後に降車確認へ切り替える。", "実装確認"],
    ["FR-021", "当日名簿更新", "運行中の送迎に対し、園児IDリストで当日名簿を更新できる。", "実装確認"],
    ["FR-022", "確認済み園児除外防止", "乗車済み・降車済みの園児を当日名簿更新で外そうとした場合は拒否する。", "実装確認"],
    ["FR-023", "QR乗降記録", "園児QR文字列と乗車・降車の区分を受け取り、該当園児の時刻・担当者を記録する。", "実装確認"],
    ["FR-024", "対象外QR警告", "当日名簿にいない園児QRを読み取った場合、当日名簿変更を案内し記録しない。", "実装確認"],
    ["FR-025", "手動乗降記録", "カメラが使えない場合、園児ごとの手動操作で乗降を記録できる。", "実装確認"],
    ["FR-026", "手動記録の監査", "手動記録はQRなし操作として監査ログに残す。", "実装確認"],
    ["FR-027", "未確認人数算出", "乗車工程は対象人数-乗車済み、降車工程は乗車済み-降車済みで未確認を算出する。", "実装確認"],
    ["FR-028", "送迎中止", "乗降・車内撮影前の送迎のみ中止できる。", "実装確認"],
    ["FR-029", "ステータス取得", "送迎IDごとに、便名、車両名、人数、園児別乗降状態、安全確認状態を取得できる。", "実装確認"],
    ["FR-030", "記録詳細取得", "送迎IDごとに、乗降詳細、安全確認、動画証跡を取得できる。", "実装確認"],
    ["FR-031", "車内確認記録", "完了前チェックで職員ID、確認種別、QR、GPSを記録できる。", "実装確認"],
    ["FR-032", "GPS保存", "端末が位置情報を許可した場合、緯度・経度を安全確認記録へ保存する。", "実装確認"],
    ["FR-033", "動画撮影制御", "5秒経過後に停止可能、30秒で自動停止する車内撮影UIを提供する。", "実装確認"],
    ["FR-034", "動画アップロード", "動画ファイル、撮影秒数、Content-Typeを受け取り、100MB以下の動画として保存する。", "実装確認"],
    ["FR-035", "動画メタデータ保存", "動画ID、元ファイル名、storage_key、保存先、形式、AI状態を保存する。", "実装確認"],
    ["FR-036", "動画取得", "認証済み職員が同一園の動画のみダウンロードできる。", "実装確認"],
    ["FR-037", "動画保存先検証", "保存先がUPLOAD_DIR配下でない場合は取得を拒否する。", "実装確認"],
    ["FR-038", "AI補助要求", "動画IDに対してAI補助要求を行い、状態とメッセージを記録する。", "土台のみ"],
    ["FR-039", "送迎完了", "全員降車、車内確認、動画1件を満たす場合のみ完了できる。", "実装確認"],
    ["FR-040", "強制完了", "管理者は必要時に送迎を強制完了できる。", "実装確認"],
    ["FR-041", "第三者確認", "別職員ID・PINで第三者確認を記録できる。", "実装確認"],
    ["FR-042", "進行中一覧", "運行中の送迎を一覧し、選択して再開できる。", "実装確認"],
    ["FR-043", "過去記録検索", "送迎記録を期間や状態で検索できる。", "実装確認"],
    ["FR-044", "監査ログ検索", "管理者は操作、対象種別、対象ID、文字列、期間で監査ログを検索できる。", "実装確認"],
    ["FR-045", "オフライン保留", "通信不能時の乗降イベントを端末ローカルに保存する。", "一部実装"],
    ["FR-046", "オフライン同期", "通信復帰時にclient_event_id付きイベントを同期し、重複登録を防止する。", "一部実装"],
    ["FR-047", "保護者連絡先", "管理者は保護者名、メール、LINE希望、通知同意、対象園児を管理できる。", "実装確認"],
    ["FR-048", "LINE連携案内", "保護者ごとに期限付き・一回限りのLINE連携QR/リンクを発行する。", "実装確認"],
    ["FR-049", "LINE署名検証", "LINE Webhook受信時にX-Line-Signatureを検証する。", "実装確認"],
    ["FR-050", "LINE連携", "保護者が「連携 <token>」を送信すると、期限・未使用・同意状態を検証してLINEユーザーIDを紐づける。", "実装確認"],
    ["FR-051", "LINE解除", "管理者操作またはunfollowイベントでLINE連携を解除・無効化できる。", "実装確認"],
    ["FR-052", "降車通知キュー", "降車イベントごとに保護者・チャネル単位で冪等な通知キューを作成する。", "実装確認"],
    ["FR-053", "通知送信", "NOTIFICATION_FEATURE_ENABLED=trueの場合、LINE・メールの通知送信を実行する。", "実装確認・実運用要確認"],
    ["FR-054", "通知再送", "失敗した通知を管理者が個別に再送できる。ただし機密リンク付きLINE案内は再発行とする。", "実装確認"],
    ["FR-055", "メール通知連携", "EMAIL_WEBHOOK_URLへ件名、本文、連携URL、QR画像Data URL等を送信できる。", "実装確認・事業者要確認"],
    ["FR-056", "PWA提供", "スマートフォンブラウザから利用できるReact PWAとして提供する。", "実装確認"],
    ["FR-057", "公開画面接続", "GitHub Pagesの画面からVITE_API_BASE_URLでRender上のFastAPIへ接続する。", "実装確認"],
    ["FR-058", "PDF出力", "送迎記録をPDFで出力できる。", "未実装"],
    ["FR-059", "CSV出力", "送迎記録をCSVで出力できる。", "未実装"],
    ["FR-060", "未確認アラーム", "時間超過時に未確認警告や管理者通知を出す。", "未実装"],
]
add_table(doc, ["ID", "機能", "要件", "状態"], functional_rows, [2.2, 4.2, 17.0, 3.0])

add_heading(doc, "8. 非機能要件", 1)
nonfunctional_rows = [
    ["NFR-001", "安全性", "システムは安全確認を補助するものであり、職員の目視確認と園の緊急手順を必須とする。", "必須"],
    ["NFR-002", "完了制御", "未降車、車内確認未実施、動画未保存の状態では送迎完了できないこと。", "必須"],
    ["NFR-003", "可用性", "送迎時間帯に利用できること。Render、GitHub Pages、DB、端末通信の障害時運用は別途定義する。", "要確認"],
    ["NFR-004", "性能", "通常操作のAPI応答は現場で待たされない水準とする。具体値は負荷試験後に確定する。", "要確認"],
    ["NFR-005", "スマートフォン対応", "園にあるスマートフォンで片手操作しやすい大きなボタンと日本語中心の画面にする。", "必須"],
    ["NFR-006", "QR読取", "ブラウザのBarcodeDetectorまたはjsQRでQRを読み取れること。カメラ不可時は手入力を許容する。", "必須"],
    ["NFR-007", "認証", "JWTを利用し、トークン期限を環境変数で設定できること。", "必須"],
    ["NFR-008", "PIN保護", "PINはPBKDF2-SHA256でハッシュ化して保存すること。", "必須"],
    ["NFR-009", "権限制御", "管理機能、通知送信、動画判定などはロールに応じて制御すること。", "必須"],
    ["NFR-010", "テナント分離", "全主要データはorganization_idで分離すること。", "必須"],
    ["NFR-011", "個人情報保護", "園児名、保護者連絡先、LINEユーザーID、動画等は個人情報として扱うこと。", "必須"],
    ["NFR-012", "動画アクセス制御", "動画は認証済み、同一園の職員だけが取得できること。", "必須"],
    ["NFR-013", "動画保存", "Render Persistent Diskまたは本番用ストレージへ保存し、再デプロイで消失しない設計にすること。", "導入前必須"],
    ["NFR-014", "動画容量", "アップロード動画は100MB以下とする。保存容量・保存期限・削除手順は本番前に決めること。", "導入前必須"],
    ["NFR-015", "監査性", "ログイン、乗降、手動操作、車内確認、通知、動画、管理操作を監査可能にすること。", "必須"],
    ["NFR-016", "通知冪等性", "同じ降車イベントで同じ保護者・チャネルに重複通知しないこと。", "必須"],
    ["NFR-017", "通信障害耐性", "短時間の通信断では端末内保留と再同期で乗降記録を扱えること。", "一部実装"],
    ["NFR-018", "データバックアップ", "DB、動画、通知履歴、監査ログのバックアップ方式を定義すること。", "未実装"],
    ["NFR-019", "障害復旧", "Render、DB、ストレージ、LINE/メール障害時の復旧手順を定義すること。", "未実装"],
    ["NFR-020", "多要素認証", "管理者向けMFAやトークン失効を検討すること。", "未実装"],
    ["NFR-021", "DBマイグレーション", "本番運用に耐えるDBマイグレーション手順を整備すること。", "未実装"],
    ["NFR-022", "脆弱性診断", "本番前に認証、権限、動画取得、Webhook、Secret管理の診断を行うこと。", "要確認"],
    ["NFR-023", "負荷試験", "送迎集中時間、複数端末、動画アップロード時の負荷試験を行うこと。", "要確認"],
    ["NFR-024", "法令・運用適合", "園の安全管理責任者、関係機関、専門家と運用方法を確認すること。", "導入前必須"],
    ["NFR-025", "アクセシビリティ", "大きな文字・ボタン、日本語中心、確認状態の明確な表示にすること。", "必須"],
]
add_table(doc, ["ID", "分類", "非機能要件", "状態"], nonfunctional_rows, [2.4, 4.0, 17.5, 3.0])

add_heading(doc, "9. 画面要件", 1)
add_table(
    doc,
    ["画面", "利用者", "主な表示・操作", "補足"],
    [
        ["ログイン", "全職員", "職員ID、PIN、管理者PIN復旧導線", "認証成功後にホームへ遷移する。"],
        ["ホーム", "全職員", "園名、本日の送迎状況、運行開始導線、未同期件数", "進行中送迎の把握を目的とする。"],
        ["運行", "operator/admin", "便選択、当日名簿、QR読取、手動乗降、人数照合、車内撮影、第三者確認、完了", "送迎中の主要画面。"],
        ["園児", "admin中心", "園児一覧、登録・編集、QR文字列確認", "職員が読み取り対象を確認できる。"],
        ["記録", "admin/operator/verifier", "送迎記録、進行中再開、乗降詳細、安全確認、動画・AI補助", "記録閲覧と再開導線を持つ。"],
        ["LINE", "admin", "保護者通知先、LINE連携QR、通知同意、通知履歴、再送", "実運用前にテストアカウントで確認する。"],
        ["設定", "admin", "園、職員、車両、便、通常名簿", "管理操作はAPIでも権限制御する。"],
    ],
    [3.0, 3.0, 15.0, 5.0],
)

add_heading(doc, "10. 外部連携・環境要件", 1)
add_table(
    doc,
    ["区分", "項目", "要件・設定"],
    [
        ["フロントエンド", "GitHub Pages", "React PWAを配信する。APIはVITE_API_BASE_URLで指定する。"],
        ["バックエンド", "Render / FastAPI", "Renderサービス名例はmamoru-bus-api。healthCheckPathは/health。"],
        ["データベース", "PostgreSQL / SQLite", "本番はDATABASE_URLで接続。開発既定はSQLite。"],
        ["動画保存", "UPLOAD_DIR", "Renderでは/var/data/mamoru-bus-uploads。Persistent Disk容量はrender.yaml上1GB。"],
        ["通知", "LINE Messaging API", "LINE_CHANNEL_ACCESS_TOKEN、LINE_CHANNEL_SECRET、LINE_BASIC_ID、LINE_ORGANIZATION_IDを設定する。"],
        ["通知", "バナナ幼稚園", "LINE公式アカウントBasic IDは@785ntzvyとして設定されている。"],
        ["通知", "メールWebhook", "EMAIL_WEBHOOK_URL、EMAIL_FROM_ADDRESSを設定し、QR案内や降車通知を送信する。"],
        ["認証", "JWT_SECRET", "本番では十分に長いSecretを環境変数で管理する。"],
        ["認証", "TOKEN_EXPIRE_MINUTES", "既定は480分。園の運用に合わせて有効期限を調整する。"],
        ["復旧", "ADMIN_PIN_RECOVERY_TOKEN", "緊急時のみ設定し、使用後はRenderから削除して再デプロイする。"],
        ["CORS", "CORS_ORIGINS", "公開フロントエンドURLのみを許可する。render.yamlではhttps://kazunyon.github.io。"],
    ],
    [3.0, 5.0, 18.0],
)

add_heading(doc, "11. データ要件", 1)
add_table(
    doc,
    ["データ", "主な項目", "要件"],
    [
        ["Organization", "id, name, created_at", "園単位のテナント情報を保持する。"],
        ["Staff", "organization_id, name, role, password_hash, is_active", "職員、権限、PINハッシュ、有効状態を保持する。"],
        ["Child", "organization_id, name, class_name, qr_token", "園児とQR文字列を保持し、園内でQRを一意にする。"],
        ["Vehicle", "organization_id, name, plate_number, is_active", "車両名、ナンバー、非表示化状態を保持する。"],
        ["BusRoute", "organization_id, name, direction, vehicle_id, is_active", "便名、方向、車両紐づけを保持する。"],
        ["RouteChild", "route_id, child_id", "通常名簿として便と園児の関係を保持する。"],
        ["BusTrip", "organization_id, route_id, vehicle_id, direction, status, started_at, completed_at", "日々の送迎単位を保持する。"],
        ["TripAttendance", "trip_id, child_id, boarded_at, alighted_at, boarded_by, alighted_by", "園児ごとの乗降時刻と担当者を保持する。"],
        ["VehicleSafetyCheck", "organization_id, trip_id, check_type, staff_id, staff_name, qr_token, latitude, longitude, created_at", "車内確認とGPSを保持する。"],
        ["GuardianContact / ChildGuardian", "保護者名、メール、同意、LINE希望、対象園児", "保護者通知先と園児の関係を保持する。"],
        ["LineLinkRequest / LineContact", "token_hash, status, expires_at, line_user_id", "LINE連携トークンとLINEユーザーIDの紐づけを保持する。"],
        ["NotificationQueue", "event_key, channel, recipient, status, attempt_count, provider_response", "通知キュー、履歴、再送状態を保持する。"],
        ["AuditLog", "actor_id, action, resource_type, resource_id, detail, created_at", "操作証跡を保持する。"],
        ["SyncEvent", "client_event_id, outcome", "オフライン同期の重複防止と結果を保持する。"],
        ["VideoEvidence", "trip_id, uploaded_by, file_name, storage_key, content_type, ai_status, ai_result", "動画証跡とAI補助状態を保持する。"],
    ],
    [4.0, 10.5, 12.0],
)

add_heading(doc, "12. 運用・移行要件", 1)
add_table(
    doc,
    ["ID", "要件", "詳細"],
    [
        ["OP-01", "導入前データ登録", "園、職員、車両、便、通常名簿、園児QR、保護者通知先を本番運用前に登録・確認する。"],
        ["OP-02", "QR配布・貼付", "園児QRを読み取りやすい形で準備し、紛失・破損時の再発行手順を決める。"],
        ["OP-03", "職員教育", "乗車、降車、人数照合、車内目視確認、動画撮影、異常時対応を職員へ説明する。"],
        ["OP-04", "通信断運用", "オフライン保留が使える範囲と、長時間圏外・端末故障時の紙運用を決める。"],
        ["OP-05", "動画保存運用", "保存期限、削除権限、容量監視、バックアップ、暗号化、閲覧手順を決める。"],
        ["OP-06", "通知運用", "テストLINE・メールで連携、降車通知、片側失敗、再送、友だち解除を確認してから自動配信を有効化する。"],
        ["OP-07", "復旧運用", "管理者PIN復旧トークンは緊急時のみ設定し、使用後に削除する。"],
        ["OP-08", "監査運用", "事故・問い合わせ・運用確認時に参照する監査ログ項目と保管期間を決める。"],
        ["OP-09", "移行", "既存の園児名簿・職員・車両・便データの投入方法を決める。CSV取込は現時点では未実装のため手入力または別手順とする。"],
        ["OP-10", "本番判定", "実機、実回線、実園での操作確認、負荷試験、脆弱性診断、運用承認後に本番利用可否を判断する。"],
    ],
    [2.4, 7.0, 17.0],
)

add_heading(doc, "13. 未決定・要確認事項", 1)
open_rows = [
    ["OPEN-01", "本番利用可否", "現時点では試作版であり、実送迎業務に使用できる状態かは未承認。園責任者・関係機関・専門家の確認が必要。"],
    ["OPEN-02", "動画保存方式", "Render Persistent Diskを使うか、S3/R2等の外部ストレージへ移行するか未決定。"],
    ["OPEN-03", "動画保存期限", "動画の保存期間、削除手順、容量監視、バックアップ、暗号化の方式が未決定。"],
    ["OPEN-04", "通知本番設定", "LINE Developers本番チャネル、メール配信事業者、実在保護者への送信承認が未確認。"],
    ["OPEN-05", "第三者確認の扱い", "第三者確認を送迎完了の必須条件にするか、任意記録のままにするか運用判断が必要。"],
    ["OPEN-06", "本人確認強化", "第三者確認の本人性、運転担当との分離、多要素認証、トークン失効の方式が未決定。"],
    ["OPEN-07", "未確認アラーム", "時間超過時の警告、管理者通知、通知先、しきい値が未実装・未決定。"],
    ["OPEN-08", "AIプロバイダー", "実際の子ども検出・動画判定を行うAIプロバイダー、精度基準、責任分界が未決定。"],
    ["OPEN-09", "帳票出力", "PDF・CSV出力の帳票様式、出力対象、保存期間が未実装・未決定。"],
    ["OPEN-10", "バックアップ・DR", "DB、動画、通知、監査ログのバックアップと障害復旧手順が未実装・未決定。"],
    ["OPEN-11", "負荷・セキュリティ試験", "複数端末、動画アップロード、Webhook、権限境界の本格試験が未実施。"],
    ["OPEN-12", "README差分", "READMEに記載のAPIパスと実装APIの完全一致は継続確認が必要。現コードではヘルスチェックは /health。"],
]
add_table(doc, ["ID", "事項", "内容"], open_rows, [2.5, 6.0, 18.0])

add_heading(doc, "14. 用語集", 1)
add_table(
    doc,
    ["用語", "意味"],
    [
        ["まもるバス", "送迎バスでの園児置き去り防止を支援するPWA。"],
        ["PWA", "スマートフォンブラウザからアプリのように利用できるWebアプリ。"],
        ["園", "organization_idで表現される利用単位。園児・職員・記録等を分離する。"],
        ["便", "バスルート。便名、方向、車両、通常名簿を持つ。"],
        ["通常名簿", "便ごとに通常乗車する園児の一覧。"],
        ["当日名簿", "送迎開始後、その日の欠席・臨時乗車を反映した園児一覧。"],
        ["乗車確認", "園児がバスに乗ったことをQRまたは手動で記録する操作。"],
        ["降車確認", "園児がバスから降りたことをQRまたは手動で記録する操作。"],
        ["完了前チェック", "全員降車後に車内目視確認・動画撮影へ進める状態。ACTIVEは安全確認完了ではない。"],
        ["車内確認", "職員が車内最後尾まで目視確認し、必要に応じてGPS付きで記録する操作。"],
        ["動画証跡", "車内確認時に撮影・保存する5〜30秒の動画。"],
        ["第三者確認", "運転担当等とは別の職員がID・PINで任意に残す確認記録。"],
        ["監査ログ", "誰が、いつ、何を操作したかを確認するための記録。"],
        ["LINE連携", "保護者通知先とLINEユーザーIDを期限付きQR/リンクで紐づける仕組み。"],
        ["AI補助", "動画確認を支援する将来拡張。現時点では人による再確認を促す土台。"],
    ],
    [5.0, 21.0],
)

add_heading(doc, "15. 参考・確認元", 1)
add_table(
    doc,
    ["確認元", "確認内容"],
    [
        ["README.md", "システム目的、実装状況、運用方針、安全上の考え方、環境変数、今後の優先開発。"],
        ["backend/main.py", "FastAPIルート、DBモデル、認証、権限制御、送迎、通知、LINE Webhook、監査、同期、動画API。"],
        ["src/App.tsx", "画面構成、利用者操作、QR読取、オフライン保留、動画撮影、通知設定画面。"],
        [".env.example", "本番・連携用環境変数の例。"],
        ["render.yaml", "Renderサービス、DB、Persistent Disk、環境変数設定。"],
        ["package.json", "React、TypeScript、Vite、Tailwind CSS、PWA、jsQR等の利用技術。"],
    ],
    [6.0, 20.0],
)

doc.core_properties.title = "まもるバス 要件定義書"
doc.core_properties.subject = "送迎バス安全確認PWA 要件定義"
doc.core_properties.author = "Codex"
doc.core_properties.comments = "BIZ UDゴシック 9pt / A4 landscape"
doc.core_properties.created = datetime(2026, 7, 30, 0, 0, 0)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT.resolve())

