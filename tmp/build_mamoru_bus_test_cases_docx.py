from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "BIZ UDゴシック"
OUT = Path(r"C:\home\github\child_presence_detection\outputs\まもるバス_総合テスト仕様書_200件_01_20260729.docx")


@dataclass
class Case:
    chapter: int
    category: str
    item: str
    precondition: str
    steps: str
    expected: str


def set_run_font(run, size=9, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, color=None):
    cell.text = ""
    for i, part in enumerate(str(text).split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(part)
        set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    width = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_paragraph(doc, text="", style=None, size=9, bold=False, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, widths, header_fill="1F4E79"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        shade_cell(cell, header_fill)
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(255, 255, 255))
    repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx])
            set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT)
    return table


def s(*items):
    return "\n".join(f"{i + 1}. {item}" for i, item in enumerate(items))


def build_cases():
    cases: list[Case] = []

    def add(ch, cat, item, pre, steps, expected):
        cases.append(Case(ch, cat, item, pre, steps, expected))

    # 1. 起動・環境・認証 10
    for item, pre, steps, exp in [
        ("フロントエンド初期表示", "テスト環境でフロントエンドを起動済み。API URL は VITE_API_BASE_URL または開発既定値を使用する。", s("ブラウザでまもるバスのURLを開く", "初期画面のタイトルとログインフォームを確認する"), "送迎バス安全確認、まもるバス、職員ID、PIN、ログインボタンが表示され、画面描画エラーが出ない。"),
        ("バックエンドヘルスチェック", "FastAPI を起動済み。", s("GET /health を実行する", "HTTPステータスとレスポンス本文を確認する"), "HTTP 200、{\"status\":\"ok\"} を返す。README上の /api/health 記載は現行実装との差異として要確認。"),
        ("正常ログイン", "初期データ: 田中 先生 ID=1 PIN=1234、管理者 ID=3 PIN=admin1234 が有効。", s("職員IDに 1 を入力する", "PINに 1234 を入力してログインする"), "POST /api/auth/login が 200 を返し、JWT、職員名、role=operator が保存され、ホームへ遷移する。"),
        ("管理者ログイン", "管理者 ID=3、PIN=admin1234 が有効。", s("職員IDに 3 を入力する", "PINに admin1234 を入力してログインする", "設定タブを開く"), "role=admin としてログインし、設定・LINE管理に必要な管理操作ボタンが表示される。"),
        ("誤PINログイン拒否", "有効な職員IDが存在する。", s("職員IDに 1 を入力する", "PINに wrong-pin を入力してログインする"), "HTTP 401、メッセージ「職員IDまたはPINが正しくありません」。トークンは保存されずログイン後画面へ遷移しない。"),
        ("存在しない職員ID", "職員ID 9999 は未登録。", s("職員IDに 9999 を入力する", "任意のPINでログインする"), "HTTP 401。職員の存在有無を推測できる詳細情報を返さない。"),
        ("無効職員のログイン拒否", "管理者が職員を is_active=false に変更済み。", s("無効化した職員IDでログインする", "正しいPINを入力する"), "HTTP 401。無効職員は認証できず、監査ログや画面に成功扱いで残らない。"),
        ("期限切れトークン", "TOKEN_EXPIRE_MINUTES を短くしたテスト環境または期限切れJWTを準備。", s("期限切れJWTで GET /api/auth/me を呼び出す", "画面を再読み込みする"), "APIは401を返し、画面はログアウト状態へ戻る。メッセージは再ログインを促す。"),
        ("ログアウト", "任意の職員でログイン済み。", s("画面右上の職員名ボタンを押す", "localStorage の mamoru-bus-token を確認する"), "トークンが削除され、ログイン画面へ戻る。送迎中データはサーバー側で勝手に完了・削除されない。"),
        ("初期データ取得失敗", "ログイン成功後に /api/bus-routes、/api/vehicles、/api/children のいずれかを500にする。", s("ログイン後の初期表示を行う", "画面メッセージを確認する"), "「一部の初期データを取得できませんでした。運行画面を開き直してください。」が表示される。"),
    ]:
        add(1, "起動・認証", item, pre, steps, exp)

    # 2. ホーム・ダッシュボード 8
    for item, pre, steps, exp in [
        ("本日の送迎集計", "JST当日の運行中1件、完了1件、中止1件を登録済み。", s("ホームを開く", "本日の便、運行中、未確認の数を確認する"), "中止を除外し、本日の便=2、運行中=1、未確認は運行中の boarded-alighted 合計で表示される。"),
        ("園名表示", "園情報の name を「バナナ幼稚園」に設定済み。", s("ホームを表示する", "先頭カードの園名を確認する"), "organization_name が表示され、別園の名称は混在しない。"),
        ("未確認強調", "未確認数が1以上の運行中送迎が存在する。", s("ホームを表示する", "未確認メトリクスの色と数値を確認する"), "未確認が赤系表示になり、0件の場合は通常色へ戻る。"),
        ("運行画面遷移", "ログイン済み、登録済み便が存在する。", s("ホームの「運行画面を開く」を押す", "バス選択または進行中送迎表示を確認する"), "進行中送迎がなければバス選択、あれば該当送迎を再開表示する。"),
        ("下部ナビゲーション", "ログイン済み。", s("ホーム、運行、園児、記録、LINE、設定を順に押す", "各画面の見出しを確認する"), "選択した画面へ遷移し、アクティブ表示が切り替わる。権限不足画面は専用メッセージを出す。"),
        ("未同期ボタン表示", "localStorage に mamoru-bus-offline-events を1件保存済み。", s("ログインしてホームを開く", "未同期の記録ボタンを確認する"), "「未同期の記録 1 件 - 同期する」が表示される。同期成功後は0件になり非表示になる。"),
        ("日付境界", "UTC保存の started_at がJST当日境界の前後にあるデータを準備。", s("GET /api/dashboard を実行する", "date と対象件数を確認する"), "JST 0:00から翌0:00未満だけを本日として集計する。"),
        ("API接続障害表示", "ログイン後、/api/dashboard を通信失敗にする。", s("画面を表示する", "メッセージ領域を確認する"), "操作不能な例外で画面全体が崩れず、再操作可能な状態を維持する。"),
    ]:
        add(2, "ホーム", item, pre, steps, exp)

    # 3. 設定・マスタ 24
    settings_cases = [
        ("園名更新", "園情報", "管理者でログイン。現在名=デモ園。", s("設定を開く", "園情報に バナナ幼稚園 を入力", "保存を押す"), "PUT /api/organization が成功し、ホーム・設定の園名が更新される。監査ログ organization.update が残る。"),
        ("重複園名拒否", "園情報", "別organizationに同名が存在するテストDB。", s("同じ園名で保存する"), "HTTP 409「この園名は登録済みです」。既存園名は変更されない。"),
        ("園児登録", "園児", "管理者でログイン。", s("園児画面で名前=あおい ちゃん、クラス=年中、QR=child-aoi を入力", "登録する"), "POST /api/children が201。園児一覧とバス名簿選択に表示される。"),
        ("園児名必須", "園児", "管理者でログイン。", s("名前を空欄、QR=child-empty-name で登録する"), "422または画面必須制御で登録不可。空名レコードは作成されない。"),
        ("園児QR重複拒否", "園児", "QR=child-sakura の園児が登録済み。", s("別園児に同じQRを設定して登録する"), "HTTP 409「このQRコードは登録済みです」。既存園児は上書きされない。"),
        ("園児QR更新", "園児", "園児 さくら ちゃん が登録済み。", s("QRを child-sakura-new に変更する", "一覧と運行QR読取で確認する"), "PUT /api/children/{id} が成功し、新QRだけで乗降記録できる。旧QRは404になる。"),
        ("職員登録", "職員", "管理者でログイン。", s("職員名=鈴木 先生、権限=運転担当、PIN=operator123 を登録"), "POST /api/staff が201。PINは平文で返らず、ログイン可能になる。"),
        ("職員PIN最小長", "職員", "管理者でログイン。", s("新規職員PINに 1234 を入力して登録する"), "新規登録は min_length=8 により422。短いPINの新規作成はできない。"),
        ("職員ロール変更", "職員", "運転担当職員が登録済み。", s("設定でロールを verifier に変更する", "該当職員でログインする"), "role が verifier になり、管理画面操作は403、第三者確認は利用できる。"),
        ("最後の管理者保護", "職員", "有効な管理者が1名のみ。", s("最後の管理者を運転担当に変更する", "または無効化する"), "HTTP 409「最後の管理者は変更・無効化できません」。"),
        ("自分自身の無効化拒否", "職員", "管理者本人でログイン。", s("自分の職員行で無効化を押す"), "HTTP 409「自分自身は無効化できません」。ログイン継続。"),
        ("車両登録", "車両", "管理者でログイン。", s("車両名=1号車、ナンバー=品川 500 あ 1234 を登録"), "POST /api/vehicles が201。便登録の車両選択肢へ表示される。"),
        ("車両名必須", "車両", "管理者でログイン。", s("車両名を空欄、ナンバーのみで登録する"), "422または画面必須制御で登録不可。"),
        ("車両編集", "車両", "1号車が登録済み。", s("車両名=1号車A、ナンバーを更新する"), "PUT /api/vehicles/{id} が成功し、運行選択と記録詳細の今後表示に反映される。"),
        ("車両削除は非表示", "車両", "完了済み送迎で使った車両が存在する。", s("該当車両を削除する", "記録詳細を開く"), "車両は is_active=false となり今後の選択肢から消える。過去記録の車両名は保持される。"),
        ("同名車両の復元", "車両", "削除済みの 1号車 が存在する。", s("車両名=1号車 で再登録する"), "新規IDではなく削除済み車両が復元され、plate_number が更新される。"),
        ("便登録", "便", "車両と園児が登録済み。", s("便名=ひまわり園 送迎便、方向=帰り、車両=2号車で登録", "通常名簿を選ぶ"), "POST /api/bus-routes が201。便一覧と運行画面のバス選択に表示される。"),
        ("存在しない車両を便に指定", "便", "vehicle_id=9999 は存在しない。", s("APIで vehicle_id=9999 の便登録を行う"), "HTTP 404「車両が見つかりません」。便は作成されない。"),
        ("通常名簿保存", "通常名簿", "便と園児3名が登録済み。", s("通常名簿で2名を選択", "通常名簿を保存"), "route_children が選択2名に置換され、運行開始時の対象園児が2名になる。"),
        ("通常名簿重複ID排除", "通常名簿", "園児ID=1が登録済み。", s("APIで child_ids=[1,1,1] を送信する"), "重複は1件に正規化され、同一園児の名簿行が複数作成されない。"),
        ("便削除は非表示", "便", "過去記録で使った便が存在する。", s("便を削除する", "過去記録を開く"), "便は is_active=false、通常名簿は削除、今後の運行選択から消える。過去記録の便名は保持される。"),
        ("便名編集", "便", "便が登録済み。", s("便名を 2号車 帰り に変更する"), "PUT /api/bus-routes/{id} が成功し、運行選択の表示順と名称に反映される。"),
        ("設定権限拒否", "権限", "運転担当でログイン。", s("設定画面またはAPIで園児・職員・車両・便を登録する"), "管理者専用APIは403「この操作を行う権限がありません」。画面でも管理変更を許可しない。"),
        ("設定初期読み込み", "設定", "管理者でログイン。", s("設定画面を開く", "園情報、職員、車両、便、通常名簿を確認する"), "GET /api/organization、/api/staff、/api/vehicles、/api/bus-routes、/api/children の内容が園単位で表示される。"),
    ]
    for row in settings_cases:
        add(3, row[1], row[0], row[2], row[3], row[4])

    # 4. 送迎開始・中止・再開 16
    trip_start_cases = [
        ("帰り便開始", "送迎開始", "帰り便の通常名簿に2名登録済み。", s("運行画面を開く", "該当便を選択する"), "BusTrip が status=運行中 で作成され、通常名簿2名が boarded_at 設定済みで開始する。画面は降車モード。"),
        ("往路便開始", "送迎開始", "往路便の通常名簿に2名登録済み。", s("往路便を選択する"), "BusTrip が作成され、TripAttendance は未乗車状態。画面は乗車モード。"),
        ("登録便なし", "送迎開始", "有効な便が0件。", s("運行画面を開く"), "「登録済みのバスがありません。設定でバスを登録してください。」を表示し、送迎開始ボタンは出ない。"),
        ("削除済み便の開始拒否", "送迎開始", "is_active=false の便IDを準備。", s("APIで削除済み route_id を指定して POST /api/trips"), "HTTP 404「バスが見つかりません」。"),
        ("削除済み車両の開始拒否", "送迎開始", "is_active=false の vehicle_id を準備。", s("APIで削除済み vehicle_id を指定して POST /api/trips"), "HTTP 404「車両が見つかりません」。"),
        ("車両未設定便の開始", "送迎開始", "vehicle_id=null の便を登録済み。", s("該当便を選択する"), "送迎開始でき、車両名は「号車未設定」または便名由来表示になる。記録は消えない。"),
        ("進行中送迎の自動表示", "再開", "運行中の送迎が1件存在する。", s("ホームから運行画面を開く"), "GET /api/trips?status_filter=運行中 の先頭を refresh し、「進行中の送迎を表示しています」を表示する。"),
        ("記録画面から再開", "再開", "運行中の送迎が記録一覧に存在する。", s("記録タブを開く", "運行中の送迎の再開操作を行う"), "運行画面へ遷移し、車両名・便名・現在の乗降状態が復元される。"),
        ("一時保存して戻る", "再開", "運行中送迎を表示中。", s("一時保存してホームへ戻るを押す", "再度運行画面を開く"), "送迎は status=運行中 のまま保持され、再表示できる。"),
        ("開始直後の中止", "中止", "帰り便開始直後。boarded_by は通常名簿のみ、降車・安全確認なし。", s("バスを選び直す", "確認ダイアログで中止して選び直す"), "POST /api/trips/{id}/cancel が成功し status=中止。ホーム集計・記録一覧から除外される。"),
        ("乗車後の中止拒否", "中止", "往路便で1名を乗車記録済み。", s("バスを選び直す", "中止実行"), "HTTP 409「乗降または安全確認を記録した送迎は中止できません...」。運行中のまま。"),
        ("降車後の中止拒否", "中止", "帰り便で1名を降車記録済み。", s("中止実行"), "HTTP 409。実記録を含む送迎は中止できない。"),
        ("車内確認後の中止拒否", "中止", "tail_qr の VehicleSafetyCheck が存在する。", s("中止実行"), "HTTP 409。安全確認済み送迎は完了または管理者処理へ誘導する。"),
        ("完了済み送迎の中止拒否", "中止", "status=完了 の送迎。", s("POST /api/trips/{id}/cancel を実行"), "HTTP 409「この送迎は運行中ではありません」。"),
        ("複数進行中の扱い", "再開", "運行中送迎が複数存在する異常運用データ。", s("運行画面を開く", "記録画面を確認する"), "運行画面は一覧先頭を表示する。記録画面から各送迎の再開可否を確認し、運用上は重複開始を要確認とする。"),
        ("監査ログ作成", "監査", "任意の送迎を開始・中止する。", s("POST /api/trips", "POST /api/trips/{id}/cancel", "GET /api/audit-logs"), "trip.create、trip.cancel の監査ログが actor_id、resource_id、detail 付きで保存される。"),
    ]
    for row in trip_start_cases:
        add(4, row[1], row[0], row[2], row[3], row[4])

    # 5. 当日名簿変更 12
    roster_cases = [
        ("欠席園児の除外", "正常系", "帰り便開始直後。対象3名、実降車なし。", s("当日変更を開く", "欠席の園児を外す", "当日の名簿を保存"), "TripAttendance から未確認園児が削除され、対象人数と未確認数が減る。"),
        ("臨時乗車園児の追加", "正常系", "通常名簿外だが同園に登録済みの園児がいる。", s("当日変更で園児を追加", "保存後にQRで乗車または降車確認"), "TripAttendance が追加され、QR読取対象になる。"),
        ("確認済み園児の除外拒否", "異常系", "対象園児の boarded_at または alighted_at が設定済み。", s("当日変更で確認済み園児を外す", "保存する"), "HTTP 409「確認済みの園児は名簿から外せません」。名簿は変わらない。"),
        ("完了済み送迎の名簿変更拒否", "異常系", "送迎 status=完了。", s("PUT /api/trips/{id}/roster を実行する"), "HTTP 409「完了した送迎の名簿は変更できません」。"),
        ("存在しない園児ID", "異常系", "child_id=9999 は存在しない。", s("child_ids に 9999 を含めて保存"), "HTTP 404「園児が見つかりません」。"),
        ("別園園児IDの拒否", "園データ分離", "別organizationの園児IDを準備。", s("当日名簿に別園園児IDを指定する"), "HTTP 404。別園園児を名簿へ混入できない。"),
        ("重複IDの正規化", "整合性", "同一園児IDを複数含むリクエストを準備。", s("child_ids=[1,1,2] を送信する"), "名簿は1,2の2名のみになる。TripAttendance の一意制約違反が起きない。"),
        ("空名簿保存", "境界値", "運行中送迎、実記録なし。", s("全園児のチェックを外す", "保存する"), "対象人数0として保存される。運用上許可するかは要確認。完了条件は未確認0だが車内確認・動画は必要。"),
        ("往路乗車後の名簿追加", "正常系", "往路で1名乗車済み。別園児を追加する。", s("当日変更で未乗車園児を追加", "保存後に乗車QRを読む"), "追加園児は未確認として表示され、乗車後に確認済みへ変わる。"),
        ("帰り便開始直後の通常名簿扱い", "整合性", "帰り便を通常名簿2名で開始。", s("ステータスAPIを確認する"), "2名とも boarded_at が設定済み、boarded=2、alighted=0、unconfirmed=2。"),
        ("画面表示の即時更新", "UI", "名簿変更画面を開いている。", s("対象園児を変更して保存", "確認状況カードを確認"), "対象、確認済み、未確認、園児一覧が refresh 後の値に一致する。"),
        ("名簿変更監査", "監査", "当日名簿を変更する。", s("PUT /api/trips/{id}/roster", "監査ログを確認"), "trip.roster.update に child_ids が記録される。"),
    ]
    for row in roster_cases:
        add(5, row[1], row[0], row[2], row[3], row[4])

    # 6. QR・手動乗降 22
    scan_cases = [
        ("園児QRで乗車", "QR乗車", "往路送迎中、園児 child-sakura は未乗車。", s("乗車QRを読み取る", "ステータスを更新する"), "boarded_at と boarded_by が保存され、画面の確認済みが1増える。監査ログ trip.乗車 が残る。"),
        ("園児QRで降車", "QR降車", "帰り送迎中、園児 child-sakura は乗車済み未降車。", s("降車QRを読み取る"), "alighted_at と alighted_by が保存され、未確認が1減る。降車通知キュー作成処理が呼ばれる。"),
        ("未登録QR", "異常系", "qr_token=unknown-child は未登録。", s("QR文字列 unknown-child を送信"), "HTTP 404「QRコードが登録されていません」。乗降記録は作成されない。"),
        ("通常名簿外QR", "異常系", "同園園児だが当日名簿に含まれないQR。", s("QRを読み取る"), "HTTP 409「この園児は通常名簿にいません。当日の園児変更で追加してください」。"),
        ("乗車の二重読取", "重複", "対象園児はすでに boarded_at 設定済み。", s("同じ園児QRを乗車モードで再読取"), "HTTP 409「この園児はすでに乗車済みです」。時刻は上書きされない。"),
        ("降車の二重読取", "重複", "対象園児はすでに alighted_at 設定済み。", s("同じ園児QRを降車モードで再読取"), "HTTP 409「この園児はすでに降車済みです」。"),
        ("未乗車園児の降車拒否", "整合性", "往路または異常データで boarded_at が null。", s("降車モードでQRを読む"), "HTTP 409「乗車記録がないため降車できません」。"),
        ("完了済み送迎へのQR拒否", "異常系", "status=完了 の送迎。", s("POST /api/trips/{id}/scans を実行"), "HTTP 409「この送迎は完了しています」。"),
        ("手入力QR送信", "代替操作", "カメラ利用不可。QR文字列は child-sakura。", s("QR読取モーダルで手入力欄に child-sakura を入力", "送信を押す"), "カメラなしでも同じ scan API が呼ばれ、正常に乗降記録される。"),
        ("カメラ権限拒否", "UI", "ブラウザでカメラ権限を拒否。", s("QR読取を開始する"), "「カメラを利用できません。権限を許可するか、QR文字列を入力してください。」を表示する。"),
        ("QRなし乗車", "手動", "運転担当または管理者でログイン。対象園児は未乗車。", s("園児行の QRなしで乗車 を押す", "確認ダイアログで実行"), "boarded_by に「職員名（QRなし）」が保存され、園児行にQRなしバッジが表示される。"),
        ("QRなし降車", "手動", "対象園児は乗車済み未降車。", s("QRなしで降車を押す", "確認して実行"), "alighted_by に「職員名（QRなし）」が保存され、降車通知作成も通常降車と同じく行われる。"),
        ("手動操作の権限拒否", "権限", "verifier 権限でログイン。", s("POST /api/trips/{id}/manual-attendance を実行"), "HTTP 403。第三者確認者は手動乗降を登録できない。"),
        ("手動対象外園児", "異常系", "当日名簿にない園児ID。", s("manual-attendance に対象外 child_id を送信"), "HTTP 404「この園児は当日の名簿にいません」。"),
        ("手動乗車二重登録", "重複", "対象園児は手動またはQRで乗車済み。", s("再度 QRなしで乗車"), "HTTP 409「この園児はすでに乗車済みです」。"),
        ("手動降車二重登録", "重複", "対象園児は降車済み。", s("再度 QRなしで降車"), "HTTP 409「この園児はすでに降車済みです」。"),
        ("手動未乗車降車拒否", "整合性", "boarded_at が null。", s("QRなしで降車を実行"), "HTTP 409「乗車記録がないため降車できません」。"),
        ("読み取り後モーダル閉鎖", "UI", "QRモーダル表示中。", s("有効QRを読み取る"), "読取後にモーダルが閉じ、メッセージ「乗車を記録しました」または「降車を記録しました」が出る。"),
        ("QR空文字送信抑止", "入力制御", "QRモーダル表示中。", s("手入力欄を空欄または空白だけにする", "送信を押す"), "onRead は呼ばれず、空QRのAPI送信は行われない。"),
        ("担当者名記録", "監査", "佐藤 先生以外で乗降操作。", s("乗降を記録する", "記録詳細を開く"), "boarded_by/alighted_by に操作職員名が表示され、監査ログ actor_id と一致する。"),
        ("降車通知との連動", "通知", "対象園児に通知同意済み保護者が紐づく。", s("降車QRを記録する", "通知履歴を確認する"), "child.alighted.v1 の通知キューが保護者単位・チャネル単位で作成される。"),
        ("別園QRの拒否", "園データ分離", "別organizationの child-sakura QR が存在。", s("自園ログインで別園QRを読む"), "自園 organization_id で検索するため404または対象外。別園の園児名は返さない。"),
    ]
    for row in scan_cases:
        add(6, row[1], row[0], row[2], row[3], row[4])

    # 7. 人数照合・工程遷移・完了条件 16
    flow_cases = [
        ("往路開始時の人数", "人数照合", "往路、対象2名、未乗車。", s("送迎開始直後の確認状況を見る"), "確認済み0、対象2、未確認2。mode は乗車。"),
        ("往路全員乗車後の降車切替", "工程遷移", "往路、対象2名。", s("2名全員の乗車QRを読む", "画面の見出しを確認"), "modeForTrip により降車モードへ切替。見出しは帰りの送迎、ボタンは降車QRを読み取る。"),
        ("帰り開始時の人数", "人数照合", "帰り便、通常名簿2名。", s("帰り便を開始する"), "boarded=2、alighted=0、unconfirmed=2。"),
        ("降車途中の未確認", "人数照合", "帰り便2名のうち1名降車済み。", s("確認状況カードを見る"), "確認済み1、対象2、未確認1。未降車警告が表示される。"),
        ("全員降車後ACTIVE", "工程遷移", "帰り便2名が全員降車済み。", s("確認状況を更新する"), "帰りの完了前チェックが ACTIVE になり、車内撮影ボタンが有効になる。"),
        ("未降車時の撮影ボタン無効", "入力制御", "unconfirmed > 0。", s("完了前チェックを見る", "車内撮影して送迎を完了するボタンを確認"), "ボタンは disabled。安全確認・完了操作を開始できない。"),
        ("未降車完了拒否", "完了条件", "未降車1名、tail_qrと動画があっても未完了。", s("POST /api/trips/{id}/complete を実行"), "HTTP 409「未降車の園児がいるため完了できません」。"),
        ("車内確認なし完了拒否", "完了条件", "全員降車、動画あり、tail_qrなし。", s("完了APIを実行"), "HTTP 409「車内確認が必要です」。"),
        ("動画なし完了拒否", "完了条件", "全員降車、tail_qrあり、動画なし。", s("完了APIを実行"), "HTTP 409「5秒以上の車内撮影が必要です」。"),
        ("第三者確認なし完了許可", "完了条件", "全員降車、tail_qrあり、動画1件あり、third_partyなし。", s("完了APIを実行"), "status=完了。第三者確認は任意であり必須条件ではない。"),
        ("完了後の状態更新", "完了", "完了条件を満たす送迎。", s("画面から車内撮影完了後に完了する", "ホームに戻る"), "trip.status=完了、completed_at が設定され、ホーム集計の完了件数が増える。"),
        ("完了済みの再完了", "異常系", "status=完了 の送迎。", s("再度 /api/trips/{id}/complete を実行"), "実装上は trip_for_org 後の条件評価となるため挙動要確認。二重完了防止の明示仕様が必要。"),
        ("対象0名の完了前チェック", "境界値", "当日名簿0名。", s("送迎状態を確認する", "完了操作条件を確認"), "unconfirmed=0 だが、車内確認と5〜30秒動画は引き続き必要。運用許可は要確認。"),
        ("画面とAPIの人数一致", "整合性", "複数園児の乗降記録が混在。", s("GET /api/trips/{id}/status を取得", "画面の3指標と比較"), "APIの boarded/alighted/unconfirmed と画面の確認済み/対象/未確認が一致する。"),
        ("記録順序に依存しない集計", "整合性", "手動とQRが混在。", s("乗車・降車を異なる順で登録", "ステータスを確認"), "各園児の boarded_at/alighted_at の有無だけで集計し、登録方法に依存しない。"),
        ("ACTIVEの意味表示", "UI", "全員降車済み。", s("完了前チェックの表示文を確認する"), "ACTIVE は開始可能を示すだけで、安全確認完了扱いではない説明とボタン状態になっている。"),
    ]
    for row in flow_cases:
        add(7, row[1], row[0], row[2], row[3], row[4])

    # 8. 車内撮影・動画証跡・AI補助/GPS 22
    video_cases = [
        ("撮影開始", "動画", "全員降車済みでACTIVE。MediaRecorder対応ブラウザ。", s("車内撮影して送迎を完了するを押す", "録画画面を確認"), "車内撮影（5〜30秒）画面が開き、車両名と撮影目的が表示される。"),
        ("5秒未満STOP不可", "動画", "録画開始直後。", s("録画開始から4秒以内にSTOP状態を見る"), "STOPできない、または停止しても「5秒以上撮影してください」でやり直しになる。"),
        ("5秒撮影成功", "動画", "ACTIVE状態。", s("5秒以上撮影してSTOP", "アップロード完了を待つ"), "duration_seconds=5以上で動画保存、AI補助要求、tail_qr作成、送迎完了へ進む。"),
        ("30秒自動終了", "動画", "録画中。", s("30秒まで待機する"), "30秒で自動停止し、duration_seconds は30以下で送信される。"),
        ("4秒動画API拒否", "バリデーション", "認証済み、動画ファイルを準備。", s("duration_seconds=4 で POST /api/trips/{id}/videos"), "HTTP 422「車内動画は5秒以上30秒以内で撮影してください」。"),
        ("31秒動画API拒否", "バリデーション", "認証済み、動画ファイルを準備。", s("duration_seconds=31 で動画アップロード"), "HTTP 422。同じ送迎に不正動画証跡は作成されない。"),
        ("動画形式拒否", "バリデーション", "text/plain ファイルを準備。", s("content_type=text/plain で動画APIに送る"), "HTTP 415「動画ファイルを指定してください」。"),
        ("100MB超過拒否", "境界値", "100MBを超えるvideoファイルを準備。", s("動画APIにアップロードする"), "HTTP 413「動画は100MB以下にしてください」。途中ファイルは削除される。"),
        ("保存キー作成", "動画", "有効な動画をアップロード。", s("レスポンスとDBを確認する"), "VideoEvidence に organization_id/trip_id/uploaded_by/file_name/storage_key/content_type が保存され、storage_key は org/UUID.拡張子。"),
        ("動画一覧表示", "記録詳細", "送迎に動画1件あり。", s("記録詳細の動画・AI補助を開く"), "動画ID、storage_key、storage_path、content_type、AI状態、AIメッセージが表示される。"),
        ("動画ダウンロード", "動画取得", "同園の動画ファイルがUPLOAD_DIRに存在。", s("記録詳細の動画を開く", "GET /api/videos/{id}/download を確認"), "認証付きで動画ファイルを返す。video.download 監査ログが残る。"),
        ("別園動画取得拒否", "園データ分離", "別organizationの動画IDを準備。", s("自園トークンで GET /api/videos/{id}/download"), "HTTP 404「動画が見つかりません」。保存キーやパスを漏らさない。"),
        ("保存先改ざん拒否", "セキュリティ", "storage_key が UPLOAD_DIR 外へ解決される異常DBを準備。", s("動画ダウンロードAPIを実行"), "HTTP 400「動画の保存先が不正です」。"),
        ("ファイル欠損時", "異常系", "VideoEvidence はあるが実ファイルが存在しない。", s("動画ダウンロードAPIを実行"), "HTTP 404「動画ファイルが見つかりません。再デプロイ等で削除された可能性があります」。"),
        ("AI補助要求", "AI", "動画1件が保存済み。", s("POST /api/videos/{id}/analyze", "レスポンスを確認"), "ai_status=needs_human_review、職員による再確認を促す ai_result が保存される。"),
        ("AI未接続表示", "AI", "analyze 実行済み。", s("記録詳細を開く"), "AIの最終判断ではなく、人による目視確認が必要である旨を表示する。"),
        ("AI権限", "権限", "admin/operator/verifier それぞれでログイン。", s("動画AI補助要求を実行"), "3ロールは実行可能。未認証または無効職員は401。"),
        ("位置情報あり車内確認", "GPS", "ブラウザで位置情報許可。", s("車内撮影完了時に位置情報を許可", "記録詳細を確認"), "VehicleSafetyCheck に latitude/longitude が保存される。"),
        ("位置情報なし車内確認", "GPS", "位置情報を拒否またはタイムアウト。", s("車内撮影を完了する"), "「位置情報なしで記録します」を表示し、latitude/longitudeなしでtail_qrを保存できる。"),
        ("撮影キャンセル", "UI", "録画画面を表示。", s("閉じるを押す"), "動画、tail_qr、完了処理は作成されず、送迎は運行中のまま。"),
        ("アップロードタイムアウト", "異常系", "動画API応答を120秒超に遅延。", s("撮影完了後アップロードを待つ"), "「動画の保存に時間がかかっています...」を表示し、再撮影可能な状態へ戻る。"),
        ("複数動画", "境界値", "同一送迎に動画を複数アップロード。", s("動画を2件保存", "ステータスと記録詳細を確認"), "video_evidence_count=2、latest_video_id は最新動画。完了条件は1件以上で満たす。"),
    ]
    for row in video_cases:
        add(8, row[1], row[0], row[2], row[3], row[4])

    # 9. 第三者確認・強制完了 10
    approval_cases = [
        ("第三者確認正常", "第三者確認", "運転担当で送迎中。別の verifier 職員が有効。", s("第三者職員IDとPINを入力", "確認を記録する"), "VehicleSafetyCheck check_type=third_party が作成され、third_party_confirmed=true。"),
        ("管理者による第三者確認", "第三者確認", "別の admin 職員が有効。", s("admin ID/PINで第三者確認"), "verifierまたはadminは第三者確認者として受け付けられる。"),
        ("誤PIN拒否", "異常系", "verifier IDは存在。", s("誤PINで第三者確認を実行"), "HTTP 401「第三者確認者の認証に失敗しました」。記録は作成されない。"),
        ("operatorは第三者確認不可", "権限", "operator 職員を確認者に指定。", s("第三者確認を実行"), "HTTP 401。verifier/admin 以外は確認者にできない。"),
        ("本人確認拒否", "職務分離", "乗車担当者名と同名の verifier/admin を指定。", s("第三者確認を実行"), "HTTP 409「運転担当者本人は第三者確認できません」。"),
        ("無効職員拒否", "異常系", "is_active=false の verifier。", s("第三者確認を実行"), "HTTP 401。無効職員では記録できない。"),
        ("第三者確認は完了必須でない", "仕様確認", "third_party なし、完了条件3点は満たす。", s("完了APIを実行"), "送迎は完了する。第三者確認は任意機能として扱う。"),
        ("強制完了正常", "管理者処理", "未降車が残る stranded trip。管理者でログイン。", s("POST /api/trips/{id}/force-complete", "監査ログを確認"), "status=完了、forced=true。監査ログに unconfirmed/boarded/alighted が残る。"),
        ("強制完了権限拒否", "権限", "operatorでログイン。", s("force-complete APIを実行"), "HTTP 403。管理者以外は強制完了できない。"),
        ("完了済み強制完了拒否", "異常系", "status=完了 の送迎。", s("force-complete APIを実行"), "HTTP 409「この送迎はすでに完了しています」。"),
    ]
    for row in approval_cases:
        add(9, row[1], row[0], row[2], row[3], row[4])

    # 10. 過去記録・監査ログ 12
    record_cases = [
        ("記録一覧表示", "記録", "完了済み送迎と運行中送迎が存在。", s("記録タブを開く", "一覧の便名・車両名・状態を確認"), "GET /api/trips が中止を除外し、最大200件を開始日時降順で返す。"),
        ("期間検索 from_at", "検索", "複数日の送迎記録。", s("from_at に当日0:00を指定して取得"), "指定以降の送迎だけ表示される。JST/UTC変換の仕様は要確認。"),
        ("期間検索 to_at", "検索", "複数日の送迎記録。", s("to_at を指定して取得"), "指定以前の送迎だけ表示される。"),
        ("状態絞込", "検索", "運行中・完了が混在。", s("status_filter=運行中 でGET /api/trips"), "運行中のみ返る。中止は常に除外される。"),
        ("記録詳細", "詳細", "乗降、安全確認、動画がある完了送迎。", s("GET /api/trips/{id}/record", "画面で詳細を確認"), "trip、attendance、safety_checks、videos が返り、担当者名・時刻・GPS・動画情報を確認できる。"),
        ("削除済み車両名保持", "履歴", "車両削除後の過去送迎。", s("記録詳細を開く"), "過去記録の vehicle_id に紐づく車両名を表示し、今後選択肢から消えていても履歴を壊さない。"),
        ("削除済み便名保持", "履歴", "便削除後の過去送迎。", s("記録詳細を開く"), "過去記録の route_id に紐づく便名を表示する。"),
        ("監査ログ一覧", "監査", "管理者でログイン。複数操作済み。", s("GET /api/audit-logs を実行"), "organization_id で絞られ、created_at降順、既定100件で返る。"),
        ("監査ログ条件検索", "監査", "trip.complete の監査ログが存在。", s("action=trip.complete で検索", "query_text でも検索"), "指定条件に一致する監査ログのみ返る。limit は最大500に丸められる。"),
        ("一般職員の監査ログ拒否", "権限", "operatorでログイン。", s("GET /api/audit-logs を実行"), "HTTP 403。監査ログは管理者のみ。"),
        ("別園記録アクセス拒否", "園データ分離", "別organizationの trip_id を準備。", s("自園トークンで /api/trips/{id}/record を実行"), "HTTP 404「運行便が見つかりません」。"),
        ("記録の実施前結果欄", "帳票運用", "本仕様書をテスト実施前に配布。", s("結果列を確認する"), "結果列は空欄。テスト実施後にOK/NG/保留等を記入できる。"),
    ]
    for row in record_cases:
        add(10, row[1], row[0], row[2], row[3], row[4])

    # 11. オフライン同期・同時操作 10
    offline_cases = [
        ("オフライン乗降保留", "オフライン", "送迎中にブラウザをoffline状態にする。", s("QRを読み取る", "メッセージとlocalStorageを確認"), "mamoru-bus-offline-events に client_event_id/trip_id/qr_token/event_type が保存される。"),
        ("オンライン復帰同期", "同期", "未同期イベントが1件、navigator.onLine=true。", s("オンラインイベントを発火", "同期後の画面を確認"), "POST /api/sync が成功し、localStorage キューが空になる。"),
        ("同期重複防止", "同期", "同じ client_event_id を2回送る。", s("POST /api/sync を2回実行"), "2回目は outcome=already_processed。乗降時刻は二重登録されない。"),
        ("同期時の業務エラー保持", "同期", "未登録QRの保留イベント。", s("POST /api/sync を実行"), "results に rejected:QRコードが登録されていません を返し、SyncEvent に outcome を保存する。"),
        ("同期API認証", "権限", "トークンなし。", s("POST /api/sync を実行"), "HTTP 401。未認証で保留記録を同期できない。"),
        ("複数端末同時QR", "同時操作", "2端末で同じ園児を同時に乗車読取。", s("ほぼ同時にscan APIを送る"), "片方のみ成功し、もう片方は409「すでに乗車済み」。DBの一貫性を確認する。"),
        ("複数端末名簿変更競合", "同時操作", "端末A/Bで当日名簿を同時編集。", s("Aが園児追加", "Bが古い状態で保存"), "最終保存内容はAPI送信値で置換される。競合警告がないため運用上要確認。"),
        ("送迎完了と降車の競合", "同時操作", "端末Aが最後の降車、端末Bが完了操作。", s("同時にAPIを実行"), "未降車判定がDB最新状態で評価される。完了後の追加乗降は409になる。"),
        ("端末再起動後の保留", "オフライン", "未同期イベントをlocalStorageに保存済み。", s("ブラウザを再起動", "ログイン後に未同期件数を確認"), "localStorage の保留件数が復元され、同期ボタンが表示される。"),
        ("長時間圏外", "運用確認", "複数件をオフライン保留。", s("長時間後にオンライン復帰", "同期結果を確認"), "通信復帰で順次同期される想定。一部実装のため実機・複数端末で要確認。"),
    ]
    for row in offline_cases:
        add(11, row[1], row[0], row[2], row[3], row[4])

    # 12. LINE・メール通知 16
    notification_cases = [
        ("保護者登録", "保護者", "管理者でログイン。園児1名登録済み。", s("LINE画面を開く", "保護者名、メール、対象園児、同意を入力して登録"), "GuardianContact と ChildGuardian が作成され、email_normalized は小文字・前後空白除去になる。"),
        ("対象園児なし拒否", "入力制御", "管理者でログイン。", s("対象園児を選択せず保護者登録"), "HTTP 422「対象園児を1人以上選択してください」。"),
        ("メール形式不正", "入力制御", "管理者でログイン。", s("email=invalid-mail で登録"), "HTTP 422「メールアドレスの形式を確認してください」。"),
        ("同意なし通知拒否", "同意", "メールまたはLINE通知ON。", s("consent=false で登録"), "HTTP 422「通知同意を確認してください」。"),
        ("LINEはメール必須", "同意", "line_enabled=true、email_enabled=false。", s("保護者登録または更新を実行"), "HTTP 422「LINE通知を希望する場合はメール通知も有効にしてください」。"),
        ("メール重複拒否", "重複", "Parent@Example.JP が登録済み。", s("parent@example.jp で別保護者登録"), "HTTP 409。兄弟追加は既存保護者編集フローへ誘導するメッセージを確認。"),
        ("QR案内発行", "LINE連携", "line_enabled=true、consent済み保護者。EMAIL_WEBHOOK_URL設定済み。", s("QR案内を発行", "プレビューを確認"), "line_status=pending、24時間期限、@785ntzvy の LINE URL、QR PNG Data URL を返す。平文トークンはDB保存しない。"),
        ("案内メール未設定", "異常系", "EMAIL_WEBHOOK_URL 未設定。", s("QR案内を発行"), "QRは発行されるが email_delivery_status=failed。画面はメール配信設定確認を促す。"),
        ("メール変更で未使用案内失効", "LINE連携", "pending のLINE連携案内がある。", s("保護者メールを変更する"), "既存 LineLinkRequest は revoked、line_status は not_requested に戻る。"),
        ("署名付きWebhook正常", "LINE Webhook", "pending token と有効なX-Line-Signatureを準備。", s("LINEメッセージ「連携 <token>」をWebhook送信"), "guardian.line_status=linked、LineContact が紐づき、link_request.status=used。"),
        ("Webhook署名不正", "LINE Webhook", "X-Line-Signature が不正。", s("Webhookを送信"), "HTTP 401「LINE署名が不正です」。連携状態は変わらない。"),
        ("期限切れトークン", "LINE Webhook", "expires_at が過去の pending token。", s("連携メッセージを送る"), "link_request.status=expired、guardian.line_status=expired。返信は再発行依頼。"),
        ("友だち解除", "LINE Webhook", "linked の LineContact が存在。", s("unfollowイベントを送る"), "LineContact.is_active=false、guardian.line_status=unfollowed、監査ログ line.contact.unfollow。"),
        ("降車通知キュー", "通知", "保護者がメールON、LINE連携済み。", s("園児降車を記録する"), "同一 event_key で email と line の2件を作成。本文は安全確認の最終判断を代替しない注意を含む。"),
        ("降車通知冪等性", "通知", "同じ trip/child の降車通知を再実行。", s("queue_alighted_notifications を2回実行"), "2回目は既存 event_key/channel を検出し、新規通知は作成されない。"),
        ("通知再送制御", "通知", "failed 通知と sent 通知が存在。", s("failedを再送", "sentを再送", "line.link.v1を再送"), "failed は再送試行、sent は409、LINE連携案内は409で再発行フローを案内する。"),
    ]
    for row in notification_cases:
        add(12, row[1], row[0], row[2], row[3], row[4])

    # 13. 権限・園データ分離・セキュリティ・未実装確認 22
    security_cases = [
        ("管理者専用APIの403", "権限", "operatorでログイン。", s("POST /api/staff、POST /api/vehicles、GET /api/audit-logs を実行"), "管理者専用APIは403。一般職員が園設定・監査ログを操作できない。"),
        ("運転担当APIの許可範囲", "権限", "operatorでログイン。", s("送迎開始、乗降、当日名簿、完了を操作", "設定登録を操作"), "送迎系は許可、管理設定は403。"),
        ("第三者確認者の許可範囲", "権限", "verifierでログイン。", s("第三者確認、動画AI補助、手動乗降、設定登録を試す"), "第三者確認とAI補助は可。手動乗降と管理設定は403。"),
        ("未認証アクセス拒否", "認証", "Authorizationヘッダーなし。", s("GET /api/dashboard、GET /api/children を実行"), "HTTP 401。業務データは返らない。"),
        ("改ざんJWT拒否", "認証", "署名を改ざんしたJWT。", s("GET /api/auth/me を実行"), "HTTP 401「認証情報が無効です」。"),
        ("JWT組織不一致", "園データ分離", "staff_id と org が一致しないJWTを準備。", s("GET /api/auth/me を実行"), "HTTP 401「ログインし直してください」。"),
        ("園児一覧の園分離", "園データ分離", "別園にも園児が存在。", s("自園トークンで GET /api/children"), "自園 organization_id の園児のみ返る。"),
        ("車両・便一覧の園分離", "園データ分離", "別園の車両・便が存在。", s("GET /api/vehicles、GET /api/bus-routes"), "自園かつ is_active=true のみ返る。"),
        ("通知一覧の園分離", "園データ分離", "別園通知が存在。", s("GET /api/notifications"), "自園通知のみ返る。"),
        ("管理者PIN復旧正常", "復旧", "ADMIN_PIN_RECOVERY_TOKEN をRenderに設定。職員ID 3 を復旧対象。", s("復旧トークンと新PIN8文字以上を入力", "ID=3でログイン"), "PINが再設定され、auth.admin_pin_recovery 監査ログと使用済みtoken_hashが保存される。"),
        ("管理者PIN復旧トークン未設定", "復旧", "ADMIN_PIN_RECOVERY_TOKEN 未設定。", s("復旧画面から送信"), "HTTP 403「管理者PIN復旧は許可されていません」。"),
        ("管理者PIN復旧トークン再利用拒否", "復旧", "同じ復旧トークンを使用済み。", s("再度同じトークンで送信"), "HTTP 409「この復旧トークンは使用済みです。Renderから削除してください」。"),
        ("復旧新PIN最小長", "復旧", "復旧トークンは有効。", s("新PIN=1234567 で送信"), "422。8文字未満の新PINは受け付けない。"),
        ("CORS設定", "環境", "CORS_ORIGINS を既定値で起動。", s("localhost:5173/5174 からAPIアクセス", "未許可OriginからAPIアクセス"), "許可Originは利用可能。未許可Originの扱いはブラウザCORSで拒否されることを確認。"),
        ("アップロード保存先", "環境", "UPLOAD_DIR を /var/data/mamoru-bus-uploads に設定。", s("動画をアップロード", "保存先パスを確認"), "指定ディレクトリ配下に保存される。本番では容量、保存期限、暗号化、バックアップを要確認。"),
        ("DBマイグレーション", "移行", "旧schemaのSQLiteを準備。", s("アプリを起動", "追加カラムと初期組織を確認"), "既存データを消さず organization_id 等を補完する。PostgreSQL本番移行は要確認。"),
        ("PDF出力未実装確認", "未実装", "READMEでPDF・CSV出力は未実装。", s("画面とAPIルート一覧を確認", "PDF出力ボタン/APIの有無を確認"), "PDF出力機能は存在しない。導入前要件として要確認・別途実装が必要。"),
        ("CSV出力未実装確認", "未実装", "READMEでPDF・CSV出力は未実装。", s("記録画面とAPIルート一覧を確認"), "CSV出力機能は存在しない。帳票連携テストは未実装確認として扱う。"),
        ("未確認アラーム未実装確認", "未実装", "READMEで未確認アラームは未実装。", s("未降車状態で一定時間待つ", "通知・警告APIの有無を確認"), "時間超過アラームや管理者通知は動作しない。今後実装後に再テストが必要。"),
        ("多要素認証未実装確認", "未実装", "READMEで多要素認証は未実装。", s("ログインフローを確認"), "職員ID/PINのみ。MFA要件は要確認。"),
        ("バックアップ・復元未実装確認", "未実装", "READMEでバックアップ・障害復旧は未実装。", s("管理画面とAPIルート一覧を確認"), "アプリ内バックアップ/復元操作はない。Persistent Disk/DBバックアップ運用を要確認。"),
        ("法令・実運用前注意", "運用確認", "試作版の公開デモ。", s("README注意事項と画面文言を確認", "テスト報告へ残す"), "本システムは安全装置や法令上必要な装置を置き換えない。実園導入前に責任者・専門家確認を必須とする。"),
    ]
    for row in security_cases:
        add(13, row[1], row[0], row[2], row[3], row[4])

    assert len(cases) == 200, len(cases)
    return cases


def build_docx():
    cases = build_cases()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(0.45))

    styles = doc.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(9)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.bold = True

    doc.core_properties.title = "まもるバス 総合テスト仕様書"
    doc.core_properties.author = ""
    doc.core_properties.comments = "Generated from repository README, backend/main.py, src/App.tsx, backend/test_main.py, and reference accounting test document."

    add_paragraph(doc, "まもるバス 総合テスト仕様書", size=14, bold=True, color=(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "送迎バス安全確認PWA／重点確認200件版", size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "更新日：2026年7月29日", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, "更新概要：SV会計入力システム総合テスト仕様書の8列構成を参考に、まもるバスのREADME、React画面、FastAPI実装、既存単体テストを確認元として、導入前に確認すべき総合テストケース200件を作成した。")
    add_paragraph(doc, "作成目的：送迎開始、乗降、人数照合、車内目視確認、5〜30秒撮影、記録保存、保護者通知、権限、園単位データ分離、未実装事項を、手動実施しやすい粒度で確認できる形に統合する。")
    add_paragraph(doc, "実施方法：各行の確認を実施し、期待結果を満たす場合は結果欄にOK、差異がある場合はNGまたは保留を記入し、対象データ、画面、API応答、障害番号を追記する。未実装または仕様未確定の行は要確認として扱う。")

    add_paragraph(doc, "確認元", size=10, bold=True, color=(31, 78, 121))
    sources = [
        ("仕訳会計システム_総合テスト仕様書_200件_01.docx", "横向き、8列テストケース表、件数配分、結果欄の運用", "体裁・列構成の参考", "参考資料"),
        ("README.md", "実装状況、運用フロー、安全上の考え方、未実装範囲、環境変数", "対象機能・未実装確認に反映", "仕様"),
        ("backend/main.py", "DBモデル、API、JWT、ロール、バリデーション、監査ログ、通知、動画", "期待結果と異常系に反映", "実装"),
        ("src/App.tsx", "画面、ナビゲーション、QR読取、当日名簿、動画撮影、LINE設定", "操作手順と画面表示に反映", "画面"),
        ("backend/test_main.py", "中止、削除、動画、LINE通知の既存単体テスト", "重要回帰観点に反映", "テスト"),
    ]
    add_table(doc, ["確認元", "確認内容", "テストへの反映", "確認区分"], sources, [2.35, 3.2, 3.0, 1.0])

    add_paragraph(doc, "仕様・実装差異／要確認", size=10, bold=True, color=(31, 78, 121))
    diffs = [
        "READMEには /api/health、/api/public-settings の記載があるが、現行 backend/main.py のルート一覧では /health のみ確認できた。公開設定APIの要否は要確認。",
        "PDF・CSV出力、未確認アラーム、多要素認証、通常PIN変更、バックアップ・復元、実AIプロバイダー接続は未実装または土台のみ。期待結果では実装済み扱いしない。",
        "第三者確認は任意機能であり、送迎完了の必須条件ではない。運用で必須化する場合は仕様変更が必要。",
        "動画保存はファイル保存およびRender Persistent Disk想定。本番では保存期限、暗号化、容量監視、バックアップ、削除手順を要確認。",
    ]
    for diff in diffs:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("・" + diff)
        set_run_font(run)

    add_paragraph(doc, "1. 圧縮後の件数配分", size=10, bold=True, color=(31, 78, 121))
    distribution = [
        ("1", "起動・環境・認証", "10", "起動、ログイン、JWT、初期取得"),
        ("2", "ホーム・ダッシュボード", "8", "本日集計、未確認、ナビゲーション"),
        ("3", "設定・マスタ", "24", "園児、職員、車両、便、通常名簿"),
        ("4", "送迎開始・中止・再開", "16", "開始、中止条件、進行中再開"),
        ("5", "当日名簿変更", "12", "欠席、臨時乗車、確認済み保護"),
        ("6", "QR・手動乗降", "22", "QR、手入力、QRなし、重複、通知連動"),
        ("7", "人数照合・工程遷移・完了条件", "16", "ACTIVE、未降車、完了条件"),
        ("8", "車内撮影・動画証跡・AI補助/GPS", "22", "5〜30秒、保存、取得、AI土台、GPS"),
        ("9", "第三者確認・強制完了", "10", "任意確認、本人分離、管理者強制完了"),
        ("10", "過去記録・監査ログ", "12", "記録詳細、検索、履歴保持、監査"),
        ("11", "オフライン同期・同時操作", "10", "localStorage、sync、冪等、競合"),
        ("12", "LINE・メール通知", "16", "保護者、同意、QR連携、Webhook、再送"),
        ("13", "権限・園データ分離・セキュリティ・未実装確認", "22", "認証、園分離、復旧、未実装、運用確認"),
    ]
    add_table(doc, ["章", "対象領域", "件数", "圧縮の考え方"], distribution, [0.55, 2.4, 0.7, 5.8])

    add_paragraph(doc, "2. 総合テストケース一覧（200件）", size=10, bold=True, color=(31, 78, 121))
    chapter_titles = {
        1: "第1章 起動・環境・認証",
        2: "第2章 ホーム・ダッシュボード",
        3: "第3章 設定・マスタ",
        4: "第4章 送迎開始・中止・再開",
        5: "第5章 当日名簿変更",
        6: "第6章 QR・手動乗降",
        7: "第7章 人数照合・工程遷移・完了条件",
        8: "第8章 車内撮影・動画証跡・AI補助/GPS",
        9: "第9章 第三者確認・強制完了",
        10: "第10章 過去記録・監査ログ",
        11: "第11章 オフライン同期・同時操作",
        12: "第12章 LINE・メール通知",
        13: "第13章 権限・園データ分離・セキュリティ・未実装確認",
    }
    widths = [0.55, 0.65, 0.85, 1.35, 1.75, 2.0, 2.15, 0.65]
    headers = ["No", "元No", "分類", "テスト項目", "前提・確認データ", "操作手順", "期待結果", "結果"]
    global_no = 1
    for chapter in range(1, 14):
        p = add_paragraph(doc, chapter_titles[chapter], size=10, bold=True, color=(31, 78, 121))
        keep_with_next(p)
        chapter_cases = [c for c in cases if c.chapter == chapter]
        rows = []
        for seq, case in enumerate(chapter_cases, 1):
            rows.append([
                f"{chapter}-{seq:02d}",
                f"MAM-{global_no:03d}",
                case.category,
                case.item,
                case.precondition,
                case.steps,
                case.expected,
                "",
            ])
            global_no += 1
        add_table(doc, headers, rows, widths)
        if chapter != 13:
            doc.add_paragraph()

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("まもるバス 総合テスト仕様書 200件版")
    set_run_font(run, size=8, color=(89, 89, 89))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_docx()
    print(path)
